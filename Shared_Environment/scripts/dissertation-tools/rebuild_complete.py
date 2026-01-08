#!/usr/bin/env python3
"""
Complete rebuild of the Word document with properly captioned figures and tables.
This script:
1. Opens the converted Word document
2. Finds ALL figure placeholder paragraphs
3. Replaces each placeholder with the actual image + proper caption
4. Ensures tables have proper labels
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
import shutil

# CORRECTED figure mapping
FIGURE_MAPPING = {
    # Section 4 - System Design (32 figures)
    "Figure 4.1": ("rendered_diagrams_drawio/docker-architecture-logical.png", "Multi-Environment Docker Architecture (Logical View)"),
    "Figure 4.2": ("rendered_diagrams_drawio/docker-architecture-physical.png", "Multi-Environment Docker Architecture (Physical View)"),
    "Figure 4.3": ("rendered_diagrams_drawio/monitoring-architecture.png", "Monitoring Architecture Diagram"),
    "Figure 4.4": ("rendered_diagrams_drawio/mobile-architecture.png", "Mobile Application Architecture"),
    "Figure 4.5": ("rendered_diagrams_drawio/backend-api-architecture.png", "Backend API Architecture"),
    "Figure 4.6": ("rendered_diagrams_drawio/property-owner-workflow.png", "Property Owner Workflow"),
    "Figure 4.7": ("rendered_diagrams_drawio/yield-architecture.png", "Yield Tokenisation Architecture"),
    "Figure 4.8": ("rendered_diagrams_drawio/yield-tokenization-flow.png", "Complete Yield Tokenization Workflow"),
    "Figure 4.9": ("rendered_diagrams_drawio/repayment-processing-flow.png", "Repayment Processing Decision Tree"),
    "Figure 4.10": ("rendered_diagrams_drawio/pooling-mechanism-flow.png", "Pooling Mechanism Workflow"),
    "Figure 4.11": ("rendered_diagrams_drawio/default-handling-flow.png", "Default Handling Workflow"),
    "Figure 4.12": ("rendered_diagrams_drawio/storage-layout-diagram.png", "ERC-7201 Storage Layout"),
    "Figure 4.13": ("rendered_diagrams_drawio/analytics-architecture.png", "Analytics Architecture (The Graph Protocol)"),
    "Figure 4.14": ("rendered_diagrams_drawio/subgraph-entity-relationships.png", "Subgraph Entity Relationships"),
    "Figure 4.15": ("rendered_diagrams_drawio/governance-architecture.png", "Governance Architecture"),
    "Figure 4.16": ("rendered_diagrams_drawio/governance-proposal-flow.png", "Governance Proposal Flow"),
    "Figure 4.17": ("rendered_diagrams_drawio/secondary-market-architecture.png", "Secondary Market Architecture"),
    "Figure 4.18": ("rendered_diagrams_drawio/transfer-restriction-flow.png", "Transfer Restriction Flow"),
    "Figure 4.19": ("rendered_diagrams_drawio/kyc-architecture.png", "KYC Architecture"),
    "Figure 4.20": ("rendered_diagrams_drawio/kyc-workflow.png", "KYC Verification Workflow"),
    "Figure 4.21": ("rendered_diagrams_drawio/er-diagram.png", "Entity-Relationship Diagram"),
    "Figure 4.22": ("rendered_diagrams_drawio/erc1155-architecture.png", "ERC-1155 Architecture"),
    "Figure 4.23": ("rendered_diagrams_drawio/erc20-token-flow.png", "ERC-20 Token Flow"),
    "Figure 4.24": ("rendered_diagrams_drawio/property-nft-architecture.png", "Property NFT Architecture"),
    "Figure 4.25": ("rendered_diagrams_drawio/share-transfer-architecture.png", "Share Transfer Architecture"),
    "Figure 4.26": ("rendered_diagrams_drawio/use-case-diagram.png", "Use Case Diagram"),
    "Figure 4.27": ("rendered_diagrams_drawio/mobile-workflow.png", "Mobile Workflow"),
    "Figure 4.28": ("rendered_diagrams_drawio/wireframe-dashboard.png", "Dashboard Wireframe"),
    "Figure 4.29": ("rendered_diagrams_drawio/wireframe-property-registration.png", "Property Registration Wireframe"),
    "Figure 4.30": ("rendered_diagrams_drawio/wireframe-yield-agreement.png", "Yield Agreement Wireframe"),
    "Figure 4.31": ("rendered_diagrams_drawio/wireframe-mobile-dashboard.png", "Mobile Dashboard Wireframe"),
    "Figure 4.32": ("rendered_diagrams_drawio/wireframe-mobile-yield-agreement.png", "Mobile Yield Agreement Wireframe"),
    
    # Section 5 - Implementation (23 figures)
    "Figure 5.1": ("generated_screenshots/docker-desktop-dashboard.png", "Docker Desktop Container Status"),
    "Figure 5.2": ("generated_screenshots/grafana-dashboard.png", "Grafana Monitoring Dashboard"),
    "Figure 5.3": ("generated_screenshots/prometheus-targets.png", "Prometheus Targets"),
    "Figure 5.4": ("generated_screenshots/node-exporter-metrics.png", "Node-Exporter Metrics"),
    "Figure 5.5": ("generated_screenshots/multi-env-orchestration.png", "Multi-Environment Orchestration"),
    "Figure 5.6": ("generated_screenshots/backup-directory.png", "Backup Directory Structure"),
    "Figure 5.7": ("generated_screenshots/foundry-project-structure.png", "Foundry Project Structure"),
    "Figure 5.8": ("generated_screenshots/forge-test-output.png", "Forge Test Output"),
    "Figure 5.9": ("generated_screenshots/gas-report.png", "Gas Report"),
    "Figure 5.10": ("generated_screenshots/anvil-container-logs.png", "Anvil Container Logs"),
    "Figure 5.11": ("generated_screenshots/smart-contract-deployment.png", "Deployment Script Output"),
    "Figure 5.12": ("generated_screenshots/erc7201-storage-layout.png", "ERC-7201 Storage Layout"),
    "Figure 5.13": ("generated_screenshots/bytecode-size-verification.png", "Bytecode Size Verification"),
    "Figure 5.14": ("generated_screenshots/amoy-testnet-deployment.png", "Diamond Deployment to Amoy"),
    "Figure 5.15": ("generated_screenshots/polygonscan-confirmation.png", "PolygonScan Transaction Confirmation"),
    "Figure 5.16": ("generated_screenshots/diamond-test-results.png", "Diamond Architecture Test Results"),
    "Figure 5.17": ("generated_screenshots/property-registration-form.png", "Property Registration Form"),
    "Figure 5.18": ("generated_screenshots/yield-agreement-interface.png", "Yield Agreement Creation Interface"),
    "Figure 5.19": ("generated_screenshots/analytics-dashboard.png", "Analytics Dashboard"),
    "Figure 5.20": ("generated_screenshots/simulation-test-results.png", "Simulation Test Results"),
    "Figure 5.21": ("generated_screenshots/gas-comparison-report.png", "Gas Comparison Report"),
    "Figure 5.22": ("generated_screenshots/variance-tracking.png", "Variance Tracking Metrics"),
    "Figure 5.23": ("rendered_diagrams_drawio/yield-tokenization-flow.png", "Yield Tokenization Flowchart"),
    
    # Section 6 - Testing (24 figures)
    "Figure 6.1": ("generated_charts/fig_6_1_token_standard_gas_comparison.png", "Token Standard Gas Comparison (Bar Chart)"),
    "Figure 6.2": ("generated_charts/fig_6_2_batch_operation_scaling.png", "Batch Operation Scaling Curve (Line Chart)"),
    "Figure 6.3": ("generated_charts/fig_6_3_amoy_anvil_variance_heatmap.png", "Amoy vs Anvil Variance Heatmap"),
    "Figure 6.4": ("generated_charts/fig_6_4_volatile_simulation_radar.png", "Volatile Simulation Recovery Percentages (Radar Chart)"),
    "Figure 6.5": ("generated_charts/fig_6_5_diamond_overhead_boxplot.png", "Diamond Architecture Call Overhead Analysis (Box Plot)"),
    "Figure 6.6": ("generated_charts/fig_6_6_load_testing_scatter.png", "Load Testing Throughput vs Shareholder Count (Scatter Plot)"),
    "Figure 6.7": ("generated_charts/fig_6_7_gas_cost_projections.png", "Gas Cost Projections for Mainnet Scenarios (Multi-Series Line Chart)"),
    "Figure 6.8": ("generated_charts/fig_6_8_test_pass_rate_evolution.png", "Test Pass Rate Evolution (Stacked Area Chart)"),
    "Figure 6.9": ("generated_charts/architecture-comparison.png", "Architecture Comparison (Side-by-Side Diagram)"),
    "Figure 6.10": ("generated_charts/user-workflow-diagrams.png", "User Workflow Diagrams (Sequence Diagrams)"),
    "Figure 6.11": ("rendered_diagrams_drawio/load-test-system-architecture.png", "Load Test System Architecture"),
    "Figure 6.12": ("rendered_diagrams_drawio/load-test-workflow.png", "Load Test Workflow"),
    "Figure 6.13": ("rendered_diagrams_drawio/load-testing-results.png", "Load Testing Results Summary"),
    "Figure 6.14": ("rendered_diagrams_drawio/load-test-gas-scaling.png", "Load Test Gas Scaling"),
    "Figure 6.15": ("rendered_diagrams_drawio/load-test-shareholder-distribution.png", "Load Test Shareholder Distribution"),
    "Figure 6.16": ("rendered_diagrams_drawio/load-test-success-rates.png", "Load Test Success Rates"),
    "Figure 6.17": ("rendered_diagrams_drawio/load-test-latency-distribution.png", "Load Test Latency Distribution"),
    "Figure 6.18": ("rendered_diagrams_drawio/load-test-metrics-summary.png", "Load Test Metrics Summary"),
    "Figure 6.19": ("rendered_diagrams_drawio/load-test-cost-comparison.png", "Load Test Cost Comparison"),
    "Figure 6.20": ("rendered_diagrams_drawio/load-test-erc1155-efficiency.png", "Load Test ERC-1155 Efficiency"),
    "Figure 6.21": ("rendered_diagrams_drawio/load-test-token-recycling.png", "Load Test Token Recycling"),
    "Figure 6.22": ("rendered_diagrams_drawio/load-test-restriction-enforcement.png", "Load Test Restriction Enforcement"),
    "Figure 6.23": ("rendered_diagrams_drawio/load-test-restriction-overhead.png", "Load Test Restriction Overhead"),
    "Figure 6.24": ("rendered_diagrams_drawio/load-test-rq-validation.png", "Load Test Research Question Validation"),
    
    # Section 7 - Evaluation (8 figures)
    "Figure 7.1": ("survey_interview_charts/fig_demographic_overview.png", "Survey Respondent Demographics Overview (4-Panel Composite)"),
    "Figure 7.2": ("survey_interview_charts/fig_tokenisation_interest_analysis.png", "Tokenisation Interest Distribution and Cross-Tabulation"),
    "Figure 7.3": ("survey_interview_charts/fig_correlation_matrix.png", "Spearman Rank Correlation Matrix"),
    "Figure 7.4": ("survey_interview_charts/fig_interview_demographics.png", "Interview Participant Demographics (4-Panel Composite)"),
    "Figure 7.5": ("survey_interview_charts/fig_thematic_code_frequencies.png", "Theme Frequency Analysis"),
    "Figure 7.6": ("generated_charts/token-decision-framework.png", "Token Standard Decision Framework"),
    "Figure 7.7": ("generated_charts/fig_7_7_volatile_recovery_comparison.png", "Volatile Market Simulation Recovery Comparison"),
    "Figure 7.8": ("generated_charts/fig_7_8_testnet_local_comparison.png", "Testnet vs Local Performance Comparison"),
}

# Appendix D - Additional wireframes (19 figures)
APPENDIX_D = {
    "Figure D.1": ("rendered_diagrams_drawio/wireframe-analytics-dashboard.png", "Analytics Dashboard Wireframe"),
    "Figure D.2": ("rendered_diagrams_drawio/wireframe-governance.png", "Governance Dashboard Wireframe"),
    "Figure D.3": ("rendered_diagrams_drawio/wireframe-governance-proposal-detail.png", "Governance Proposal Detail Wireframe"),
    "Figure D.4": ("rendered_diagrams_drawio/wireframe-kyc.png", "KYC Verification Wireframe"),
    "Figure D.5": ("rendered_diagrams_drawio/wireframe-kyc-admin.png", "KYC Admin Wireframe"),
    "Figure D.6": ("rendered_diagrams_drawio/wireframe-marketplace.png", "Secondary Market Wireframe"),
    "Figure D.7": ("rendered_diagrams_drawio/wireframe-portfolio.png", "Investor Portfolio Wireframe"),
    "Figure D.8": ("rendered_diagrams_drawio/wireframe-properties-list.png", "Properties List Wireframe"),
    "Figure D.9": ("rendered_diagrams_drawio/wireframe-yield-agreement-detail.png", "Yield Agreement Detail Wireframe"),
    "Figure D.10": ("rendered_diagrams_drawio/wireframe-yield-agreements-list.png", "Yield Agreements List Wireframe"),
    "Figure D.11": ("rendered_diagrams_drawio/wireframe-mobile-analytics.png", "Mobile Analytics Wireframe"),
    "Figure D.12": ("rendered_diagrams_drawio/wireframe-mobile-governance.png", "Mobile Governance Wireframe"),
    "Figure D.13": ("rendered_diagrams_drawio/wireframe-mobile-kyc.png", "Mobile KYC Wireframe"),
    "Figure D.14": ("rendered_diagrams_drawio/wireframe-mobile-kyc-admin.png", "Mobile KYC Admin Wireframe"),
    "Figure D.15": ("rendered_diagrams_drawio/wireframe-mobile-marketplace.png", "Mobile Marketplace Wireframe"),
    "Figure D.16": ("rendered_diagrams_drawio/wireframe-mobile-portfolio.png", "Mobile Portfolio Wireframe"),
    "Figure D.17": ("rendered_diagrams_drawio/wireframe-mobile-properties-list.png", "Mobile Properties List Wireframe"),
    "Figure D.18": ("rendered_diagrams_drawio/wireframe-mobile-register-property.png", "Mobile Register Property Wireframe"),
    "Figure D.19": ("rendered_diagrams_drawio/wireframe-mobile-yield-agreements-list.png", "Mobile Yield Agreements List Wireframe"),
}

# Appendix E - Additional charts (5 figures)
APPENDIX_E = {
    "Figure E.1": ("survey_interview_charts/fig_cluster_analysis.png", "Respondent Cluster Analysis"),
    "Figure E.2": ("survey_interview_charts/fig_feature_importance.png", "Platform Feature Importance Rankings"),
    "Figure E.3": ("survey_interview_charts/fig_motivations_concerns.png", "Investment Motivations and Concerns"),
    "Figure E.4": ("survey_interview_charts/fig_landlord_fintech_comparison.png", "Landlord vs FinTech Expert Comparison"),
    "Figure E.5": ("survey_interview_charts/fig_likert_distributions.png", "Likert Scale Response Distributions"),
}

ALL_FIGURES = {**FIGURE_MAPPING, **APPENDIX_D, **APPENDIX_E}

def rebuild_document():
    """Rebuild the document with proper figure captions"""
    
    print("=" * 80)
    print("REBUILDING DOCUMENT WITH PROPER FIGURE CAPTIONS")
    print("=" * 80)
    
    # Step 1: Convert markdown to Word fresh
    print("\n1. Converting markdown to Word...")
    import subprocess
    result = subprocess.run([
        'pandoc', 
        'DissertationProgressFinal.md', 
        '-o', 'DissertationProgressFinal_temp.docx',
        '--reference-doc=reference.docx' if os.path.exists('reference.docx') else '--standalone'
    ], capture_output=True, text=True)
    
    if result.returncode != 0:
        # Try without reference doc
        result = subprocess.run([
            'pandoc', 
            'DissertationProgressFinal.md', 
            '-o', 'DissertationProgressFinal_temp.docx'
        ], capture_output=True, text=True)
    
    print("  Conversion complete.")
    
    # Step 2: Open the document
    print("\n2. Opening converted document...")
    doc = Document('DissertationProgressFinal_temp.docx')
    
    # Step 3: Find all figure placeholders
    print("\n3. Finding figure placeholders...")
    
    figures_inserted = 0
    figures_failed = []
    
    # We need to iterate through paragraphs and find Figure X.X patterns
    # Then insert the image BEFORE the caption
    
    para_idx = 0
    while para_idx < len(doc.paragraphs):
        para = doc.paragraphs[para_idx]
        text = para.text.strip()
        
        # Look for Figure X.X: pattern (the caption line)
        match = re.match(r'^(Figure\s+[0-9A-Z]+\.[0-9]+):\s*(.+)$', text)
        
        if match:
            fig_num = match.group(1)
            caption_text = match.group(2)
            
            # Check if this figure is in our mapping
            if fig_num in ALL_FIGURES:
                img_path, expected_caption = ALL_FIGURES[fig_num]
                
                if os.path.exists(img_path):
                    # Insert image BEFORE this paragraph
                    # We do this by:
                    # 1. Creating a new paragraph before this one
                    # 2. Adding the image to that paragraph
                    
                    # Get the paragraph's XML element
                    para_element = para._element
                    
                    # Create a new paragraph for the image
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    
                    # Create new paragraph
                    new_p = OxmlElement('w:p')
                    para_element.addprevious(new_p)
                    
                    # Now we need to add the image to this new paragraph
                    # This is complex, so let's use a different approach
                    
                    # Instead, let's modify the existing paragraph to contain the image
                    # Clear the paragraph
                    for run in para.runs:
                        run.clear()
                    
                    # Add image
                    run = para.add_run()
                    try:
                        # Determine width based on image type
                        if 'wireframe' in img_path.lower():
                            width = Inches(5.0)
                        elif 'screenshot' in img_path.lower():
                            width = Inches(5.5)
                        else:
                            width = Inches(6.0)
                        
                        run.add_picture(img_path, width=width)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add a line break and the caption
                        para.add_run('\n')
                        caption_run = para.add_run(f"{fig_num}: {caption_text}")
                        caption_run.bold = True
                        caption_run.font.size = Pt(10)
                        
                        figures_inserted += 1
                        print(f"  ✓ {fig_num}: Inserted {os.path.basename(img_path)}")
                        
                    except Exception as e:
                        figures_failed.append((fig_num, str(e)))
                        print(f"  ✗ {fig_num}: Failed - {e}")
                else:
                    figures_failed.append((fig_num, f"File not found: {img_path}"))
                    print(f"  ✗ {fig_num}: File not found - {img_path}")
        
        para_idx += 1
    
    # Step 4: Save the document
    print(f"\n4. Saving document...")
    doc.save('DissertationProgressFinal_v2.2.docx')
    
    # Clean up temp file
    if os.path.exists('DissertationProgressFinal_temp.docx'):
        os.remove('DissertationProgressFinal_temp.docx')
    
    print(f"\n{'=' * 80}")
    print(f"SUMMARY")
    print(f"{'=' * 80}")
    print(f"Figures inserted: {figures_inserted}")
    print(f"Figures failed: {len(figures_failed)}")
    
    if figures_failed:
        print("\nFailed figures:")
        for fig, error in figures_failed:
            print(f"  {fig}: {error}")
    
    return figures_inserted, figures_failed

if __name__ == "__main__":
    rebuild_document()
