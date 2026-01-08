#!/usr/bin/env python3
"""
Rebuild the Word document from the updated markdown with all figures and tables properly placed.
Version 3.0 - Comprehensive rebuild with all changes implemented.
"""

import os
import re
import subprocess
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_image_mapping():
    """Create mapping between figure references and image files."""
    mapping = {}
    
    # Section 4 - System Design (Draw.io diagrams)
    mapping['4.1'] = 'rendered_diagrams_drawio/docker-architecture-logical.png'
    mapping['4.2'] = 'rendered_diagrams_drawio/docker-architecture-physical.png'
    mapping['4.3'] = 'rendered_diagrams_drawio/github-repository.png'  # New figure
    mapping['4.4'] = 'rendered_diagrams_drawio/monitoring-architecture.png'
    mapping['4.5'] = 'rendered_diagrams_drawio/mobile-architecture.png'
    mapping['4.6'] = 'rendered_diagrams_drawio/backend-api-architecture.png'
    mapping['4.7'] = 'rendered_diagrams_drawio/property-owner-workflow.png'
    mapping['4.8'] = 'rendered_diagrams_drawio/yield-architecture.png'
    mapping['4.9'] = 'rendered_diagrams_drawio/yield-tokenization-flow.png'
    mapping['4.10'] = 'rendered_diagrams_drawio/repayment-processing-flow.png'
    mapping['4.11'] = 'rendered_diagrams_drawio/pooling-mechanism-flow.png'
    mapping['4.12'] = 'rendered_diagrams_drawio/default-handling-flow.png'
    mapping['4.13'] = 'rendered_diagrams_drawio/erc7201-storage-layout.png'
    mapping['4.14'] = 'rendered_diagrams_drawio/analytics-architecture.png'
    mapping['4.15'] = 'rendered_diagrams_drawio/subgraph-entities.png'
    mapping['4.16'] = 'rendered_diagrams_drawio/governance-architecture.png'
    mapping['4.17'] = 'rendered_diagrams_drawio/governance-proposal-flow.png'
    mapping['4.18'] = 'rendered_diagrams_drawio/secondary-market-architecture.png'
    mapping['4.19'] = 'rendered_diagrams_drawio/transfer-restriction-flow.png'
    mapping['4.20'] = 'rendered_diagrams_drawio/kyc-architecture.png'
    mapping['4.21'] = 'rendered_diagrams_drawio/kyc-workflow.png'
    mapping['4.22'] = 'rendered_diagrams_drawio/er-diagram.png'
    mapping['4.23'] = 'rendered_diagrams_drawio/erc1155-architecture.png'
    mapping['4.24'] = 'rendered_diagrams_drawio/erc20-token-flow.png'
    mapping['4.25'] = 'rendered_diagrams_drawio/property-nft-architecture.png'
    mapping['4.26'] = 'rendered_diagrams_drawio/share-transfer-architecture.png'
    mapping['4.27'] = 'rendered_diagrams_drawio/use-case-diagram.png'
    mapping['4.28'] = 'rendered_diagrams_drawio/mobile-workflow.png'
    mapping['4.29'] = 'rendered_wireframes/wireframe-dashboard.png'
    mapping['4.30'] = 'rendered_wireframes/wireframe-property-registration.png'
    mapping['4.31'] = 'rendered_wireframes/wireframe-yield-agreement.png'
    mapping['4.32'] = 'rendered_wireframes/wireframe-mobile-dashboard.png'
    mapping['4.33'] = 'rendered_wireframes/wireframe-mobile-yield-agreement.png'
    
    # Section 5 - Implementation (Screenshots)
    mapping['5.1'] = 'generated_screenshots/docker-desktop-dashboard.png'
    mapping['5.2'] = 'generated_screenshots/grafana-dashboard.png'
    mapping['5.3'] = 'generated_screenshots/prometheus-targets.png'
    mapping['5.4'] = 'generated_screenshots/node-exporter-metrics.png'
    mapping['5.5'] = 'generated_screenshots/multi-env-orchestration.png'
    mapping['5.6'] = 'generated_screenshots/backup-directory.png'
    mapping['5.7'] = 'generated_screenshots/foundry-project-structure.png'
    mapping['5.8'] = 'generated_screenshots/forge-test-output.png'
    mapping['5.9'] = 'generated_screenshots/gas-report.png'
    mapping['5.10'] = 'generated_screenshots/anvil-container-logs.png'
    mapping['5.11'] = 'generated_screenshots/smart-contract-deployment.png'
    mapping['5.12'] = 'generated_screenshots/erc7201-storage-layout.png'
    mapping['5.13'] = 'generated_screenshots/bytecode-size-verification.png'
    mapping['5.14'] = 'generated_screenshots/amoy-testnet-deployment.png'
    mapping['5.15'] = 'generated_screenshots/polygonscan-confirmation.png'
    mapping['5.16'] = 'generated_screenshots/diamond-test-results.png'
    mapping['5.17'] = 'generated_screenshots/property-registration-form.png'
    mapping['5.18'] = 'generated_screenshots/yield-agreement-interface.png'
    mapping['5.19'] = 'generated_screenshots/analytics-dashboard.png'
    mapping['5.20'] = 'generated_screenshots/simulation-test-results.png'
    mapping['5.21'] = 'generated_screenshots/gas-comparison-report.png'
    mapping['5.22'] = 'generated_screenshots/variance-tracking.png'
    mapping['5.23'] = 'rendered_diagrams_drawio/yield-tokenization-flow.png'
    
    # Section 6 - Testing (Charts)
    mapping['6.1'] = 'generated_charts/token-standard-gas-comparison.png'
    mapping['6.2'] = 'generated_charts/batch-operation-scaling.png'
    mapping['6.3'] = 'generated_charts/amoy-anvil-variance.png'
    mapping['6.4'] = 'generated_charts/volatile-simulation-recovery.png'
    mapping['6.5'] = 'generated_charts/diamond-architecture-overhead.png'
    mapping['6.6'] = 'rendered_diagrams_drawio/load-test-shareholder-distribution.png'
    mapping['6.7'] = 'generated_charts/gas-cost-projections.png'
    mapping['6.8'] = 'generated_charts/test-pass-rate-evolution.png'
    mapping['6.9'] = 'rendered_diagrams_drawio/load-test-metrics-summary.png'
    mapping['6.10'] = 'rendered_diagrams_drawio/use-case-diagram.png'
    mapping['6.11'] = 'rendered_diagrams_drawio/load-test-system-architecture.png'
    mapping['6.12'] = 'rendered_diagrams_drawio/load-test-workflow.png'
    mapping['6.13'] = 'rendered_diagrams_drawio/load-testing-results.png'
    mapping['6.14'] = 'rendered_diagrams_drawio/load-test-gas-scaling.png'
    mapping['6.15'] = 'rendered_diagrams_drawio/load-test-shareholder-distribution.png'
    mapping['6.16'] = 'rendered_diagrams_drawio/load-test-success-rates.png'
    mapping['6.17'] = 'rendered_diagrams_drawio/load-test-latency-distribution.png'
    mapping['6.18'] = 'rendered_diagrams_drawio/load-test-metrics-summary.png'
    mapping['6.19'] = 'rendered_diagrams_drawio/load-test-cost-comparison.png'
    mapping['6.20'] = 'rendered_diagrams_drawio/load-test-erc1155-efficiency.png'
    mapping['6.21'] = 'rendered_diagrams_drawio/load-test-token-recycling.png'
    mapping['6.22'] = 'rendered_diagrams_drawio/load-test-restriction-enforcement.png'
    mapping['6.23'] = 'rendered_diagrams_drawio/load-test-restriction-overhead.png'
    mapping['6.24'] = 'rendered_diagrams_drawio/load-test-rq-validation.png'
    
    # Section 7 - Evaluation (Survey/Interview charts)
    mapping['7.1'] = 'survey_charts/fig_demographics_overview.png'
    mapping['7.2'] = 'survey_charts/fig_interest_crosstab.png'
    mapping['7.3'] = 'survey_charts/fig_correlation_matrix.png'
    mapping['7.4'] = 'survey_charts/fig_cluster_analysis.png'
    mapping['7.5'] = 'survey_charts/fig_feature_importance.png'
    mapping['7.6'] = 'survey_charts/fig_motivations_concerns.png'
    mapping['7.7'] = 'interview_charts/fig_interview_demographics.png'
    mapping['7.8'] = 'interview_charts/fig_theme_frequency.png'
    mapping['7.9'] = 'interview_charts/fig_landlord_fintech_comparison.png'
    mapping['7.10'] = 'interview_charts/fig_likert_distributions.png'
    mapping['7.11'] = 'generated_charts/token-standard-decision-framework.png'
    mapping['7.12'] = 'generated_charts/volatile-market-recovery-comparison.png'
    mapping['7.13'] = 'generated_charts/testnet-local-comparison.png'
    
    # Appendix F - Wireframes
    mapping['F.1'] = 'rendered_wireframes/wireframe-analytics-dashboard.png'
    mapping['F.2'] = 'rendered_wireframes/wireframe-governance-dashboard.png'
    mapping['F.3'] = 'rendered_wireframes/wireframe-governance-proposal.png'
    mapping['F.4'] = 'rendered_wireframes/wireframe-kyc-verification.png'
    mapping['F.5'] = 'rendered_wireframes/wireframe-kyc-admin.png'
    mapping['F.6'] = 'rendered_wireframes/wireframe-secondary-market.png'
    mapping['F.7'] = 'rendered_wireframes/wireframe-investor-portfolio.png'
    mapping['F.8'] = 'rendered_wireframes/wireframe-properties-list.png'
    mapping['F.9'] = 'rendered_wireframes/wireframe-yield-agreement-detail.png'
    mapping['F.10'] = 'rendered_wireframes/wireframe-yield-agreements-list.png'
    mapping['F.11'] = 'rendered_wireframes/wireframe-mobile-analytics.png'
    mapping['F.12'] = 'rendered_wireframes/wireframe-mobile-governance.png'
    mapping['F.13'] = 'rendered_wireframes/wireframe-mobile-kyc.png'
    mapping['F.14'] = 'rendered_wireframes/wireframe-mobile-kyc-admin.png'
    mapping['F.15'] = 'rendered_wireframes/wireframe-mobile-marketplace.png'
    mapping['F.16'] = 'rendered_wireframes/wireframe-mobile-portfolio.png'
    mapping['F.17'] = 'rendered_wireframes/wireframe-mobile-properties-list.png'
    mapping['F.18'] = 'rendered_wireframes/wireframe-mobile-register-property.png'
    mapping['F.19'] = 'rendered_wireframes/wireframe-mobile-yield-agreements-list.png'
    
    return mapping

def find_image_file(figure_num, mapping):
    """Find the actual image file for a figure number."""
    if figure_num in mapping:
        path = mapping[figure_num]
        if os.path.exists(path):
            return path
    
    # Try alternative paths
    alternatives = [
        f'rendered_diagrams_drawio/',
        f'generated_screenshots/',
        f'generated_charts/',
        f'survey_charts/',
        f'interview_charts/',
        f'rendered_wireframes/',
    ]
    
    # Try to find by partial name match
    for alt_dir in alternatives:
        if os.path.exists(alt_dir):
            for f in os.listdir(alt_dir):
                if f.endswith('.png'):
                    return os.path.join(alt_dir, f)
    
    return None

def main():
    print("=" * 60)
    print("REBUILDING WORD DOCUMENT v3.0")
    print("=" * 60)
    
    # Step 1: Convert markdown to Word using pandoc
    print("\n1. Converting markdown to Word with pandoc...")
    result = subprocess.run([
        'pandoc', 
        'DissertationProgressFinal.md',
        '-o', 'DissertationProgressFinal_v2.3.docx',
        '--from', 'markdown',
        '--to', 'docx',
        '--toc',
        '--toc-depth=4'
    ], capture_output=True, text=True)
    
    if result.returncode != 0:
        print(f"Error: {result.stderr}")
        return
    
    print("   ✓ Markdown converted to Word")
    
    # Step 2: Open the document and insert images
    print("\n2. Opening document to insert images...")
    doc = Document('DissertationProgressFinal_v2.3.docx')
    
    mapping = get_image_mapping()
    
    # Count figures inserted
    figures_inserted = 0
    figures_not_found = []
    
    # Find all figure placeholders and insert images
    print("\n3. Inserting images at figure placeholders...")
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Look for figure patterns
        figure_match = re.match(r'\*\*Figure ([0-9]+\.[0-9]+|[A-Z]\.[0-9]+):', text)
        if figure_match:
            figure_num = figure_match.group(1)
            
            # Find the image file
            if figure_num in mapping:
                img_path = mapping[figure_num]
                if os.path.exists(img_path):
                    # Insert the image before this paragraph
                    run = para.insert_paragraph_before().add_run()
                    try:
                        run.add_picture(img_path, width=Inches(5.5))
                        figures_inserted += 1
                        if figures_inserted % 20 == 0:
                            print(f"   Inserted {figures_inserted} figures...")
                    except Exception as e:
                        print(f"   Warning: Could not insert {img_path}: {e}")
                        figures_not_found.append(figure_num)
                else:
                    figures_not_found.append(figure_num)
            else:
                figures_not_found.append(figure_num)
    
    print(f"\n   ✓ Inserted {figures_inserted} figures")
    
    if figures_not_found:
        print(f"\n   ⚠ {len(figures_not_found)} figures not found:")
        for fig in figures_not_found[:10]:
            print(f"      - Figure {fig}")
        if len(figures_not_found) > 10:
            print(f"      ... and {len(figures_not_found) - 10} more")
    
    # Step 3: Save the document
    print("\n4. Saving document...")
    doc.save('DissertationProgressFinal_v2.3.docx')
    
    # Get file size
    file_size = os.path.getsize('DissertationProgressFinal_v2.3.docx') / (1024 * 1024)
    
    print("\n" + "=" * 60)
    print("DOCUMENT REBUILD COMPLETE")
    print("=" * 60)
    print(f"Output: DissertationProgressFinal_v2.3.docx")
    print(f"Size: {file_size:.1f} MB")
    print(f"Figures inserted: {figures_inserted}")
    print(f"Figures not found: {len(figures_not_found)}")

if __name__ == "__main__":
    main()



