#!/usr/bin/env python3
"""
Rebuild the Word document from the updated markdown with all figures and tables properly placed.
Version 5.0 - Corrected file paths based on actual directory structure.
"""

import os
import re
import subprocess
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def get_image_mapping():
    """Create mapping between figure references and image files - CORRECTED PATHS."""
    mapping = {}
    
    # Section 4 - System Design (Draw.io diagrams)
    mapping['4.1'] = 'rendered_diagrams_drawio/docker-architecture-logical.png'
    mapping['4.2'] = 'rendered_diagrams_drawio/docker-architecture-physical.png'
    mapping['4.3'] = 'rendered_diagrams_drawio/docker-architecture-logical.png'  # GitHub repo placeholder
    mapping['4.4'] = 'rendered_diagrams_drawio/monitoring-architecture.png'
    mapping['4.5'] = 'rendered_diagrams_drawio/mobile-architecture.png'
    mapping['4.6'] = 'rendered_diagrams_drawio/backend-api-architecture.png'
    mapping['4.7'] = 'rendered_diagrams_drawio/property-owner-workflow.png'
    mapping['4.8'] = 'rendered_diagrams_drawio/yield-architecture.png'
    mapping['4.9'] = 'rendered_diagrams_drawio/yield-tokenization-flow.png'
    mapping['4.10'] = 'rendered_diagrams_drawio/repayment-processing-flow.png'
    mapping['4.11'] = 'rendered_diagrams_drawio/pooling-mechanism-flow.png'
    mapping['4.12'] = 'rendered_diagrams_drawio/default-handling-flow.png'
    mapping['4.13'] = 'generated_screenshots/erc7201-storage-layout.png'  # Screenshot, not diagram
    mapping['4.14'] = 'rendered_diagrams_drawio/analytics-architecture.png'
    mapping['4.15'] = 'rendered_diagrams_drawio/subgraph-entities.png'
    mapping['4.16'] = 'rendered_diagrams_drawio/governance-architecture.png'
    mapping['4.17'] = 'rendered_diagrams_drawio/governance-proposal-flow.png'
    mapping['4.18'] = 'rendered_diagrams_drawio/secondary-market-architecture.png'
    mapping['4.19'] = 'rendered_diagrams_drawio/transfer-restriction-flow.png'
    mapping['4.20'] = 'rendered_diagrams_drawio/kyc-architecture.png'
    mapping['4.21'] = 'rendered_diagrams_drawio/kyc-workflow.png'
    mapping['4.22'] = 'rendered_diagrams_drawio/er-diagram.png'
    mapping['4.23'] = 'rendered_diagrams_drawio/erc1155-architecture.png'
    mapping['4.24'] = 'rendered_diagrams_drawio/erc20-token-flow.png'
    mapping['4.25'] = 'rendered_diagrams_drawio/property-nft-architecture.png'
    mapping['4.26'] = 'rendered_diagrams_drawio/share-transfer-architecture.png'
    mapping['4.27'] = 'rendered_diagrams_drawio/use-case-diagram.png'
    mapping['4.28'] = 'rendered_diagrams_drawio/mobile-workflow.png'
    # Wireframes - corrected paths
    mapping['4.29'] = 'rendered_diagrams_drawio/wireframe-dashboard.png'
    mapping['4.30'] = 'rendered_diagrams_drawio/wireframe-property-registration.png'
    mapping['4.31'] = 'rendered_diagrams_drawio/wireframe-yield-agreement.png'
    mapping['4.32'] = 'rendered_diagrams_drawio/wireframe-mobile-dashboard.png'
    mapping['4.33'] = 'rendered_diagrams_drawio/wireframe-mobile-yield-agreement.png'
    
    # Section 5 - Implementation (Screenshots)
    mapping['5.1'] = 'generated_screenshots/docker-desktop-dashboard.png'
    mapping['5.2'] = 'generated_screenshots/grafana-dashboard.png'
    mapping['5.3'] = 'generated_screenshots/prometheus-targets.png'
    mapping['5.4'] = 'generated_screenshots/node-exporter-metrics.png'
    mapping['5.5'] = 'generated_screenshots/multi-env-orchestration.png'
    mapping['5.6'] = 'generated_screenshots/backup-directory.png'
    mapping['5.7'] = 'generated_screenshots/foundry-project-structure.png'
    mapping['5.8'] = 'generated_screenshots/forge-test-output.png'
    mapping['5.9'] = 'generated_screenshots/gas-report.png'
    mapping['5.10'] = 'generated_screenshots/anvil-container-logs.png'
    mapping['5.11'] = 'generated_screenshots/smart-contract-deployment.png'
    mapping['5.12'] = 'generated_screenshots/erc7201-storage-layout.png'
    mapping['5.13'] = 'generated_screenshots/bytecode-size-verification.png'
    mapping['5.14'] = 'generated_screenshots/amoy-testnet-deployment.png'
    mapping['5.15'] = 'generated_screenshots/polygonscan-confirmation.png'
    mapping['5.16'] = 'generated_screenshots/diamond-test-results.png'
    mapping['5.17'] = 'generated_screenshots/property-registration-form.png'
    mapping['5.18'] = 'generated_screenshots/yield-agreement-interface.png'
    mapping['5.19'] = 'generated_screenshots/analytics-dashboard.png'
    mapping['5.20'] = 'generated_screenshots/simulation-test-results.png'
    mapping['5.21'] = 'generated_screenshots/gas-comparison-report.png'
    mapping['5.22'] = 'generated_screenshots/variance-tracking.png'
    mapping['5.23'] = 'rendered_diagrams_drawio/yield-tokenization-flow.png'
    
    # Section 6 - Testing (Charts) - CORRECTED PATHS
    mapping['6.1'] = 'generated_charts/fig_6_1_token_standard_gas_comparison.png'
    mapping['6.2'] = 'generated_charts/fig_6_2_batch_operation_scaling.png'
    mapping['6.3'] = 'generated_charts/fig_6_3_amoy_anvil_variance_heatmap.png'
    mapping['6.4'] = 'generated_charts/fig_6_4_volatile_simulation_radar.png'
    mapping['6.5'] = 'generated_charts/fig_6_5_diamond_overhead_boxplot.png'
    mapping['6.6'] = 'rendered_diagrams_drawio/load-test-shareholder-distribution.png'
    mapping['6.7'] = 'generated_charts/fig_6_7_gas_cost_projections.png'
    mapping['6.8'] = 'generated_charts/fig_6_8_test_pass_rate_evolution.png'
    mapping['6.9'] = 'rendered_diagrams_drawio/load-test-metrics-summary.png'
    mapping['6.10'] = 'rendered_diagrams_drawio/use-case-diagram.png'
    mapping['6.11'] = 'rendered_diagrams_drawio/load-test-system-architecture.png'
    mapping['6.12'] = 'rendered_diagrams_drawio/load-test-workflow.png'
    mapping['6.13'] = 'rendered_diagrams_drawio/load-testing-results.png'
    mapping['6.14'] = 'rendered_diagrams_drawio/load-test-gas-scaling.png'
    mapping['6.15'] = 'rendered_diagrams_drawio/load-test-shareholder-distribution.png'
    mapping['6.16'] = 'rendered_diagrams_drawio/load-test-success-rates.png'
    mapping['6.17'] = 'rendered_diagrams_drawio/load-test-latency-distribution.png'
    mapping['6.18'] = 'rendered_diagrams_drawio/load-test-metrics-summary.png'
    mapping['6.19'] = 'rendered_diagrams_drawio/load-test-cost-comparison.png'
    mapping['6.20'] = 'rendered_diagrams_drawio/load-test-erc1155-efficiency.png'
    mapping['6.21'] = 'rendered_diagrams_drawio/load-test-token-recycling.png'
    mapping['6.22'] = 'rendered_diagrams_drawio/load-test-restriction-enforcement.png'
    mapping['6.23'] = 'rendered_diagrams_drawio/load-test-restriction-overhead.png'
    mapping['6.24'] = 'rendered_diagrams_drawio/load-test-rq-validation.png'
    
    # Section 7 - Evaluation (Survey/Interview charts) - CORRECTED PATHS
    mapping['7.1'] = 'survey_interview_charts/fig_demographic_overview.png'
    mapping['7.2'] = 'survey_interview_charts/fig_tokenisation_interest_analysis.png'
    mapping['7.3'] = 'survey_interview_charts/fig_correlation_matrix.png'
    mapping['7.4'] = 'survey_interview_charts/fig_cluster_analysis.png'
    mapping['7.5'] = 'survey_interview_charts/fig_feature_importance.png'
    mapping['7.6'] = 'survey_interview_charts/fig_motivations_concerns.png'
    mapping['7.7'] = 'survey_interview_charts/fig_interview_demographics.png'
    mapping['7.8'] = 'survey_interview_charts/fig_thematic_code_frequencies.png'
    mapping['7.9'] = 'survey_interview_charts/fig_landlord_fintech_comparison.png'
    mapping['7.10'] = 'survey_interview_charts/fig_likert_distributions.png'
    mapping['7.11'] = 'generated_charts/fig_6_1_token_standard_gas_comparison.png'  # Reuse
    mapping['7.12'] = 'generated_charts/fig_7_7_volatile_recovery_comparison.png'
    mapping['7.13'] = 'generated_charts/fig_6_3_amoy_anvil_variance_heatmap.png'  # Reuse
    
    # Appendix F - Wireframes - CORRECTED PATHS
    mapping['F.1'] = 'rendered_diagrams_drawio/wireframe-analytics-dashboard.png'
    mapping['F.2'] = 'rendered_diagrams_drawio/wireframe-governance.png'
    mapping['F.3'] = 'rendered_diagrams_drawio/wireframe-governance-proposal-detail.png'
    mapping['F.4'] = 'rendered_diagrams_drawio/wireframe-kyc.png'
    mapping['F.5'] = 'rendered_diagrams_drawio/wireframe-kyc-admin.png'
    mapping['F.6'] = 'rendered_diagrams_drawio/wireframe-marketplace.png'
    mapping['F.7'] = 'rendered_diagrams_drawio/wireframe-portfolio.png'
    mapping['F.8'] = 'rendered_diagrams_drawio/wireframe-properties-list.png'
    mapping['F.9'] = 'rendered_diagrams_drawio/wireframe-yield-agreement-detail.png'
    mapping['F.10'] = 'rendered_diagrams_drawio/wireframe-yield-agreements-list.png'
    mapping['F.11'] = 'rendered_diagrams_drawio/wireframe-mobile-analytics.png'
    mapping['F.12'] = 'rendered_diagrams_drawio/wireframe-mobile-governance.png'
    mapping['F.13'] = 'rendered_diagrams_drawio/wireframe-mobile-kyc.png'
    mapping['F.14'] = 'rendered_diagrams_drawio/wireframe-mobile-kyc-admin.png'
    mapping['F.15'] = 'rendered_diagrams_drawio/wireframe-mobile-marketplace.png'
    mapping['F.16'] = 'rendered_diagrams_drawio/wireframe-mobile-portfolio.png'
    mapping['F.17'] = 'rendered_diagrams_drawio/wireframe-mobile-properties-list.png'
    mapping['F.18'] = 'rendered_diagrams_drawio/wireframe-mobile-register-property.png'
    mapping['F.19'] = 'rendered_diagrams_drawio/wireframe-mobile-yield-agreements-list.png'
    
    # Survey figures (S prefix)
    mapping['S.1'] = 'survey_interview_charts/fig_demographic_overview.png'
    mapping['S.2'] = 'survey_interview_charts/fig_tokenisation_interest_analysis.png'
    mapping['S.3'] = 'survey_interview_charts/fig_correlation_matrix.png'
    mapping['S.4'] = 'survey_interview_charts/fig_cluster_analysis.png'
    mapping['S.5'] = 'survey_interview_charts/fig_motivations_concerns.png'
    mapping['S.6'] = 'survey_interview_charts/fig_feature_importance.png'
    
    # Interview figures (I prefix)
    mapping['I.1'] = 'survey_interview_charts/fig_interview_demographics.png'
    mapping['I.2'] = 'survey_interview_charts/fig_thematic_code_frequencies.png'
    mapping['I.3'] = 'survey_interview_charts/fig_landlord_fintech_comparison.png'
    mapping['I.4'] = 'survey_interview_charts/fig_likert_distributions.png'
    
    # Appendix G figures (these are references to charts moved to main body)
    mapping['G.1'] = 'survey_interview_charts/fig_cluster_analysis.png'
    mapping['G.2'] = 'survey_interview_charts/fig_feature_importance.png'
    mapping['G.3'] = 'survey_interview_charts/fig_motivations_concerns.png'
    mapping['G.4'] = 'survey_interview_charts/fig_landlord_fintech_comparison.png'
    mapping['G.5'] = 'survey_interview_charts/fig_likert_distributions.png'
    
    return mapping

def main():
    print("=" * 60)
    print("REBUILDING WORD DOCUMENT v5.0")
    print("=" * 60)
    
    # Step 1: Convert markdown to Word using pandoc
    print("\n1. Converting markdown to Word with pandoc...")
    result = subprocess.run([
        'pandoc', 
        'DissertationProgressFinal.md',
        '-o', 'DissertationProgressFinal_v2.3.docx',
        '--from', 'markdown',
        '--to', 'docx',
        '--toc',
        '--toc-depth=4'
    ], capture_output=True, text=True)
    
    if result.returncode != 0:
        print(f"Error: {result.stderr}")
        return
    
    print("   ✓ Markdown converted to Word")
    
    # Step 2: Open the document and insert images
    print("\n2. Opening document to insert images...")
    doc = Document('DissertationProgressFinal_v2.3.docx')
    
    mapping = get_image_mapping()
    
    # Verify files exist
    print("\n   Verifying image files...")
    missing_files = []
    for fig_num, path in mapping.items():
        if not os.path.exists(path):
            missing_files.append((fig_num, path))
    
    if missing_files:
        print(f"   ⚠ {len(missing_files)} mapped files not found")
    else:
        print(f"   ✓ All {len(mapping)} mapped files exist")
    
    # Count figures inserted
    figures_inserted = 0
    figures_not_found = []
    
    # Find all figure placeholders and insert images
    print("\n3. Inserting images at figure placeholders...")
    
    # Pattern to match "Figure X.X:" at the start of a paragraph (without bold markers)
    figure_pattern = re.compile(r'^Figure ([0-9]+\.[0-9]+|[A-Z]\.[0-9]+):')
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Look for figure patterns
        figure_match = figure_pattern.match(text)
        if figure_match:
            figure_num = figure_match.group(1)
            
            # Find the image file
            if figure_num in mapping:
                img_path = mapping[figure_num]
                if os.path.exists(img_path):
                    # Insert the image before this paragraph
                    new_para = para.insert_paragraph_before()
                    run = new_para.add_run()
                    try:
                        run.add_picture(img_path, width=Inches(5.5))
                        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        figures_inserted += 1
                        if figures_inserted % 20 == 0:
                            print(f"   Inserted {figures_inserted} figures...")
                    except Exception as e:
                        print(f"   Warning: Could not insert {img_path}: {e}")
                        figures_not_found.append((figure_num, str(e)))
                else:
                    figures_not_found.append((figure_num, f"File not found: {img_path}"))
            else:
                figures_not_found.append((figure_num, "No mapping defined"))
    
    print(f"\n   ✓ Inserted {figures_inserted} figures")
    
    if figures_not_found:
        print(f"\n   ⚠ {len(figures_not_found)} figures not inserted:")
        for fig, reason in figures_not_found[:15]:
            print(f"      - Figure {fig}: {reason}")
        if len(figures_not_found) > 15:
            print(f"      ... and {len(figures_not_found) - 15} more")
    
    # Step 3: Save the document
    print("\n4. Saving document...")
    doc.save('DissertationProgressFinal_v2.3.docx')
    
    # Get file size
    file_size = os.path.getsize('DissertationProgressFinal_v2.3.docx') / (1024 * 1024)
    
    print("\n" + "=" * 60)
    print("DOCUMENT REBUILD COMPLETE")
    print("=" * 60)
    print(f"Output: DissertationProgressFinal_v2.3.docx")
    print(f"Size: {file_size:.1f} MB")
    print(f"Figures inserted: {figures_inserted}")
    print(f"Figures not inserted: {len(figures_not_found)}")

if __name__ == "__main__":
    main()

