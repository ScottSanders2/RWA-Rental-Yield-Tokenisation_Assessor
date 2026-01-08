#!/usr/bin/env python3
"""
Insert ALL figures into DissertationProgressFinal.docx - Complete version
Includes: 65 Draw.io + 10 Survey + 10 Charts + 23 Screenshots = 108 images
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from pathlib import Path

# Paths
doc_path = "DissertationProgressFinal.docx"
output_path = "DissertationProgressFinal_Complete.docx"

# Image directories
drawio_dir = Path("rendered_diagrams_drawio")
survey_dir = Path("survey_interview_charts")
charts_dir = Path("generated_charts")
screenshots_dir = Path("generated_screenshots")

# Load document
print(f"Loading {doc_path}...")
doc = Document(doc_path)

# Build complete image inventory
all_images = {}

# Draw.io diagrams (65)
for f in drawio_dir.glob("*.png"):
    all_images[f.stem] = f
    all_images[f.stem.replace("-", "_")] = f

# Survey/Interview charts (10)
for f in survey_dir.glob("*.png"):
    all_images[f.stem] = f
    all_images[f.stem.replace("_", "-")] = f

# Generated charts (10)
if charts_dir.exists():
    for f in charts_dir.glob("*.png"):
        all_images[f.stem] = f

# Generated screenshots (23)
for f in screenshots_dir.glob("*.png"):
    all_images[f.stem] = f
    all_images[f.stem.replace("-", "_")] = f

print(f"Total images available: {len(all_images)}")

# Keyword mappings for screenshots
screenshot_mappings = {
    "docker desktop dashboard": "docker-desktop-dashboard",
    "container resource": "container-resource-allocation",
    "volume configuration": "volume-configuration",
    "health check": "health-check-configuration",
    "network isolation": "network-isolation-setup",
    "smart contract deployment": "smart-contract-deployment",
    "backend api test": "backend-api-test-results",
    "frontend build": "frontend-build-output",
    "mobile app screenshot": "mobile-app-screenshots",
    "analytics dashboard": "analytics-dashboard",
    "load testing result": "load-testing-results-detailed",
    "diamond pattern": "diamond-pattern-migration",
    "amoy testnet": "amoy-testnet-deployment",
    "iteration 1": "iteration-1-test-results",
    "iteration 3": "iteration-3-test-results",
    "iteration 5": "iteration-5-test-results",
    "iteration 7": "iteration-7-test-results",
    "iteration 10": "iteration-10-test-results",
    "iteration 11": "iteration-11-test-results",
    "iteration 12": "iteration-12-test-results",
    "iteration 14": "iteration-14-test-results",
    "iteration 15": "iteration-15-test-results",
    "iteration 16": "iteration-16-test-results",
}

# Survey/interview mappings
survey_mappings = {
    "demographic overview": "fig_demographic_overview",
    "survey demographics": "fig_demographic_overview",
    "tokenisation interest": "fig_tokenisation_interest_analysis",
    "correlation matrix": "fig_correlation_matrix",
    "spearman": "fig_correlation_matrix",
    "cluster analysis": "fig_cluster_analysis",
    "motivations concerns": "fig_motivations_concerns",
    "feature importance": "fig_feature_importance",
    "interview demographics": "fig_interview_demographics",
    "thematic code": "fig_thematic_code_frequencies",
    "landlord fintech": "fig_landlord_fintech_comparison",
    "landlord expert": "fig_landlord_fintech_comparison",
    "likert distribution": "fig_likert_distributions",
}

# Pattern to extract .mmd filename from placeholder
mmd_pattern = re.compile(r'([a-z0-9-]+)\.mmd', re.IGNORECASE)

figures_inserted = 0
figures_not_found = []

print("\nProcessing document paragraphs...")

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    if '[PLACEHOLDER' in text.upper():
        image_path = None
        description = text
        
        # Try to extract .mmd filename first
        mmd_match = mmd_pattern.search(text)
        if mmd_match:
            filename = mmd_match.group(1)
            if filename in all_images:
                image_path = all_images[filename]
                description = filename
        
        # If not found, try screenshot mappings
        if not image_path:
            text_lower = text.lower()
            for keyword, filename in screenshot_mappings.items():
                if keyword in text_lower:
                    if filename in all_images:
                        image_path = all_images[filename]
                        description = filename
                        break
        
        # If not found, try survey mappings
        if not image_path:
            for keyword, filename in survey_mappings.items():
                if keyword in text_lower:
                    if filename in all_images:
                        image_path = all_images[filename]
                        description = filename
                        break
        
        # If still not found, try direct filename match
        if not image_path:
            for name, path in all_images.items():
                if name in text_lower or name.replace("-", " ") in text_lower:
                    image_path = path
                    description = name
                    break
        
        if image_path and image_path.exists():
            # Clear paragraph and insert image
            para.clear()
            run = para.add_run()
            
            # Insert image (5.5 inches wide)
            run.add_picture(str(image_path), width=Inches(5.5))
            
            figures_inserted += 1
            print(f"  ✓ {description} -> {image_path.name}")
        else:
            figures_not_found.append(text[:80])
            print(f"  ✗ Not found: {text[:60]}...")

# Save document
print(f"\nSaving to {output_path}...")
doc.save(output_path)

print(f"\n{'='*60}")
print(f"SUMMARY")
print(f"{'='*60}")
print(f"Figures inserted: {figures_inserted}")
print(f"Figures not found: {len(figures_not_found)}")

if figures_not_found:
    print(f"\nPlaceholders without matching images ({len(figures_not_found)}):")
    for desc in figures_not_found[:30]:
        print(f"  - {desc}")

print(f"\n✅ Document saved: {output_path}")
print(f"   File size: {os.path.getsize(output_path) / 1024 / 1024:.1f} MB")
