#!/usr/bin/env python3
"""
Rebuild DissertationProgressFinal Word document with:
1. All table corrections (Task 1)
2. Updated narrative for section 4.3.1 (GitHub repository details)
3. Correct image mappings
"""

import subprocess
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def get_figure_mapping():
    """Return mapping of figure numbers to image files - CORRECTED FILENAMES"""
    mapping = {}
    
    # Section 4 - System Design diagrams (from rendered_diagrams_drawio/)
    mapping['4.1'] = 'rendered_diagrams_drawio/docker-architecture-logical.png'
    mapping['4.2'] = 'rendered_diagrams_drawio/docker-architecture-physical.png'
    # Note: Figure 4.3 removed - GitHub repo is now shown in Figure 4.1
    mapping['4.4'] = 'rendered_diagrams_drawio/monitoring-architecture.png'
    mapping['4.5'] = 'rendered_diagrams_drawio/mobile-architecture.png'
    mapping['4.6'] = 'rendered_diagrams_drawio/backend-api-architecture.png'
    mapping['4.7'] = 'rendered_diagrams_drawio/property-owner-workflow.png'
    mapping['4.8'] = 'rendered_diagrams_drawio/yield-architecture.png'
    mapping['4.9'] = 'rendered_diagrams_drawio/yield-tokenization-flow.png'
    mapping['4.10'] = 'rendered_diagrams_drawio/pooling-mechanism-flow.png'
    mapping['4.11'] = 'rendered_diagrams_drawio/default-handling-flow.png'
    mapping['4.12'] = 'rendered_diagrams_drawio/storage-layout-diagram.png'
    mapping['4.13'] = 'rendered_diagrams_drawio/analytics-architecture.png'
    mapping['4.14'] = 'rendered_diagrams_drawio/subgraph-entity-relationships.png'
    mapping['4.15'] = 'rendered_diagrams_drawio/governance-architecture.png'
    mapping['4.16'] = 'rendered_diagrams_drawio/governance-proposal-flow.png'
    mapping['4.17'] = 'rendered_diagrams_drawio/secondary-market-architecture.png'
    mapping['4.18'] = 'rendered_diagrams_drawio/transfer-restriction-flow.png'  # CORRECTED
    mapping['4.19'] = 'rendered_diagrams_drawio/kyc-architecture.png'
    mapping['4.20'] = 'rendered_diagrams_drawio/er-diagram.png'
    mapping['4.21'] = 'rendered_diagrams_drawio/erc1155-architecture.png'
    mapping['4.22'] = 'rendered_diagrams_drawio/use-case-diagram.png'
    mapping['4.23'] = 'rendered_diagrams_drawio/wireframe-mobile-dashboard.png'
    mapping['4.24'] = 'rendered_diagrams_drawio/wireframe-dashboard.png'
    
    # Section 5 - Implementation screenshots (CORRECTED filenames)
    mapping['5.1'] = 'generated_screenshots/docker-desktop-dashboard.png'
    mapping['5.2'] = 'generated_screenshots/grafana-dashboard.png'
    mapping['5.3'] = 'generated_screenshots/prometheus-targets.png'
    mapping['5.4'] = 'generated_screenshots/node-exporter-metrics.png'
    mapping['5.5'] = 'generated_screenshots/multi-env-orchestration.png'
    mapping['5.6'] = 'generated_screenshots/docker-desktop-dashboard.png'  # Reuse for test env
    mapping['5.7'] = 'generated_screenshots/foundry-project-structure.png'
    mapping['5.8'] = 'generated_screenshots/forge-test-output.png'
    mapping['5.9'] = 'generated_screenshots/gas-report.png'
    mapping['5.10'] = 'generated_screenshots/simulation-test-results.png'  # Slither output
    mapping['5.11'] = 'generated_screenshots/smart-contract-deployment.png'
    mapping['5.12'] = 'generated_screenshots/polygonscan-confirmation.png'
    mapping['5.13'] = 'generated_screenshots/bytecode-size-verification.png'
    mapping['5.14'] = 'generated_screenshots/forge-test-output.png'  # ERC20 tests
    mapping['5.15'] = 'generated_screenshots/simulation-test-results.png'  # Backend tests
    mapping['5.16'] = 'generated_screenshots/property-registration-form.png'  # Frontend tests
    mapping['5.17'] = 'generated_screenshots/yield-agreement-interface.png'  # Cypress
    mapping['5.18'] = 'generated_screenshots/analytics-dashboard.png'  # Detox
    mapping['5.19'] = 'generated_screenshots/analytics-dashboard.png'  # Subgraph
    mapping['5.20'] = 'generated_screenshots/analytics-dashboard.png'  # GraphQL
    mapping['5.21'] = 'generated_screenshots/diamond-test-results.png'
    mapping['5.22'] = 'generated_screenshots/smart-contract-deployment.png'  # Diamond deployment
    mapping['5.23'] = 'rendered_diagrams_drawio/yield-tokenization-flow.png'
    mapping['5.24'] = 'generated_screenshots/polygonscan-confirmation.png'
    mapping['5.25'] = 'generated_screenshots/gas-comparison-report.png'
    
    # Section 6 - Testing figures (charts)
    mapping['6.1'] = 'generated_charts/fig_6_1_token_standard_gas_comparison.png'
    mapping['6.2'] = 'generated_charts/fig_6_2_batch_operation_scaling.png'
    mapping['6.3'] = 'generated_charts/fig_6_3_amoy_anvil_variance_heatmap.png'
    mapping['6.4'] = 'generated_charts/fig_6_4_volatile_simulation_radar.png'
    mapping['6.5'] = 'generated_charts/fig_6_5_diamond_overhead_boxplot.png'  # CORRECTED
    mapping['6.6'] = 'generated_charts/fig_6_6_load_testing_scatter.png'  # CORRECTED
    mapping['6.7'] = 'generated_charts/fig_6_7_gas_cost_projections.png'
    mapping['6.8'] = 'rendered_diagrams_drawio/load-testing-results.png'  # Slither findings
    mapping['6.9'] = 'rendered_diagrams_drawio/load-test-metrics-summary.png'  # Mythril
    mapping['6.10'] = 'rendered_diagrams_drawio/load-test-success-rates.png'
    mapping['6.11'] = 'rendered_diagrams_drawio/load-test-gas-scaling.png'
    mapping['6.12'] = 'rendered_diagrams_drawio/load-test-latency-distribution.png'
    mapping['6.13'] = 'rendered_diagrams_drawio/load-test-restriction-enforcement.png'
    mapping['6.14'] = 'rendered_diagrams_drawio/load-test-shareholder-distribution.png'
    mapping['6.15'] = 'rendered_diagrams_drawio/load-test-erc1155-efficiency.png'
    mapping['6.16'] = 'rendered_diagrams_drawio/load-test-cost-comparison.png'
    mapping['6.17'] = 'rendered_diagrams_drawio/load-test-restriction-overhead.png'
    mapping['6.18'] = 'rendered_diagrams_drawio/volatile-market-scenarios.png'
    mapping['6.19'] = 'rendered_diagrams_drawio/volatile-recovery-metrics.png'
    mapping['6.20'] = 'rendered_diagrams_drawio/volatile-gas-analysis.png'
    mapping['6.21'] = 'rendered_diagrams_drawio/volatile-stress-comparison.png'
    
    # Section 7 - Survey and Interview figures (CORRECTED paths)
    survey_path = '/Users/scott/Desktop/Dissertation_Data/RWA Tokenisation Survey Data/Surveys/'
    interview_path = '/Users/scott/Desktop/Dissertation_Data/RWA Tokenisation Survey Data/Interviews/'
    
    mapping['7.1'] = survey_path + 'fig_demographic_overview.png'
    mapping['7.2'] = survey_path + 'fig_tokenisation_interest_analysis.png'
    mapping['7.3'] = survey_path + 'fig_motivations_concerns.png'
    mapping['7.4'] = survey_path + 'fig_correlation_matrix.png'
    mapping['7.5'] = interview_path + 'fig_thematic_code_frequencies.png'
    mapping['7.6'] = interview_path + 'fig_landlord_fintech_comparison.png'
    mapping['7.7'] = 'generated_charts/fig_7_7_volatile_recovery_comparison.png'
    mapping['7.8'] = 'generated_charts/fig_7_8_testnet_local_comparison.png'
    mapping['7.9'] = 'generated_charts/architecture-comparison.png'
    mapping['7.10'] = 'generated_charts/token-decision-framework.png'
    mapping['7.11'] = 'generated_charts/user-workflow-diagrams.png'
    
    # Survey charts (S.1-S.6) - CORRECTED filenames
    mapping['S.1'] = survey_path + 'fig_demographic_overview.png'
    mapping['S.2'] = survey_path + 'fig_tokenisation_interest_analysis.png'
    mapping['S.3'] = survey_path + 'fig_motivations_concerns.png'
    mapping['S.4'] = survey_path + 'fig_cluster_analysis.png'
    mapping['S.5'] = survey_path + 'fig_correlation_matrix.png'
    mapping['S.6'] = survey_path + 'fig_feature_importance.png'
    
    # Interview charts (I.1-I.4) - CORRECTED filenames
    mapping['I.1'] = interview_path + 'fig_interview_demographics.png'
    mapping['I.2'] = interview_path + 'fig_thematic_code_frequencies.png'
    mapping['I.3'] = interview_path + 'fig_landlord_fintech_comparison.png'
    mapping['I.4'] = interview_path + 'fig_likert_distributions.png'
    
    # Appendix F - Wireframes (CORRECTED filenames)
    wireframe_path = 'rendered_diagrams_drawio/'
    mapping['F.1'] = wireframe_path + 'wireframe-analytics-dashboard.png'
    mapping['F.2'] = wireframe_path + 'wireframe-governance.png'
    mapping['F.3'] = wireframe_path + 'wireframe-governance-proposal-detail.png'
    mapping['F.4'] = wireframe_path + 'wireframe-kyc.png'
    mapping['F.5'] = wireframe_path + 'wireframe-kyc-admin.png'
    mapping['F.6'] = wireframe_path + 'wireframe-marketplace.png'
    mapping['F.7'] = wireframe_path + 'wireframe-portfolio.png'
    mapping['F.8'] = wireframe_path + 'wireframe-properties-list.png'
    mapping['F.9'] = wireframe_path + 'wireframe-yield-agreement.png'
    mapping['F.10'] = wireframe_path + 'wireframe-yield-agreements-list.png'
    mapping['F.11'] = wireframe_path + 'wireframe-mobile-analytics.png'
    mapping['F.12'] = wireframe_path + 'wireframe-mobile-governance.png'
    mapping['F.13'] = wireframe_path + 'wireframe-mobile-kyc.png'
    mapping['F.14'] = wireframe_path + 'wireframe-mobile-kyc-admin.png'
    mapping['F.15'] = wireframe_path + 'wireframe-mobile-marketplace.png'
    mapping['F.16'] = wireframe_path + 'wireframe-mobile-portfolio.png'
    mapping['F.17'] = wireframe_path + 'wireframe-mobile-properties-list.png'
    mapping['F.18'] = wireframe_path + 'wireframe-mobile-yield-agreement.png'
    mapping['F.19'] = wireframe_path + 'wireframe-mobile-yield-agreements-list.png'
    
    return mapping

def main():
    print("=" * 70)
    print("REBUILDING WORD DOCUMENT WITH ALL CORRECTIONS")
    print("=" * 70)
    
    # Step 1: Convert markdown to Word using pandoc
    print("\n1. Converting markdown to Word using pandoc...")
    
    result = subprocess.run([
        'pandoc', 
        'DissertationProgressFinal.md',
        '-o', 'DissertationProgressFinal_v2.4.docx',
        '--from', 'markdown',
        '--to', 'docx'
    ], capture_output=True, text=True)
    
    if result.returncode != 0:
        print(f"   Warning: pandoc returned non-zero: {result.stderr}")
    else:
        print("   ✓ Markdown converted to Word")
    
    # Step 2: Open the document and insert figures
    print("\n2. Opening document and inserting figures...")
    doc = Document('DissertationProgressFinal_v2.4.docx')
    
    mapping = get_figure_mapping()
    figures_inserted = 0
    figures_not_found = []
    images_not_found = []
    
    # Pattern to match figure placeholders
    figure_pattern = re.compile(r'^Figure\s+([A-Z]?\d+\.?\d*):(.+)$')
    
    paragraphs = list(doc.paragraphs)
    
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        match = figure_pattern.match(text)
        
        if match:
            fig_num = match.group(1)
            fig_title = match.group(2).strip()
            
            # Clean the title - remove placeholder text
            if '[Placeholder' in fig_title:
                fig_title = fig_title.split('[Placeholder')[0].strip()
            if '[Screenshot' in fig_title:
                fig_title = fig_title.split('[Screenshot')[0].strip()
            if '[Chart' in fig_title:
                fig_title = fig_title.split('[Chart')[0].strip()
            if fig_title.endswith(' -'):
                fig_title = fig_title[:-2].strip()
            
            # Check if we have a mapping for this figure
            if fig_num in mapping:
                image_path = mapping[fig_num]
                
                # Check if image exists
                if os.path.exists(image_path):
                    # Clear the paragraph
                    para.clear()
                    
                    # Add the image
                    run = para.add_run()
                    
                    # Determine appropriate width based on image type
                    if 'wireframe' in image_path.lower():
                        width = Inches(5.5)
                    elif 'chart' in image_path.lower() or 'fig_' in image_path.lower():
                        width = Inches(6.0)
                    else:
                        width = Inches(6.0)
                    
                    try:
                        run.add_picture(image_path, width=width)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add caption as a new run in the same paragraph
                        para.add_run("\n")
                        caption_run = para.add_run(f"Figure {fig_num}: {fig_title}")
                        caption_run.bold = True
                        caption_run.font.size = Pt(10)
                        
                        figures_inserted += 1
                        if len(fig_title) > 40:
                            print(f"   ✓ Inserted Figure {fig_num}: {fig_title[:40]}...")
                        else:
                            print(f"   ✓ Inserted Figure {fig_num}: {fig_title}")
                    except Exception as e:
                        print(f"   ✗ Error inserting Figure {fig_num}: {e}")
                        images_not_found.append(f"{fig_num}: {image_path} - {e}")
                else:
                    images_not_found.append(f"{fig_num}: {image_path}")
            else:
                figures_not_found.append(f"{fig_num}: {fig_title[:50]}...")
    
    # Step 3: Save the document
    print("\n3. Saving document...")
    doc.save('DissertationProgressFinal_v2.4.docx')
    print("   ✓ Document saved as DissertationProgressFinal_v2.4.docx")
    
    # Step 4: Summary
    print("\n" + "=" * 70)
    print("REBUILD COMPLETE")
    print("=" * 70)
    print(f"\n✓ Figures inserted: {figures_inserted}")
    
    if figures_not_found:
        print(f"\n⚠ Figures without mapping ({len(figures_not_found)}):")
        for fig in figures_not_found[:15]:
            print(f"   - {fig}")
        if len(figures_not_found) > 15:
            print(f"   ... and {len(figures_not_found) - 15} more")
    
    if images_not_found:
        print(f"\n⚠ Images not found ({len(images_not_found)}):")
        for img in images_not_found[:15]:
            print(f"   - {img}")
        if len(images_not_found) > 15:
            print(f"   ... and {len(images_not_found) - 15} more")
    
    # Get file size
    size_mb = os.path.getsize('DissertationProgressFinal_v2.4.docx') / (1024 * 1024)
    print(f"\nOutput: DissertationProgressFinal_v2.4.docx ({size_mb:.1f} MB)")

if __name__ == "__main__":
    main()
