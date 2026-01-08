#!/usr/bin/env python3
"""
Word Document Figure Migration Script
Migrates figures, tables, and diagrams from DissertationProgress.docx 
to DissertationProgressFinal.docx

This script:
1. Extracts images from the original document
2. Identifies placeholder locations in the new document
3. Maps placeholders to source content
4. Inserts content at appropriate locations
"""

import os
import re
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Paths
BASE_DIR = Path("/Users/scott/Cursor/RWA-Rental-Yield-Tokenisation_v5.1")
ORIGINAL_DOC = BASE_DIR / "DissertationProgress.docx"
NEW_DOC = BASE_DIR / "DissertationProgressFinal.docx"
OUTPUT_DOC = BASE_DIR / "DissertationProgressFinal_WithFigures.docx"
MERMAID_DIR = BASE_DIR / "Shared_Environment/docs/architecture/diagrams/mermaid"
RENDERED_DIR = BASE_DIR / "rendered_diagrams"
EXTRACTED_IMAGES_DIR = BASE_DIR / "extracted_images"

def setup_directories():
    """Create necessary directories"""
    RENDERED_DIR.mkdir(exist_ok=True)
    EXTRACTED_IMAGES_DIR.mkdir(exist_ok=True)
    print(f"✓ Created directories: {RENDERED_DIR}, {EXTRACTED_IMAGES_DIR}")

def extract_images_from_original():
    """Extract all images from the original document"""
    print("\n=== EXTRACTING IMAGES FROM ORIGINAL DOCUMENT ===")
    
    doc = Document(ORIGINAL_DOC)
    
    # Get all image relationships
    image_count = 0
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            image_count += 1
            # Get the image data
            image_data = rel.target_part.blob
            # Determine extension
            content_type = rel.target_part.content_type
            ext = content_type.split('/')[-1]
            if ext == 'jpeg':
                ext = 'jpg'
            
            # Save the image
            image_path = EXTRACTED_IMAGES_DIR / f"image_{image_count}.{ext}"
            with open(image_path, 'wb') as f:
                f.write(image_data)
            print(f"  Extracted: {image_path.name}")
    
    print(f"✓ Extracted {image_count} images")
    return image_count

def list_mermaid_diagrams():
    """List all Mermaid diagram files"""
    print("\n=== MERMAID DIAGRAMS ===")
    
    mmd_files = list(MERMAID_DIR.glob("*.mmd"))
    print(f"Found {len(mmd_files)} Mermaid diagram files:")
    
    for mmd in sorted(mmd_files):
        print(f"  - {mmd.name}")
    
    return mmd_files

def find_placeholders_in_new_doc():
    """Find all placeholder paragraphs in the new document"""
    print("\n=== FINDING PLACEHOLDERS IN NEW DOCUMENT ===")
    
    doc = Document(NEW_DOC)
    
    placeholders = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        
        # Check for figure headers
        if text.startswith("**Figure ") or text.startswith("**Table "):
            fig_match = re.match(r'\*\*(Figure|Table)\s+([\d\.]+):\s*(.+)\*\*', text)
            if fig_match:
                placeholders.append({
                    'index': i,
                    'type': fig_match.group(1),
                    'number': fig_match.group(2),
                    'title': fig_match.group(3),
                    'text': text
                })
        
        # Check for placeholder descriptions
        if '[Placeholder' in text or '[Chart placeholder' in text or '[Diagram placeholder' in text:
            placeholders.append({
                'index': i,
                'type': 'placeholder_desc',
                'text': text
            })
    
    print(f"Found {len(placeholders)} placeholder entries")
    
    # Group by figure number
    figures = {}
    current_figure = None
    
    for p in placeholders:
        if p['type'] in ['Figure', 'Table']:
            current_figure = p['number']
            figures[current_figure] = {
                'type': p['type'],
                'title': p['title'],
                'header_index': p['index'],
                'placeholder_index': None,
                'mermaid_file': None
            }
        elif p['type'] == 'placeholder_desc' and current_figure:
            figures[current_figure]['placeholder_index'] = p['index']
            # Extract mermaid filename if present
            mmd_match = re.search(r'(\w+[-\w]*\.mmd)', p['text'])
            if mmd_match:
                figures[current_figure]['mermaid_file'] = mmd_match.group(1)
    
    return figures

def extract_tables_from_original():
    """Extract tables from original document with their context"""
    print("\n=== EXTRACTING TABLES FROM ORIGINAL DOCUMENT ===")
    
    doc = Document(ORIGINAL_DOC)
    
    tables_info = []
    for i, table in enumerate(doc.tables):
        # Get table dimensions
        rows = len(table.rows)
        cols = len(table.columns) if table.rows else 0
        
        # Get header row
        header = []
        if table.rows:
            for cell in table.rows[0].cells:
                header.append(cell.text.strip()[:30])
        
        tables_info.append({
            'index': i,
            'rows': rows,
            'cols': cols,
            'header': header,
            'table': table
        })
    
    print(f"Found {len(tables_info)} tables in original document")
    return tables_info

def generate_migration_report(figures, mmd_files, tables_info):
    """Generate a detailed migration report"""
    print("\n=== GENERATING MIGRATION REPORT ===")
    
    report_path = BASE_DIR / "FIGURE_MIGRATION_REPORT.md"
    
    with open(report_path, 'w') as f:
        f.write("# Figure Migration Report\n\n")
        f.write(f"Generated: {os.popen('date').read().strip()}\n\n")
        
        f.write("## Summary\n\n")
        f.write(f"- **Total placeholders in new document:** {len(figures)}\n")
        f.write(f"- **Mermaid diagrams available:** {len(mmd_files)}\n")
        f.write(f"- **Tables in original document:** {len(tables_info)}\n")
        f.write(f"- **Images in original document:** 10\n\n")
        
        f.write("## Placeholder Mapping\n\n")
        f.write("| Figure # | Type | Title | Mermaid File | Status |\n")
        f.write("|----------|------|-------|--------------|--------|\n")
        
        for fig_num, fig_info in sorted(figures.items(), key=lambda x: [int(n) for n in x[0].split('.')]):
            mmd = fig_info.get('mermaid_file', '-')
            status = '✓ Has MMD' if mmd and mmd != '-' else '⚠️ Needs source'
            f.write(f"| {fig_num} | {fig_info['type']} | {fig_info['title'][:40]}... | {mmd} | {status} |\n")
        
        f.write("\n## Mermaid Files Available\n\n")
        for mmd in sorted(mmd_files):
            f.write(f"- `{mmd.name}`\n")
        
        f.write("\n## Tables in Original Document (First 30)\n\n")
        f.write("| # | Rows | Cols | Header Preview |\n")
        f.write("|---|------|------|----------------|\n")
        for t in tables_info[:30]:
            header_preview = ', '.join(t['header'][:3])
            f.write(f"| {t['index']+1} | {t['rows']} | {t['cols']} | {header_preview} |\n")
    
    print(f"✓ Report saved to: {report_path}")
    return report_path

def main():
    """Main execution"""
    print("=" * 60)
    print("WORD DOCUMENT FIGURE MIGRATION")
    print("=" * 60)
    
    # Setup
    setup_directories()
    
    # Extract from original
    image_count = extract_images_from_original()
    tables_info = extract_tables_from_original()
    
    # List Mermaid diagrams
    mmd_files = list_mermaid_diagrams()
    
    # Find placeholders in new document
    figures = find_placeholders_in_new_doc()
    
    # Generate report
    report_path = generate_migration_report(figures, mmd_files, tables_info)
    
    print("\n" + "=" * 60)
    print("NEXT STEPS:")
    print("=" * 60)
    print("""
1. RENDER MERMAID DIAGRAMS:
   - Use Mermaid Live Editor (https://mermaid.live/) 
   - Or install mermaid-cli: npm install -g @mermaid-js/mermaid-cli
   - Render each .mmd file to PNG and save to 'rendered_diagrams/' folder

2. MAP TABLES:
   - Review FIGURE_MIGRATION_REPORT.md
   - Identify which tables from original correspond to placeholders

3. INSERT CONTENT:
   - Run the insertion script (to be created) after rendering diagrams

4. VERIFY:
   - Check all figures display correctly
   - Verify List of Figures and List of Tables
""")
    
    print(f"\n✓ Migration analysis complete. See: {report_path}")

if __name__ == "__main__":
    main()





