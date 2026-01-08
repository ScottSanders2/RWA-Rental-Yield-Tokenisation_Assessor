#!/usr/bin/env python3
"""
Rebuild DissertationProgressFinal Word document v7 with:
1. All table corrections
2. Updated figure numbering after removals
3. Correct image mappings for all figures
4. New Draw.io images (Foundry Structure, ERC-7201 Storage Layout)
"""

import subprocess
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Base paths
BASE_PATH = '/Users/scott/Cursor/RWA-Rental-Yield-Tokenisation_v5.1'
RENDERED_DIAGRAMS = f'{BASE_PATH}/rendered_diagrams_drawio'
GENERATED_SCREENSHOTS = f'{BASE_PATH}/generated_screenshots'
GENERATED_CHARTS = f'{BASE_PATH}/generated_charts'
DRAWIO_PNG = f'{BASE_PATH}/Shared_Environment/docs/architecture/diagrams/png'
SURVEY_PATH = '/Users/scott/Desktop/Dissertation_Data/RWA Tokenisation Survey Data/Surveys'
INTERVIEW_PATH = '/Users/scott/Desktop/Dissertation_Data/RWA Tokenisation Survey Data/Interviews'

def get_figure_mapping():
    """Return mapping of figure numbers to image files - UPDATED FOR V7"""
    mapping = {}
    
    # =========================================================================
    # SECTION 4 - System Design diagrams (33 figures: 4.1-4.33)
    # =========================================================================
    mapping['4.1'] = f'{RENDERED_DIAGRAMS}/docker-architecture-logical.png'
    mapping['4.2'] = f'{RENDERED_DIAGRAMS}/docker-architecture-physical.png'
    # Note: 4.3 was removed (GitHub repo is now shown in Figure 4.1)
    mapping['4.4'] = f'{RENDERED_DIAGRAMS}/monitoring-architecture.png'
    mapping['4.5'] = f'{RENDERED_DIAGRAMS}/mobile-architecture.png'
    mapping['4.6'] = f'{RENDERED_DIAGRAMS}/backend-api-architecture.png'
    mapping['4.7'] = f'{RENDERED_DIAGRAMS}/yield-architecture.png'  # Property Owner Workflow / Yield Architecture
    mapping['4.8'] = f'{RENDERED_DIAGRAMS}/yield-tokenization-flow.png'
    mapping['4.9'] = f'{RENDERED_DIAGRAMS}/default-handling-flow.png'  # Repayment Processing
    mapping['4.10'] = f'{RENDERED_DIAGRAMS}/pooling-mechanism-flow.png'
    mapping['4.11'] = f'{RENDERED_DIAGRAMS}/default-handling-flow.png'
    mapping['4.12'] = f'{RENDERED_DIAGRAMS}/storage-layout-diagram.png'  # ERC-7201
    mapping['4.13'] = f'{RENDERED_DIAGRAMS}/analytics-architecture.png'
    mapping['4.14'] = f'{RENDERED_DIAGRAMS}/subgraph-entity-relationships.png'
    mapping['4.16'] = f'{RENDERED_DIAGRAMS}/governance-architecture.png'
    mapping['4.17'] = f'{RENDERED_DIAGRAMS}/governance-proposal-flow.png'
    mapping['4.18'] = f'{RENDERED_DIAGRAMS}/secondary-market-architecture.png'
    mapping['4.19'] = f'{RENDERED_DIAGRAMS}/transfer-restriction-flow.png'
    mapping['4.20'] = f'{RENDERED_DIAGRAMS}/kyc-architecture.png'
    mapping['4.21'] = f'{RENDERED_DIAGRAMS}/kyc-workflow.png'
    mapping['4.22'] = f'{RENDERED_DIAGRAMS}/er-diagram.png'
    mapping['4.23'] = f'{RENDERED_DIAGRAMS}/erc1155-architecture.png'
    mapping['4.24'] = f'{RENDERED_DIAGRAMS}/erc20-token-flow.png'
    mapping['4.25'] = f'{RENDERED_DIAGRAMS}/property-nft-architecture.png'
    mapping['4.26'] = f'{RENDERED_DIAGRAMS}/share-transfer-architecture.png'
    mapping['4.27'] = f'{RENDERED_DIAGRAMS}/use-case-diagram.png'
    mapping['4.28'] = f'{RENDERED_DIAGRAMS}/mobile-workflow.png'
    mapping['4.29'] = f'{RENDERED_DIAGRAMS}/wireframe-dashboard.png'
    mapping['4.30'] = f'{RENDERED_DIAGRAMS}/wireframe-property-registration.png'
    mapping['4.31'] = f'{RENDERED_DIAGRAMS}/wireframe-yield-agreement.png'
    mapping['4.32'] = f'{RENDERED_DIAGRAMS}/wireframe-mobile-dashboard.png'
    mapping['4.33'] = f'{RENDERED_DIAGRAMS}/wireframe-mobile-yield-agreement.png'
    
    # =========================================================================
    # SECTION 5 - Implementation (18 figures: 5.1-5.18 after removals)
    # =========================================================================
    mapping['5.1'] = f'{GENERATED_SCREENSHOTS}/docker-desktop-dashboard.png'
    mapping['5.2'] = f'{GENERATED_SCREENSHOTS}/grafana-dashboard.png'
    mapping['5.3'] = f'{GENERATED_SCREENSHOTS}/prometheus-targets.png'
    mapping['5.4'] = f'{GENERATED_SCREENSHOTS}/backup-directory.png'
    mapping['5.5'] = f'{DRAWIO_PNG}/foundry-smart-contract-project-structure.drawio.png'  # NEW Draw.io
    mapping['5.6'] = f'{GENERATED_SCREENSHOTS}/gas-report.png'
    mapping['5.7'] = f'{DRAWIO_PNG}/erc7201-storage-layout.drawio.png'  # NEW Draw.io
    mapping['5.8'] = f'{GENERATED_SCREENSHOTS}/bytecode-size-verification.png'
    mapping['5.9'] = f'{GENERATED_SCREENSHOTS}/property-registration-form.png'
    mapping['5.10'] = f'{GENERATED_SCREENSHOTS}/yield-agreement-interface.png'
    mapping['5.11'] = f'{GENERATED_SCREENSHOTS}/analytics-dashboard.png'
    mapping['5.12'] = f'{GENERATED_SCREENSHOTS}/amoy-testnet-deployment.png'
    mapping['5.13'] = f'{GENERATED_SCREENSHOTS}/polygonscan-confirmation.png'
    mapping['5.14'] = f'{GENERATED_SCREENSHOTS}/diamond-test-results.png'
    mapping['5.15'] = f'{GENERATED_SCREENSHOTS}/simulation-test-results.png'
    mapping['5.16'] = f'{GENERATED_SCREENSHOTS}/gas-comparison-report.png'
    mapping['5.17'] = f'{GENERATED_SCREENSHOTS}/variance-tracking.png'
    mapping['5.18'] = f'{RENDERED_DIAGRAMS}/yield-tokenization-flow.png'
    
    # =========================================================================
    # SECTION 6 - Testing (24 figures: 6.1-6.24)
    # =========================================================================
    mapping['6.1'] = f'{GENERATED_CHARTS}/fig_6_1_token_standard_gas_comparison.png'
    mapping['6.2'] = f'{GENERATED_CHARTS}/fig_6_2_batch_operation_scaling.png'
    mapping['6.3'] = f'{GENERATED_CHARTS}/fig_6_3_amoy_anvil_variance_heatmap.png'
    mapping['6.4'] = f'{GENERATED_CHARTS}/fig_6_4_volatile_simulation_radar.png'
    mapping['6.5'] = f'{GENERATED_CHARTS}/fig_6_5_diamond_overhead_boxplot.png'
    mapping['6.6'] = f'{GENERATED_CHARTS}/fig_6_6_load_testing_scatter.png'
    mapping['6.7'] = f'{GENERATED_CHARTS}/fig_6_7_gas_cost_projections.png'
    mapping['6.8'] = f'{GENERATED_CHARTS}/fig_6_8_test_pass_rate_evolution.png'
    mapping['6.9'] = f'{GENERATED_CHARTS}/architecture-comparison.png'
    mapping['6.10'] = f'{GENERATED_CHARTS}/user-workflow-diagrams.png'
    mapping['6.11'] = f'{RENDERED_DIAGRAMS}/load-test-system-architecture.png'
    mapping['6.12'] = f'{RENDERED_DIAGRAMS}/load-test-workflow.png'
    mapping['6.13'] = f'{RENDERED_DIAGRAMS}/load-testing-results.png'
    mapping['6.14'] = f'{RENDERED_DIAGRAMS}/load-test-gas-scaling.png'
    mapping['6.15'] = f'{RENDERED_DIAGRAMS}/load-test-shareholder-distribution.png'
    mapping['6.16'] = f'{RENDERED_DIAGRAMS}/load-test-success-rates.png'
    mapping['6.17'] = f'{RENDERED_DIAGRAMS}/load-test-latency-distribution.png'
    mapping['6.18'] = f'{RENDERED_DIAGRAMS}/load-test-metrics-summary.png'
    mapping['6.19'] = f'{RENDERED_DIAGRAMS}/load-test-cost-comparison.png'
    mapping['6.20'] = f'{RENDERED_DIAGRAMS}/load-test-erc1155-efficiency.png'
    mapping['6.21'] = f'{RENDERED_DIAGRAMS}/load-test-token-recycling.png'
    mapping['6.22'] = f'{RENDERED_DIAGRAMS}/load-test-restriction-enforcement.png'
    mapping['6.23'] = f'{RENDERED_DIAGRAMS}/load-test-restriction-overhead.png'
    mapping['6.24'] = f'{RENDERED_DIAGRAMS}/load-test-rq-validation.png'
    
    # =========================================================================
    # SECTION 7 - Evaluation (13 figures: 7.1-7.13)
    # =========================================================================
    mapping['7.1'] = f'{SURVEY_PATH}/fig_demographic_overview.png'
    mapping['7.2'] = f'{SURVEY_PATH}/fig_tokenisation_interest_analysis.png'
    mapping['7.3'] = f'{SURVEY_PATH}/fig_correlation_matrix.png'
    mapping['7.4'] = f'{INTERVIEW_PATH}/fig_interview_demographics.png'
    mapping['7.5'] = f'{INTERVIEW_PATH}/fig_thematic_code_frequencies.png'
    mapping['7.11'] = f'{GENERATED_CHARTS}/token-decision-framework.png'
    mapping['7.12'] = f'{GENERATED_CHARTS}/fig_7_7_volatile_recovery_comparison.png'
    mapping['7.13'] = f'{GENERATED_CHARTS}/fig_7_8_testnet_local_comparison.png'
    
    # =========================================================================
    # SURVEY FIGURES (S.1-S.6)
    # =========================================================================
    mapping['S.1'] = f'{SURVEY_PATH}/fig_demographic_overview.png'
    mapping['S.2'] = f'{SURVEY_PATH}/fig_tokenisation_interest_analysis.png'
    mapping['S.3'] = f'{SURVEY_PATH}/fig_correlation_matrix.png'
    mapping['S.4'] = f'{SURVEY_PATH}/fig_cluster_analysis.png'
    mapping['S.5'] = f'{SURVEY_PATH}/fig_motivations_concerns.png'
    mapping['S.6'] = f'{SURVEY_PATH}/fig_feature_importance.png'
    
    # =========================================================================
    # INTERVIEW FIGURES (I.1-I.4)
    # =========================================================================
    mapping['I.1'] = f'{INTERVIEW_PATH}/fig_interview_demographics.png'
    mapping['I.2'] = f'{INTERVIEW_PATH}/fig_thematic_code_frequencies.png'
    mapping['I.3'] = f'{INTERVIEW_PATH}/fig_landlord_fintech_comparison.png'
    mapping['I.4'] = f'{INTERVIEW_PATH}/fig_likert_distributions.png'
    
    # =========================================================================
    # APPENDIX F - Wireframes (F.1-F.19)
    # =========================================================================
    mapping['F.1'] = f'{RENDERED_DIAGRAMS}/wireframe-analytics-dashboard.png'
    mapping['F.2'] = f'{RENDERED_DIAGRAMS}/wireframe-governance.png'
    mapping['F.3'] = f'{RENDERED_DIAGRAMS}/wireframe-governance-proposal-detail.png'
    mapping['F.4'] = f'{RENDERED_DIAGRAMS}/wireframe-kyc.png'
    mapping['F.5'] = f'{RENDERED_DIAGRAMS}/wireframe-kyc-admin.png'
    mapping['F.6'] = f'{RENDERED_DIAGRAMS}/wireframe-marketplace.png'
    mapping['F.7'] = f'{RENDERED_DIAGRAMS}/wireframe-portfolio.png'
    mapping['F.8'] = f'{RENDERED_DIAGRAMS}/wireframe-properties-list.png'
    mapping['F.9'] = f'{RENDERED_DIAGRAMS}/wireframe-yield-agreement-detail.png'
    mapping['F.10'] = f'{RENDERED_DIAGRAMS}/wireframe-yield-agreements-list.png'
    mapping['F.11'] = f'{RENDERED_DIAGRAMS}/wireframe-mobile-analytics.png'
    mapping['F.12'] = f'{RENDERED_DIAGRAMS}/wireframe-mobile-governance.png'
    mapping['F.13'] = f'{RENDERED_DIAGRAMS}/wireframe-mobile-kyc.png'
    mapping['F.14'] = f'{RENDERED_DIAGRAMS}/wireframe-mobile-kyc-admin.png'
    mapping['F.15'] = f'{RENDERED_DIAGRAMS}/wireframe-mobile-marketplace.png'
    mapping['F.16'] = f'{RENDERED_DIAGRAMS}/wireframe-mobile-portfolio.png'
    mapping['F.17'] = f'{RENDERED_DIAGRAMS}/wireframe-mobile-properties-list.png'
    mapping['F.18'] = f'{RENDERED_DIAGRAMS}/wireframe-mobile-register-property.png'
    mapping['F.19'] = f'{RENDERED_DIAGRAMS}/wireframe-mobile-yield-agreements-list.png'
    
    # =========================================================================
    # APPENDIX G - Additional Research Charts (G.1-G.5)
    # =========================================================================
    mapping['G.1'] = f'{SURVEY_PATH}/fig_cluster_analysis.png'
    mapping['G.2'] = f'{SURVEY_PATH}/fig_feature_importance.png'
    mapping['G.3'] = f'{SURVEY_PATH}/fig_motivations_concerns.png'
    mapping['G.4'] = f'{INTERVIEW_PATH}/fig_landlord_fintech_comparison.png'
    mapping['G.5'] = f'{INTERVIEW_PATH}/fig_likert_distributions.png'
    
    return mapping

def main():
    os.chdir(BASE_PATH)
    
    print("=" * 70)
    print("REBUILDING WORD DOCUMENT V7 - FINAL VERSION")
    print("=" * 70)
    
    # Step 1: Convert markdown to Word using pandoc
    print("\n1. Converting markdown to Word using pandoc...")
    
    result = subprocess.run([
        'pandoc', 
        'DissertationProgressFinal.md',
        '-o', 'DissertationProgressFinal_v2.3.docx',
        '--from', 'markdown',
        '--to', 'docx'
    ], capture_output=True, text=True)
    
    if result.returncode != 0:
        print(f"   Warning: pandoc returned non-zero: {result.stderr}")
    else:
        print("   ✓ Markdown converted to Word")
    
    # Step 2: Open the document and insert figures
    print("\n2. Opening document and inserting figures...")
    doc = Document('DissertationProgressFinal_v2.3.docx')
    
    mapping = get_figure_mapping()
    figures_inserted = 0
    figures_not_found = []
    images_not_found = []
    
    # Pattern to match figure placeholders (pandoc removes ** from bold)
    # Matches: Figure 4.1, Figure S.1, Figure F.10, Figure G.5, etc.
    figure_pattern = re.compile(r'^Figure\s+([A-Z]?\d*\.?\d+):(.+)$')
    
    paragraphs = list(doc.paragraphs)
    
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        match = figure_pattern.match(text)
        
        if match:
            fig_num = match.group(1)
            fig_title = match.group(2).strip()
            
            # Clean the title - remove placeholder text
            if '[Placeholder' in fig_title:
                fig_title = fig_title.split('[Placeholder')[0].strip()
            if '[Screenshot' in fig_title:
                fig_title = fig_title.split('[Screenshot')[0].strip()
            if '[Chart' in fig_title:
                fig_title = fig_title.split('[Chart')[0].strip()
            if '[Diagram' in fig_title:
                fig_title = fig_title.split('[Diagram')[0].strip()
            if fig_title.endswith(' -'):
                fig_title = fig_title[:-2].strip()
            
            # Check if we have a mapping for this figure
            if fig_num in mapping:
                image_path = mapping[fig_num]
                
                # Check if image exists
                if os.path.exists(image_path):
                    # Clear the paragraph
                    para.clear()
                    
                    # Add the image
                    run = para.add_run()
                    
                    # Determine appropriate width based on image type
                    if 'wireframe' in image_path.lower():
                        width = Inches(5.5)
                    elif 'chart' in image_path.lower() or 'fig_' in image_path.lower():
                        width = Inches(6.0)
                    else:
                        width = Inches(6.0)
                    
                    try:
                        run.add_picture(image_path, width=width)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add caption as a new run in the same paragraph
                        para.add_run("\n")
                        caption_run = para.add_run(f"Figure {fig_num}: {fig_title}")
                        caption_run.bold = True
                        caption_run.font.size = Pt(10)
                        
                        figures_inserted += 1
                        if len(fig_title) > 50:
                            print(f"   ✓ Figure {fig_num}: {fig_title[:50]}...")
                        else:
                            print(f"   ✓ Figure {fig_num}: {fig_title}")
                    except Exception as e:
                        print(f"   ✗ Error inserting Figure {fig_num}: {e}")
                        images_not_found.append(f"{fig_num}: {image_path} - {e}")
                else:
                    images_not_found.append(f"{fig_num}: {image_path}")
                    print(f"   ⚠ Image not found for Figure {fig_num}: {image_path}")
            else:
                figures_not_found.append(f"{fig_num}: {fig_title[:50]}...")
    
    # Step 3: Save the document
    print("\n3. Saving document...")
    doc.save('DissertationProgressFinal_v2.3.docx')
    print("   ✓ Document saved as DissertationProgressFinal_v2.3.docx")
    
    # Step 4: Summary
    print("\n" + "=" * 70)
    print("REBUILD COMPLETE")
    print("=" * 70)
    print(f"\n✓ Figures inserted: {figures_inserted}")
    
    if figures_not_found:
        print(f"\n⚠ Figures without mapping ({len(figures_not_found)}):")
        for fig in figures_not_found[:20]:
            print(f"   - {fig}")
        if len(figures_not_found) > 20:
            print(f"   ... and {len(figures_not_found) - 20} more")
    
    if images_not_found:
        print(f"\n⚠ Images not found ({len(images_not_found)}):")
        for img in images_not_found[:20]:
            print(f"   - {img}")
        if len(images_not_found) > 20:
            print(f"   ... and {len(images_not_found) - 20} more")
    
    # Get file size
    size_mb = os.path.getsize('DissertationProgressFinal_v2.3.docx') / (1024 * 1024)
    print(f"\nOutput: DissertationProgressFinal_v2.3.docx ({size_mb:.1f} MB)")
    
    # Calculate expected vs actual
    expected_figures = len(mapping)
    print(f"\nExpected figures in mapping: {expected_figures}")
    print(f"Figures successfully inserted: {figures_inserted}")
    
    if figures_inserted < expected_figures:
        print(f"\n⚠ {expected_figures - figures_inserted} figures may be missing from document")

if __name__ == "__main__":
    main()

