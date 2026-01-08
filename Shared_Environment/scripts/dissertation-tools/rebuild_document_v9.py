#!/usr/bin/env python3
"""
Rebuild DissertationProgressFinal Word document from restructured markdown.
Version 9: Creates v2.6 with restructured Section 4.3 and 4.5.
"""

import subprocess
import os
import re
from pathlib import Path

# Paths
WORKSPACE = Path("/Users/scott/Cursor/RWA-Rental-Yield-Tokenisation_v5.1")
MARKDOWN_FILE = WORKSPACE / "DissertationProgressFinal_v2.6.md"
OUTPUT_DOCX = WORKSPACE / "DissertationProgressFinal_v2.6.docx"
TEMP_DOCX = WORKSPACE / "temp_converted_v9.docx"

# Image directories
DIAGRAMS_DIR = WORKSPACE / "rendered_diagrams_drawio"
GENERATED_SCREENSHOTS = WORKSPACE / "generated_screenshots"
GENERATED_CHARTS = WORKSPACE / "generated_charts"
SURVEY_CHARTS = WORKSPACE / "survey_interview_charts"
PNG_DIR = WORKSPACE / "Shared_Environment/docs/architecture/diagrams/png"

def convert_markdown_to_docx():
    """Convert markdown to Word using pandoc."""
    print("Converting markdown to Word document...")
    cmd = [
        "pandoc",
        str(MARKDOWN_FILE),
        "-o", str(TEMP_DOCX),
        "--from", "markdown",
        "--to", "docx",
        "--standalone",
        "--toc",
        "--toc-depth=3"
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print(f"Pandoc error: {result.stderr}")
        return False
    print(f"Created temporary document: {TEMP_DOCX}")
    return True

def get_image_mapping():
    """Create mapping from figure numbers to image files."""
    mapping = {}
    
    # Section 4 diagrams (from rendered_diagrams_drawio)
    mapping["4.1"] = DIAGRAMS_DIR / "docker-architecture-logical.png"
    mapping["4.2"] = DIAGRAMS_DIR / "docker-architecture-physical.png"
    mapping["4.3"] = DIAGRAMS_DIR / "docker-architecture-logical.png"  # GitHub repo is part of logical view
    mapping["4.4"] = DIAGRAMS_DIR / "monitoring-architecture.png"
    mapping["4.5"] = DIAGRAMS_DIR / "mobile-architecture.png"
    mapping["4.6"] = DIAGRAMS_DIR / "backend-api-architecture.png"
    mapping["4.7"] = DIAGRAMS_DIR / "property-owner-workflow.png"
    mapping["4.8"] = DIAGRAMS_DIR / "yield-architecture.png"
    mapping["4.9"] = DIAGRAMS_DIR / "yield-tokenization-flow.png"
    mapping["4.10"] = DIAGRAMS_DIR / "repayment-processing-flow.png"
    mapping["4.11"] = DIAGRAMS_DIR / "pooling-mechanism-flow.png"
    mapping["4.12"] = DIAGRAMS_DIR / "default-handling-flow.png"
    mapping["4.13"] = DIAGRAMS_DIR / "storage-layout-diagram.png"
    mapping["4.14"] = DIAGRAMS_DIR / "analytics-architecture.png"
    mapping["4.15"] = DIAGRAMS_DIR / "subgraph-entity-relationships.png"
    mapping["4.16"] = DIAGRAMS_DIR / "governance-architecture.png"
    mapping["4.17"] = DIAGRAMS_DIR / "governance-proposal-flow.png"
    mapping["4.18"] = DIAGRAMS_DIR / "secondary-market-architecture.png"
    mapping["4.19"] = DIAGRAMS_DIR / "transfer-restriction-flow.png"
    mapping["4.20"] = DIAGRAMS_DIR / "kyc-architecture.png"
    mapping["4.21"] = DIAGRAMS_DIR / "kyc-workflow.png"
    mapping["4.22"] = DIAGRAMS_DIR / "er-diagram.png"
    mapping["4.23"] = DIAGRAMS_DIR / "erc1155-architecture.png"
    mapping["4.24"] = DIAGRAMS_DIR / "erc20-token-flow.png"
    mapping["4.25"] = DIAGRAMS_DIR / "property-nft-architecture.png"
    mapping["4.26"] = DIAGRAMS_DIR / "share-transfer-architecture.png"
    mapping["4.27"] = DIAGRAMS_DIR / "use-case-diagram.png"
    mapping["4.28"] = DIAGRAMS_DIR / "mobile-workflow.png"
    mapping["4.29"] = DIAGRAMS_DIR / "wireframe-dashboard.png"
    mapping["4.30"] = DIAGRAMS_DIR / "wireframe-property-registration.png"
    mapping["4.31"] = DIAGRAMS_DIR / "wireframe-yield-agreement.png"
    mapping["4.32"] = DIAGRAMS_DIR / "wireframe-mobile-dashboard.png"
    mapping["4.33"] = DIAGRAMS_DIR / "wireframe-mobile-yield-agreement.png"
    
    # Section 5 screenshots
    mapping["5.1"] = GENERATED_SCREENSHOTS / "docker-desktop-dashboard.png"
    mapping["5.2"] = GENERATED_SCREENSHOTS / "grafana-dashboard.png"
    mapping["5.3"] = GENERATED_SCREENSHOTS / "prometheus-targets.png"
    mapping["5.4"] = GENERATED_SCREENSHOTS / "backup-directory.png"
    mapping["5.5"] = PNG_DIR / "foundry-smart-contract-project-structure.drawio.png"
    mapping["5.6"] = GENERATED_SCREENSHOTS / "gas-report.png"
    mapping["5.7"] = GENERATED_SCREENSHOTS / "erc7201-storage-layout.png"
    mapping["5.8"] = GENERATED_SCREENSHOTS / "bytecode-size-verification.png"
    mapping["5.9"] = GENERATED_SCREENSHOTS / "property-registration-form.png"
    mapping["5.10"] = GENERATED_SCREENSHOTS / "yield-agreement-interface.png"
    mapping["5.11"] = GENERATED_SCREENSHOTS / "analytics-dashboard.png"
    mapping["5.12"] = GENERATED_SCREENSHOTS / "smart-contract-deployment.png"
    mapping["5.13"] = GENERATED_SCREENSHOTS / "polygonscan-confirmation.png"
    mapping["5.14"] = GENERATED_SCREENSHOTS / "diamond-test-results.png"
    mapping["5.15"] = GENERATED_SCREENSHOTS / "simulation-test-results.png"
    mapping["5.16"] = GENERATED_SCREENSHOTS / "gas-comparison-report.png"
    mapping["5.17"] = GENERATED_SCREENSHOTS / "variance-tracking.png"
    mapping["5.18"] = DIAGRAMS_DIR / "yield-tokenization-flow.png"
    
    # Section 6 charts
    mapping["6.1"] = GENERATED_CHARTS / "fig_6_1_token_standard_gas_comparison.png"
    mapping["6.2"] = GENERATED_CHARTS / "fig_6_2_batch_operation_scaling.png"
    mapping["6.3"] = GENERATED_CHARTS / "fig_6_3_amoy_anvil_variance_heatmap.png"
    mapping["6.4"] = GENERATED_CHARTS / "fig_6_4_volatile_simulation_radar.png"
    mapping["6.5"] = GENERATED_CHARTS / "fig_6_5_diamond_overhead_boxplot.png"
    mapping["6.6"] = GENERATED_CHARTS / "fig_6_6_load_testing_scatter.png"
    mapping["6.7"] = GENERATED_CHARTS / "fig_6_7_gas_cost_projections.png"
    mapping["6.8"] = GENERATED_CHARTS / "fig_6_8_test_pass_rate_evolution.png"
    mapping["6.9"] = GENERATED_CHARTS / "architecture-comparison.png"
    mapping["6.10"] = GENERATED_CHARTS / "user-workflow-diagrams.png"
    mapping["6.11"] = DIAGRAMS_DIR / "load-test-system-architecture.png"
    mapping["6.12"] = DIAGRAMS_DIR / "load-test-workflow.png"
    mapping["6.13"] = DIAGRAMS_DIR / "load-testing-results.png"
    mapping["6.14"] = DIAGRAMS_DIR / "load-test-gas-scaling.png"
    mapping["6.15"] = DIAGRAMS_DIR / "load-test-shareholder-distribution.png"
    mapping["6.16"] = DIAGRAMS_DIR / "load-test-success-rates.png"
    mapping["6.17"] = DIAGRAMS_DIR / "load-test-latency-distribution.png"
    mapping["6.18"] = DIAGRAMS_DIR / "load-test-metrics-summary.png"
    mapping["6.19"] = DIAGRAMS_DIR / "load-test-cost-comparison.png"
    mapping["6.20"] = DIAGRAMS_DIR / "load-test-erc1155-efficiency.png"
    mapping["6.21"] = DIAGRAMS_DIR / "load-test-token-recycling.png"
    mapping["6.22"] = DIAGRAMS_DIR / "load-test-restriction-enforcement.png"
    mapping["6.23"] = DIAGRAMS_DIR / "load-test-restriction-overhead.png"
    mapping["6.24"] = DIAGRAMS_DIR / "load-test-rq-validation.png"
    
    # Section 7 charts
    mapping["7.1"] = SURVEY_CHARTS / "fig_demographic_overview.png"
    mapping["7.2"] = SURVEY_CHARTS / "fig_tokenisation_interest_analysis.png"
    mapping["7.3"] = SURVEY_CHARTS / "fig_correlation_matrix.png"
    mapping["7.4"] = SURVEY_CHARTS / "fig_interview_demographics.png"
    mapping["7.5"] = SURVEY_CHARTS / "fig_thematic_code_frequencies.png"
    mapping["7.11"] = GENERATED_CHARTS / "token-decision-framework.png"
    mapping["7.12"] = GENERATED_CHARTS / "fig_7_7_volatile_recovery_comparison.png"
    mapping["7.13"] = GENERATED_CHARTS / "fig_7_8_testnet_local_comparison.png"
    
    # Survey figures (S.1-S.6)
    mapping["S.1"] = SURVEY_CHARTS / "fig_demographic_overview.png"
    mapping["S.2"] = SURVEY_CHARTS / "fig_tokenisation_interest_analysis.png"
    mapping["S.3"] = SURVEY_CHARTS / "fig_correlation_matrix.png"
    mapping["S.4"] = SURVEY_CHARTS / "fig_cluster_analysis.png"
    mapping["S.5"] = SURVEY_CHARTS / "fig_motivations_concerns.png"
    mapping["S.6"] = SURVEY_CHARTS / "fig_feature_importance.png"
    
    # Interview figures (I.1-I.4)
    mapping["I.1"] = SURVEY_CHARTS / "fig_interview_demographics.png"
    mapping["I.2"] = SURVEY_CHARTS / "fig_thematic_code_frequencies.png"
    mapping["I.3"] = SURVEY_CHARTS / "fig_landlord_fintech_comparison.png"
    mapping["I.4"] = SURVEY_CHARTS / "fig_likert_distributions.png"
    
    # Appendix F wireframes (F.1-F.19)
    mapping["F.1"] = DIAGRAMS_DIR / "wireframe-analytics-dashboard.png"
    mapping["F.2"] = DIAGRAMS_DIR / "wireframe-governance.png"
    mapping["F.3"] = DIAGRAMS_DIR / "wireframe-governance-proposal-detail.png"
    mapping["F.4"] = DIAGRAMS_DIR / "wireframe-kyc.png"
    mapping["F.5"] = DIAGRAMS_DIR / "wireframe-kyc-admin.png"
    mapping["F.6"] = DIAGRAMS_DIR / "wireframe-marketplace.png"
    mapping["F.7"] = DIAGRAMS_DIR / "wireframe-portfolio.png"
    mapping["F.8"] = DIAGRAMS_DIR / "wireframe-properties-list.png"
    mapping["F.9"] = DIAGRAMS_DIR / "wireframe-yield-agreement-detail.png"
    mapping["F.10"] = DIAGRAMS_DIR / "wireframe-yield-agreements-list.png"
    mapping["F.11"] = DIAGRAMS_DIR / "wireframe-mobile-analytics.png"
    mapping["F.12"] = DIAGRAMS_DIR / "wireframe-mobile-governance.png"
    mapping["F.13"] = DIAGRAMS_DIR / "wireframe-mobile-kyc.png"
    mapping["F.14"] = DIAGRAMS_DIR / "wireframe-mobile-kyc-admin.png"
    mapping["F.15"] = DIAGRAMS_DIR / "wireframe-mobile-marketplace.png"
    mapping["F.16"] = DIAGRAMS_DIR / "wireframe-mobile-portfolio.png"
    mapping["F.17"] = DIAGRAMS_DIR / "wireframe-mobile-properties-list.png"
    mapping["F.18"] = DIAGRAMS_DIR / "wireframe-mobile-register-property.png"
    mapping["F.19"] = DIAGRAMS_DIR / "wireframe-mobile-yield-agreements-list.png"
    
    # Appendix G charts (G.1-G.5)
    mapping["G.1"] = SURVEY_CHARTS / "fig_cluster_analysis.png"
    mapping["G.2"] = SURVEY_CHARTS / "fig_feature_importance.png"
    mapping["G.3"] = SURVEY_CHARTS / "fig_motivations_concerns.png"
    mapping["G.4"] = SURVEY_CHARTS / "fig_landlord_fintech_comparison.png"
    mapping["G.5"] = SURVEY_CHARTS / "fig_likert_distributions.png"
    
    return mapping

def insert_images_into_docx():
    """Insert images into the Word document at figure placeholders."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Installing python-docx...")
        subprocess.run(["pip", "install", "--break-system-packages", "python-docx"], 
                      capture_output=True)
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    print("\nLoading Word document...")
    doc = Document(str(TEMP_DOCX))
    
    mapping = get_image_mapping()
    
    # Pattern to match figure titles - handles both bold and plain text
    figure_pattern = re.compile(r'^(?:\*\*)?Figure\s+([A-Z]?\d+\.?\d*):(.+?)(?:\*\*)?$')
    
    figures_found = 0
    figures_inserted = 0
    figures_missing = []
    
    print("\nSearching for figure placeholders...")
    
    paragraphs_to_process = list(doc.paragraphs)
    
    for i, para in enumerate(paragraphs_to_process):
        text = para.text.strip()
        
        # Check if this is a figure title
        match = figure_pattern.match(text)
        if match:
            fig_num = match.group(1)
            fig_title = match.group(2).strip()
            figures_found += 1
            
            # Clean the title - remove placeholder descriptions
            clean_title = fig_title
            if '[' in clean_title:
                clean_title = clean_title.split('[')[0].strip()
            
            # Get the image path
            if fig_num in mapping:
                img_path = mapping[fig_num]
                
                if img_path.exists():
                    # Clear the paragraph and insert image
                    para.clear()
                    
                    # Add the image
                    run = para.add_run()
                    try:
                        run.add_picture(str(img_path), width=Inches(6.0))
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        figures_inserted += 1
                        print(f"  ✓ Inserted Figure {fig_num}: {clean_title[:50]}...")
                    except Exception as e:
                        print(f"  ✗ Error inserting Figure {fig_num}: {e}")
                        # Restore the original text
                        para.add_run(f"Figure {fig_num}: {clean_title}")
                else:
                    print(f"  ✗ Image not found for Figure {fig_num}: {img_path}")
                    figures_missing.append(f"{fig_num}: {img_path}")
            else:
                print(f"  ? No mapping for Figure {fig_num}")
                figures_missing.append(f"{fig_num}: No mapping")
    
    print(f"\n--- Summary ---")
    print(f"Figures found: {figures_found}")
    print(f"Figures inserted: {figures_inserted}")
    print(f"Figures missing/failed: {len(figures_missing)}")
    
    if figures_missing:
        print("\nMissing figures:")
        for fig in figures_missing[:20]:
            print(f"  - {fig}")
        if len(figures_missing) > 20:
            print(f"  ... and {len(figures_missing) - 20} more")
    
    # Save the document
    print(f"\nSaving document to {OUTPUT_DOCX}...")
    doc.save(str(OUTPUT_DOCX))
    print("Document saved successfully!")
    
    return figures_inserted

def cleanup():
    """Remove temporary files."""
    if TEMP_DOCX.exists():
        os.remove(TEMP_DOCX)
        print(f"Cleaned up temporary file: {TEMP_DOCX}")

def main():
    print("=" * 60)
    print("Rebuilding DissertationProgressFinal Word Document")
    print("Version 2.6 - With restructured Section 4.3 and 4.5")
    print("=" * 60)
    
    # Step 1: Convert markdown to docx
    if not convert_markdown_to_docx():
        print("Failed to convert markdown to Word. Exiting.")
        return
    
    # Step 2: Insert images
    figures_inserted = insert_images_into_docx()
    
    # Step 3: Cleanup
    cleanup()
    
    print("\n" + "=" * 60)
    print(f"COMPLETE: Created {OUTPUT_DOCX}")
    print(f"Total figures inserted: {figures_inserted}")
    print("=" * 60)

if __name__ == "__main__":
    main()


