#!/usr/bin/env python3
"""
Insert all figures into DissertationProgressFinal.docx - v2
Uses exact filename matching from placeholder text.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from pathlib import Path

# Paths
doc_path = "DissertationProgressFinal.docx"
output_path = "DissertationProgressFinal_WithFigures.docx"

# Image directories
drawio_dir = Path("rendered_diagrams_drawio")
survey_dir = Path("survey_interview_charts")
charts_dir = Path("generated_charts")

# Load document
print(f"Loading {doc_path}...")
doc = Document(doc_path)

# Build complete image inventory
all_images = {}

# Draw.io diagrams
for f in drawio_dir.glob("*.png"):
    all_images[f.stem] = f
    all_images[f.stem.replace("-", "_")] = f  # Also map with underscores
print(f"Found {len(list(drawio_dir.glob('*.png')))} Draw.io PNGs")

# Survey/Interview charts
for f in survey_dir.glob("*.png"):
    all_images[f.stem] = f
    all_images[f.stem.replace("_", "-")] = f
print(f"Found {len(list(survey_dir.glob('*.png')))} Survey/Interview PNGs")

# Generated charts
if charts_dir.exists():
    for f in charts_dir.glob("*.png"):
        all_images[f.stem] = f
    print(f"Found {len(list(charts_dir.glob('*.png')))} Generated chart PNGs")

print(f"Total image mappings: {len(all_images)}")

# Pattern to extract .mmd filename from placeholder
mmd_pattern = re.compile(r'([a-z0-9-]+)\.mmd', re.IGNORECASE)
placeholder_pattern = re.compile(r'\[PLACEHOLDER[:\s]*(.+?)\]', re.IGNORECASE)

figures_inserted = 0
figures_not_found = []

print("\nProcessing document paragraphs...")

for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    if '[PLACEHOLDER' in text.upper():
        # Try to extract .mmd filename
        mmd_match = mmd_pattern.search(text)
        
        image_path = None
        description = text
        
        if mmd_match:
            # Extract filename without extension
            filename = mmd_match.group(1)
            description = filename
            
            # Look up in our image inventory
            if filename in all_images:
                image_path = all_images[filename]
        
        if image_path and image_path.exists():
            # Clear paragraph and insert image
            para.clear()
            run = para.add_run()
            
            # Insert image (5.5 inches wide)
            run.add_picture(str(image_path), width=Inches(5.5))
            
            figures_inserted += 1
            print(f"  ✓ {filename} -> {image_path.name}")
        else:
            # Try to find by keywords in the placeholder text
            text_lower = text.lower()
            found = False
            
            for name, path in all_images.items():
                if name in text_lower or name.replace("-", " ") in text_lower:
                    para.clear()
                    run = para.add_run()
                    run.add_picture(str(path), width=Inches(5.5))
                    figures_inserted += 1
                    print(f"  ✓ (keyword) {name} -> {path.name}")
                    found = True
                    break
            
            if not found:
                figures_not_found.append(text[:80])
                print(f"  ✗ Not found: {text[:60]}...")

# Save document
print(f"\nSaving to {output_path}...")
doc.save(output_path)

print(f"\n{'='*60}")
print(f"SUMMARY")
print(f"{'='*60}")
print(f"Figures inserted: {figures_inserted}")
print(f"Figures not found: {len(figures_not_found)}")

if figures_not_found:
    print(f"\nMissing figures ({len(figures_not_found)}):")
    for desc in figures_not_found[:30]:
        print(f"  - {desc}")
    if len(figures_not_found) > 30:
        print(f"  ... and {len(figures_not_found) - 30} more")

print(f"\n✅ Document saved: {output_path}")
