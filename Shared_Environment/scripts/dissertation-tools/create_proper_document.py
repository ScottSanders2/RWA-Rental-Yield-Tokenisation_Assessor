#!/usr/bin/env python3
"""
Create a proper Word document with correctly captioned figures.
Strategy:
1. Convert markdown to Word (without images)
2. Find each Figure placeholder in the Word doc
3. Replace placeholder with image + proper caption
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

# Complete figure mapping
FIGURE_MAPPING = {
    # Section 4 - System Design (32 figures)
    "Figure 4.1": ("rendered_diagrams_drawio/docker-architecture-logical.png", "Multi-Environment Docker Architecture (Logical View)"),
    "Figure 4.2": ("rendered_diagrams_drawio/docker-architecture-physical.png", "Multi-Environment Docker Architecture (Physical View)"),
    "Figure 4.3": ("rendered_diagrams_drawio/monitoring-architecture.png", "Monitoring Architecture Diagram"),
    "Figure 4.4": ("rendered_diagrams_drawio/mobile-architecture.png", "Mobile Application Architecture"),
    "Figure 4.5": ("rendered_diagrams_drawio/backend-api-architecture.png", "Backend API Architecture"),
    "Figure 4.6": ("rendered_diagrams_drawio/property-owner-workflow.png", "Property Owner Workflow"),
    "Figure 4.7": ("rendered_diagrams_drawio/yield-architecture.png", "Yield Tokenisation Architecture"),
    "Figure 4.8": ("rendered_diagrams_drawio/yield-tokenization-flow.png", "Complete Yield Tokenization Workflow"),
    "Figure 4.9": ("rendered_diagrams_drawio/repayment-processing-flow.png", "Repayment Processing Decision Tree"),
    "Figure 4.10": ("rendered_diagrams_drawio/pooling-mechanism-flow.png", "Pooling Mechanism Workflow"),
    "Figure 4.11": ("rendered_diagrams_drawio/default-handling-flow.png", "Default Handling Workflow"),
    "Figure 4.12": ("rendered_diagrams_drawio/erc7201-storage-layout.png", "ERC-7201 Storage Layout"),
    "Figure 4.13": ("rendered_diagrams_drawio/analytics-architecture.png", "Analytics Architecture (The Graph Protocol)"),
    "Figure 4.14": ("rendered_diagrams_drawio/subgraph-entity-relationships.png", "Subgraph Entity Relationships"),
    "Figure 4.15": ("rendered_diagrams_drawio/governance-architecture.png", "Governance Architecture"),
    "Figure 4.16": ("rendered_diagrams_drawio/governance-proposal-flow.png", "Governance Proposal Flow"),
    "Figure 4.17": ("rendered_diagrams_drawio/secondary-market-architecture.png", "Secondary Market Architecture"),
    "Figure 4.18": ("rendered_diagrams_drawio/transfer-restriction-flow.png", "Transfer Restriction Flow"),
    "Figure 4.19": ("rendered_diagrams_drawio/kyc-architecture.png", "KYC Architecture"),
    "Figure 4.20": ("rendered_diagrams_drawio/kyc-verification-workflow.png", "KYC Verification Workflow"),
    "Figure 4.21": ("rendered_diagrams_drawio/er-diagram.png", "Entity-Relationship Diagram"),
    "Figure 4.22": ("rendered_diagrams_drawio/erc1155-architecture.png", "ERC-1155 Architecture"),
    "Figure 4.23": ("rendered_diagrams_drawio/erc20-token-flow.png", "ERC-20 Token Flow"),
    "Figure 4.24": ("rendered_diagrams_drawio/property-nft-architecture.png", "Property NFT Architecture"),
    "Figure 4.25": ("rendered_diagrams_drawio/share-transfer-architecture.png", "Share Transfer Architecture"),
    "Figure 4.26": ("rendered_diagrams_drawio/use-case-diagram.png", "Use Case Diagram"),
    "Figure 4.27": ("rendered_diagrams_drawio/mobile-workflow.png", "Mobile Workflow"),
    "Figure 4.28": ("rendered_diagrams_drawio/dashboard-wireframe.png", "Dashboard Wireframe"),
    "Figure 4.29": ("rendered_diagrams_drawio/property-registration-wireframe.png", "Property Registration Wireframe"),
    "Figure 4.30": ("rendered_diagrams_drawio/yield-agreement-wireframe.png", "Yield Agreement Wireframe"),
    "Figure 4.31": ("rendered_diagrams_drawio/mobile-dashboard-wireframe.png", "Mobile Dashboard Wireframe"),
    "Figure 4.32": ("rendered_diagrams_drawio/mobile-yield-agreement-wireframe.png", "Mobile Yield Agreement Wireframe"),
    
    # Section 5 - Implementation (23 figures)
    "Figure 5.1": ("generated_screenshots/docker-desktop-dashboard.png", "Docker Desktop Container Status"),
    "Figure 5.2": ("generated_screenshots/grafana-dashboard.png", "Grafana Monitoring Dashboard"),
    "Figure 5.3": ("generated_screenshots/prometheus-targets.png", "Prometheus Targets"),
    "Figure 5.4": ("generated_screenshots/node-exporter-metrics.png", "Node-Exporter Metrics"),
    "Figure 5.5": ("generated_screenshots/multi-environment-orchestration.png", "Multi-Environment Orchestration"),
    "Figure 5.6": ("generated_screenshots/backup-directory.png", "Backup Directory Structure"),
    "Figure 5.7": ("generated_screenshots/foundry-project-structure.png", "Foundry Project Structure"),
    "Figure 5.8": ("generated_screenshots/forge-test-output.png", "Forge Test Output"),
    "Figure 5.9": ("generated_screenshots/gas-report.png", "Gas Report"),
    "Figure 5.10": ("generated_screenshots/anvil-container-logs.png", "Anvil Container Logs"),
    "Figure 5.11": ("generated_screenshots/deployment-script-output.png", "Deployment Script Output"),
    "Figure 5.12": ("generated_screenshots/erc7201-storage-layout.png", "ERC-7201 Storage Layout"),
    "Figure 5.13": ("generated_screenshots/bytecode-size-verification.png", "Bytecode Size Verification"),
    "Figure 5.14": ("generated_screenshots/amoy-testnet-deployment.png", "Diamond Deployment to Amoy"),
    "Figure 5.15": ("generated_screenshots/polygonscan-transaction.png", "PolygonScan Transaction Confirmation"),
    "Figure 5.16": ("generated_screenshots/diamond-test-results.png", "Diamond Architecture Test Results"),
    "Figure 5.17": ("generated_screenshots/property-registration-form.png", "Property Registration Form"),
    "Figure 5.18": ("generated_screenshots/yield-agreement-creation.png", "Yield Agreement Creation Interface"),
    "Figure 5.19": ("generated_screenshots/analytics-dashboard.png", "Analytics Dashboard"),
    "Figure 5.20": ("generated_screenshots/simulation-test-results.png", "Simulation Test Results"),
    "Figure 5.21": ("generated_screenshots/gas-comparison-report.png", "Gas Comparison Report"),
    "Figure 5.22": ("generated_screenshots/variance-tracking-metrics.png", "Variance Tracking Metrics"),
    "Figure 5.23": ("rendered_diagrams_drawio/yield-tokenization-flow.png", "Yield Tokenization Flowchart"),
    
    # Section 6 - Testing (24 figures)
    "Figure 6.1": ("generated_charts/token-standard-gas-comparison.png", "Token Standard Gas Comparison (Bar Chart)"),
    "Figure 6.2": ("generated_charts/batch-operation-scaling.png", "Batch Operation Scaling Curve (Line Chart)"),
    "Figure 6.3": ("generated_charts/amoy-anvil-variance-heatmap.png", "Amoy vs Anvil Variance Heatmap"),
    "Figure 6.4": ("generated_charts/volatile-simulation-recovery.png", "Volatile Simulation Recovery Percentages (Radar Chart)"),
    "Figure 6.5": ("generated_charts/diamond-call-overhead.png", "Diamond Architecture Call Overhead Analysis (Box Plot)"),
    "Figure 6.6": ("generated_charts/load-testing-throughput.png", "Load Testing Throughput vs Shareholder Count (Scatter Plot)"),
    "Figure 6.7": ("generated_charts/gas-cost-projections.png", "Gas Cost Projections for Mainnet Scenarios (Multi-Series Line Chart)"),
    "Figure 6.8": ("generated_charts/test-pass-rate-evolution.png", "Test Pass Rate Evolution (Stacked Area Chart)"),
    "Figure 6.9": ("rendered_diagrams_drawio/architecture-comparison.png", "Architecture Comparison (Side-by-Side Diagram)"),
    "Figure 6.10": ("rendered_diagrams_drawio/user-workflow-diagrams.png", "User Workflow Diagrams (Sequence Diagrams)"),
    "Figure 6.11": ("rendered_diagrams_drawio/load-test-architecture.png", "Load Test System Architecture"),
    "Figure 6.12": ("rendered_diagrams_drawio/load-test-workflow.png", "Load Test Workflow"),
    "Figure 6.13": ("rendered_diagrams_drawio/load-test-results-summary.png", "Load Testing Results Summary"),
    "Figure 6.14": ("rendered_diagrams_drawio/load-test-gas-scaling.png", "Load Test Gas Scaling"),
    "Figure 6.15": ("rendered_diagrams_drawio/load-test-shareholder-distribution.png", "Load Test Shareholder Distribution"),
    "Figure 6.16": ("rendered_diagrams_drawio/load-test-success-rates.png", "Load Test Success Rates"),
    "Figure 6.17": ("rendered_diagrams_drawio/load-test-latency-distribution.png", "Load Test Latency Distribution"),
    "Figure 6.18": ("rendered_diagrams_drawio/load-test-metrics-summary.png", "Load Test Metrics Summary"),
    "Figure 6.19": ("rendered_diagrams_drawio/load-test-cost-comparison.png", "Load Test Cost Comparison"),
    "Figure 6.20": ("rendered_diagrams_drawio/load-test-erc1155-efficiency.png", "Load Test ERC-1155 Efficiency"),
    "Figure 6.21": ("rendered_diagrams_drawio/load-test-token-recycling.png", "Load Test Token Recycling"),
    "Figure 6.22": ("rendered_diagrams_drawio/load-test-restriction-enforcement.png", "Load Test Restriction Enforcement"),
    "Figure 6.23": ("rendered_diagrams_drawio/load-test-restriction-overhead.png", "Load Test Restriction Overhead"),
    "Figure 6.24": ("rendered_diagrams_drawio/load-test-research-validation.png", "Load Test Research Question Validation"),
    
    # Section 7 - Evaluation (8 figures)
    "Figure 7.1": ("survey_interview_charts/survey_demographics_composite.png", "Survey Respondent Demographics Overview (4-Panel Composite)"),
    "Figure 7.2": ("survey_interview_charts/tokenisation_interest_distribution.png", "Tokenisation Interest Distribution and Cross-Tabulation"),
    "Figure 7.3": ("survey_interview_charts/spearman_correlation_matrix.png", "Spearman Rank Correlation Matrix"),
    "Figure 7.4": ("survey_interview_charts/interview_demographics_composite.png", "Interview Participant Demographics (4-Panel Composite)"),
    "Figure 7.5": ("survey_interview_charts/theme_frequency_analysis.png", "Theme Frequency Analysis"),
    "Figure 7.6": ("rendered_diagrams_drawio/token-standard-decision-framework.png", "Token Standard Decision Framework"),
    "Figure 7.7": ("generated_charts/volatile-simulation-recovery.png", "Volatile Market Simulation Recovery Comparison"),
    "Figure 7.8": ("generated_charts/testnet-local-performance.png", "Testnet vs Local Performance Comparison"),
}

# Appendix D wireframes
APPENDIX_D = {
    "Figure D.1": ("rendered_diagrams_drawio/analytics-dashboard-wireframe.png", "Analytics Dashboard Wireframe"),
    "Figure D.2": ("rendered_diagrams_drawio/compliance-dashboard-wireframe.png", "Compliance Dashboard Wireframe"),
    "Figure D.3": ("rendered_diagrams_drawio/error-states-wireframe.png", "Error States Wireframe"),
    "Figure D.4": ("rendered_diagrams_drawio/governance-dashboard-wireframe.png", "Governance Dashboard Wireframe"),
    "Figure D.5": ("rendered_diagrams_drawio/investor-dashboard-wireframe.png", "Investor Dashboard Wireframe"),
    "Figure D.6": ("rendered_diagrams_drawio/investor-portfolio-wireframe.png", "Investor Portfolio Wireframe"),
    "Figure D.7": ("rendered_diagrams_drawio/kyc-verification-wireframe.png", "KYC Verification Wireframe"),
    "Figure D.8": ("rendered_diagrams_drawio/loading-states-wireframe.png", "Loading States Wireframe"),
    "Figure D.9": ("rendered_diagrams_drawio/mobile-analytics-wireframe.png", "Mobile Analytics Wireframe"),
    "Figure D.10": ("rendered_diagrams_drawio/mobile-investor-wireframe.png", "Mobile Investor Wireframe"),
    "Figure D.11": ("rendered_diagrams_drawio/mobile-notifications-wireframe.png", "Mobile Notifications Wireframe"),
    "Figure D.12": ("rendered_diagrams_drawio/mobile-property-details-wireframe.png", "Mobile Property Details Wireframe"),
    "Figure D.13": ("rendered_diagrams_drawio/mobile-secondary-market-wireframe.png", "Mobile Secondary Market Wireframe"),
    "Figure D.14": ("rendered_diagrams_drawio/notification-center-wireframe.png", "Notification Center Wireframe"),
    "Figure D.15": ("rendered_diagrams_drawio/property-details-wireframe.png", "Property Details Wireframe"),
    "Figure D.16": ("rendered_diagrams_drawio/secondary-market-wireframe.png", "Secondary Market Wireframe"),
    "Figure D.17": ("rendered_diagrams_drawio/settings-wireframe.png", "Settings Wireframe"),
    "Figure D.18": ("rendered_diagrams_drawio/transaction-history-wireframe.png", "Transaction History Wireframe"),
    "Figure D.19": ("rendered_diagrams_drawio/yield-distribution-wireframe.png", "Yield Distribution Wireframe"),
}

# Appendix E charts
APPENDIX_E = {
    "Figure E.1": ("survey_interview_charts/respondent_cluster_analysis.png", "Respondent Cluster Analysis"),
    "Figure E.2": ("survey_interview_charts/platform_feature_importance.png", "Platform Feature Importance Rankings"),
    "Figure E.3": ("survey_interview_charts/investment_motivations_concerns.png", "Investment Motivations and Concerns"),
    "Figure E.4": ("survey_interview_charts/landlord_fintech_comparison.png", "Landlord vs FinTech Expert Comparison"),
    "Figure E.5": ("survey_interview_charts/likert_scale_distributions.png", "Likert Scale Response Distributions"),
}

# Combine all mappings
ALL_FIGURES = {**FIGURE_MAPPING, **APPENDIX_D, **APPENDIX_E}

def insert_image_with_caption(doc, para_idx, figure_num, image_path, caption_text):
    """Insert an image with a proper caption after the specified paragraph"""
    
    if not os.path.exists(image_path):
        print(f"  WARNING: Image not found: {image_path}")
        return False
    
    # Get the paragraph
    para = doc.paragraphs[para_idx]
    
    # Clear the placeholder text
    para.clear()
    
    # Add the image
    run = para.add_run()
    try:
        # Determine image width based on type
        if 'wireframe' in image_path.lower() or 'screenshot' in image_path.lower():
            width = Inches(5.5)
        else:
            width = Inches(6.0)
        
        run.add_picture(image_path, width=width)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add caption as a new paragraph after the image
        # We need to insert after this paragraph
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_para.add_run(f"{figure_num}: {caption_text}")
        caption_run.bold = True
        caption_run.font.size = Pt(10)
        
        # Move the caption paragraph to right after the image paragraph
        # This is tricky with python-docx, we'll handle it differently
        
        return True
    except Exception as e:
        print(f"  ERROR inserting {image_path}: {e}")
        return False

def find_figure_placeholders(doc):
    """Find all paragraphs that contain Figure X.X placeholders"""
    placeholders = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Look for Figure X.X pattern (bold figure labels)
        match = re.match(r'^(Figure\s+[0-9A-Z]+\.[0-9]+):\s*(.+)$', text)
        if match:
            figure_num = match.group(1)
            title = match.group(2)
            placeholders.append({
                'para_idx': i,
                'figure_num': figure_num,
                'title': title,
                'full_text': text
            })
    
    return placeholders

def create_proper_document():
    """Create a new document with properly captioned figures"""
    
    print("=" * 80)
    print("CREATING PROPERLY CAPTIONED DOCUMENT")
    print("=" * 80)
    
    # First, verify all image files exist
    print("\n1. Verifying image files...")
    missing_images = []
    for fig_num, (img_path, caption) in ALL_FIGURES.items():
        if not os.path.exists(img_path):
            missing_images.append((fig_num, img_path))
            print(f"  MISSING: {fig_num} -> {img_path}")
    
    if missing_images:
        print(f"\n  WARNING: {len(missing_images)} images are missing!")
    else:
        print(f"  All {len(ALL_FIGURES)} images verified.")
    
    # Open the document
    print("\n2. Opening document...")
    doc = Document('DissertationProgressFinal_v2.1.docx')
    
    # Find all figure placeholders
    print("\n3. Finding figure placeholders...")
    placeholders = find_figure_placeholders(doc)
    print(f"  Found {len(placeholders)} figure placeholders")
    
    for p in placeholders[:10]:
        print(f"    {p['figure_num']}: {p['title'][:50]}...")
    
    # Now we need to match placeholders to images and insert them
    print("\n4. Matching placeholders to images...")
    matched = 0
    unmatched = []
    
    for p in placeholders:
        fig_num = p['figure_num']
        if fig_num in ALL_FIGURES:
            matched += 1
        else:
            unmatched.append(fig_num)
    
    print(f"  Matched: {matched}")
    print(f"  Unmatched: {len(unmatched)}")
    if unmatched:
        for u in unmatched[:10]:
            print(f"    {u}")
    
    return placeholders

if __name__ == "__main__":
    placeholders = create_proper_document()
