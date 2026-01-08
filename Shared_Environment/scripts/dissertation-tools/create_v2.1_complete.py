#!/usr/bin/env python3
"""
Create DissertationProgressFinal_v2.1.docx with:
1. All 87 main figures (already in v2.0)
2. 19 Appendix D wireframes (ADDING NOW)
3. Additional survey/interview charts for Appendices (ADDING NOW)
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from pathlib import Path

# Load v2.0
doc = Document("DissertationProgressFinal_v2.0.docx")
output_path = "DissertationProgressFinal_v2.1.docx"

print("=" * 70)
print("CREATING v2.1 WITH ALL APPENDIX FIGURES")
print("=" * 70)

# ============================================================
# APPENDIX D: Additional UI Wireframes (19 figures)
# ============================================================
print("\n>>> Adding Appendix D: UI Wireframes")

doc.add_page_break()
heading = doc.add_heading('Appendix D: Additional UI Wireframes', level=1)

# Web wireframes
doc.add_heading('D.1 Web Application Wireframes', level=2)

web_wireframes = [
    ("wireframe-analytics-dashboard", "Analytics Dashboard"),
    ("wireframe-governance", "Governance Dashboard"),
    ("wireframe-governance-proposal-detail", "Governance Proposal Detail"),
    ("wireframe-kyc", "KYC Verification"),
    ("wireframe-kyc-admin", "KYC Admin Dashboard"),
    ("wireframe-marketplace", "Marketplace"),
    ("wireframe-portfolio", "Portfolio Dashboard"),
    ("wireframe-properties-list", "Properties List"),
    ("wireframe-yield-agreement-detail", "Yield Agreement Detail"),
    ("wireframe-yield-agreements-list", "Yield Agreements List"),
]

fig_count = 1
for img_name, caption in web_wireframes:
    img_path = Path(f"rendered_diagrams_drawio/{img_name}.png")
    if img_path.exists():
        # Caption
        cap_para = doc.add_paragraph()
        cap_para.add_run(f"Figure D.{fig_count}: {caption}").bold = True
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Image
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        
        print(f"  ✓ Figure D.{fig_count}: {caption}")
        fig_count += 1

# Mobile wireframes
doc.add_heading('D.2 Mobile Application Wireframes', level=2)

mobile_wireframes = [
    ("wireframe-mobile-analytics", "Mobile Analytics"),
    ("wireframe-mobile-governance", "Mobile Governance"),
    ("wireframe-mobile-kyc", "Mobile KYC"),
    ("wireframe-mobile-kyc-admin", "Mobile KYC Admin"),
    ("wireframe-mobile-marketplace", "Mobile Marketplace"),
    ("wireframe-mobile-portfolio", "Mobile Portfolio"),
    ("wireframe-mobile-properties-list", "Mobile Properties List"),
    ("wireframe-mobile-register-property", "Mobile Property Registration"),
    ("wireframe-mobile-yield-agreements-list", "Mobile Yield Agreements List"),
]

for img_name, caption in mobile_wireframes:
    img_path = Path(f"rendered_diagrams_drawio/{img_name}.png")
    if img_path.exists():
        cap_para = doc.add_paragraph()
        cap_para.add_run(f"Figure D.{fig_count}: {caption}").bold = True
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        
        print(f"  ✓ Figure D.{fig_count}: {caption}")
        fig_count += 1

appendix_d_count = fig_count - 1
print(f"\nAppendix D complete: {appendix_d_count} wireframes added")

# ============================================================
# APPENDIX E: Additional Survey Charts
# ============================================================
print("\n>>> Adding Appendix E: Additional Survey/Interview Charts")

doc.add_page_break()
doc.add_heading('Appendix E: Additional Research Charts', level=1)

doc.add_heading('E.1 Survey Analysis Charts', level=2)

survey_charts = [
    ("fig_cluster_analysis", "Respondent Cluster Analysis"),
    ("fig_feature_importance", "Platform Feature Importance Rankings"),
    ("fig_motivations_concerns", "Investment Motivations and Concerns"),
]

fig_count = 1
for img_name, caption in survey_charts:
    img_path = Path(f"survey_interview_charts/{img_name}.png")
    if img_path.exists():
        cap_para = doc.add_paragraph()
        cap_para.add_run(f"Figure E.{fig_count}: {caption}").bold = True
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        
        print(f"  ✓ Figure E.{fig_count}: {caption}")
        fig_count += 1

doc.add_heading('E.2 Interview Analysis Charts', level=2)

interview_charts = [
    ("fig_landlord_fintech_comparison", "Landlord vs FinTech Expert Comparison"),
    ("fig_likert_distributions", "Likert Scale Response Distributions"),
]

for img_name, caption in interview_charts:
    img_path = Path(f"survey_interview_charts/{img_name}.png")
    if img_path.exists():
        cap_para = doc.add_paragraph()
        cap_para.add_run(f"Figure E.{fig_count}: {caption}").bold = True
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        
        print(f"  ✓ Figure E.{fig_count}: {caption}")
        fig_count += 1

appendix_e_count = fig_count - 1
print(f"\nAppendix E complete: {appendix_e_count} charts added")

# Save
print(f"\nSaving to {output_path}...")
doc.save(output_path)

# Final count
from docx import Document as Doc2
doc2 = Doc2(output_path)
total_images = sum(1 for rel in doc2.part.rels.values() if "image" in rel.reltype)

print("\n" + "=" * 70)
print("v2.1 COMPLETE")
print("=" * 70)
print(f"Main document figures:    86 (87 references, 1 shared)")
print(f"Appendix D wireframes:    {appendix_d_count}")
print(f"Appendix E charts:        {appendix_e_count}")
print(f"TOTAL IMAGES:             {total_images}")
print(f"\nFile: {output_path}")
print(f"Size: {os.path.getsize(output_path) / 1024 / 1024:.1f} MB")
