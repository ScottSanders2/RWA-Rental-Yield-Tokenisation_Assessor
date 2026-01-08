#!/usr/bin/env python3
"""
Complete document rebuild with ALL figures properly captioned.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess
import re
import os

# Complete figure mapping with CORRECT filenames
FIGURES = {
    # Section 4 - System Design (32 figures)
    "Figure 4.1": "rendered_diagrams_drawio/docker-architecture-logical.png",
    "Figure 4.2": "rendered_diagrams_drawio/docker-architecture-physical.png",
    "Figure 4.3": "rendered_diagrams_drawio/monitoring-architecture.png",
    "Figure 4.4": "rendered_diagrams_drawio/mobile-architecture.png",
    "Figure 4.5": "rendered_diagrams_drawio/backend-api-architecture.png",
    "Figure 4.6": "rendered_diagrams_drawio/property-owner-workflow.png",
    "Figure 4.7": "rendered_diagrams_drawio/yield-architecture.png",
    "Figure 4.8": "rendered_diagrams_drawio/yield-tokenization-flow.png",
    "Figure 4.9": "rendered_diagrams_drawio/repayment-processing-flow.png",
    "Figure 4.10": "rendered_diagrams_drawio/pooling-mechanism-flow.png",
    "Figure 4.11": "rendered_diagrams_drawio/default-handling-flow.png",
    "Figure 4.12": "rendered_diagrams_drawio/storage-layout-diagram.png",
    "Figure 4.13": "rendered_diagrams_drawio/analytics-architecture.png",
    "Figure 4.14": "rendered_diagrams_drawio/subgraph-entity-relationships.png",
    "Figure 4.15": "rendered_diagrams_drawio/governance-architecture.png",
    "Figure 4.16": "rendered_diagrams_drawio/governance-proposal-flow.png",
    "Figure 4.17": "rendered_diagrams_drawio/secondary-market-architecture.png",
    "Figure 4.18": "rendered_diagrams_drawio/transfer-restriction-flow.png",
    "Figure 4.19": "rendered_diagrams_drawio/kyc-architecture.png",
    "Figure 4.20": "rendered_diagrams_drawio/kyc-workflow.png",
    "Figure 4.21": "rendered_diagrams_drawio/er-diagram.png",
    "Figure 4.22": "rendered_diagrams_drawio/erc1155-architecture.png",
    "Figure 4.23": "rendered_diagrams_drawio/erc20-token-flow.png",
    "Figure 4.24": "rendered_diagrams_drawio/property-nft-architecture.png",
    "Figure 4.25": "rendered_diagrams_drawio/share-transfer-architecture.png",
    "Figure 4.26": "rendered_diagrams_drawio/use-case-diagram.png",
    "Figure 4.27": "rendered_diagrams_drawio/mobile-workflow.png",
    "Figure 4.28": "rendered_diagrams_drawio/wireframe-dashboard.png",
    "Figure 4.29": "rendered_diagrams_drawio/wireframe-property-registration.png",
    "Figure 4.30": "rendered_diagrams_drawio/wireframe-yield-agreement.png",
    "Figure 4.31": "rendered_diagrams_drawio/wireframe-mobile-dashboard.png",
    "Figure 4.32": "rendered_diagrams_drawio/wireframe-mobile-yield-agreement.png",
    # Section 5 - Implementation (23 figures)
    "Figure 5.1": "generated_screenshots/docker-desktop-dashboard.png",
    "Figure 5.2": "generated_screenshots/grafana-dashboard.png",
    "Figure 5.3": "generated_screenshots/prometheus-targets.png",
    "Figure 5.4": "generated_screenshots/node-exporter-metrics.png",
    "Figure 5.5": "generated_screenshots/multi-env-orchestration.png",
    "Figure 5.6": "generated_screenshots/backup-directory.png",
    "Figure 5.7": "generated_screenshots/foundry-project-structure.png",
    "Figure 5.8": "generated_screenshots/forge-test-output.png",
    "Figure 5.9": "generated_screenshots/gas-report.png",
    "Figure 5.10": "generated_screenshots/anvil-container-logs.png",
    "Figure 5.11": "generated_screenshots/smart-contract-deployment.png",
    "Figure 5.12": "generated_screenshots/erc7201-storage-layout.png",
    "Figure 5.13": "generated_screenshots/bytecode-size-verification.png",
    "Figure 5.14": "generated_screenshots/amoy-testnet-deployment.png",
    "Figure 5.15": "generated_screenshots/polygonscan-confirmation.png",
    "Figure 5.16": "generated_screenshots/diamond-test-results.png",
    "Figure 5.17": "generated_screenshots/property-registration-form.png",
    "Figure 5.18": "generated_screenshots/yield-agreement-interface.png",
    "Figure 5.19": "generated_screenshots/analytics-dashboard.png",
    "Figure 5.20": "generated_screenshots/simulation-test-results.png",
    "Figure 5.21": "generated_screenshots/gas-comparison-report.png",
    "Figure 5.22": "generated_screenshots/variance-tracking.png",
    "Figure 5.23": "rendered_diagrams_drawio/yield-tokenization-flow.png",
    # Section 6 - Testing (24 figures)
    "Figure 6.1": "generated_charts/fig_6_1_token_standard_gas_comparison.png",
    "Figure 6.2": "generated_charts/fig_6_2_batch_operation_scaling.png",
    "Figure 6.3": "generated_charts/fig_6_3_amoy_anvil_variance_heatmap.png",
    "Figure 6.4": "generated_charts/fig_6_4_volatile_simulation_radar.png",
    "Figure 6.5": "generated_charts/fig_6_5_diamond_overhead_boxplot.png",
    "Figure 6.6": "generated_charts/fig_6_6_load_testing_scatter.png",
    "Figure 6.7": "generated_charts/fig_6_7_gas_cost_projections.png",
    "Figure 6.8": "generated_charts/fig_6_8_test_pass_rate_evolution.png",
    "Figure 6.9": "generated_charts/architecture-comparison.png",
    "Figure 6.10": "generated_charts/user-workflow-diagrams.png",
    "Figure 6.11": "rendered_diagrams_drawio/load-test-system-architecture.png",
    "Figure 6.12": "rendered_diagrams_drawio/load-test-workflow.png",
    "Figure 6.13": "rendered_diagrams_drawio/load-testing-results.png",
    "Figure 6.14": "rendered_diagrams_drawio/load-test-gas-scaling.png",
    "Figure 6.15": "rendered_diagrams_drawio/load-test-shareholder-distribution.png",
    "Figure 6.16": "rendered_diagrams_drawio/load-test-success-rates.png",
    "Figure 6.17": "rendered_diagrams_drawio/load-test-latency-distribution.png",
    "Figure 6.18": "rendered_diagrams_drawio/load-test-metrics-summary.png",
    "Figure 6.19": "rendered_diagrams_drawio/load-test-cost-comparison.png",
    "Figure 6.20": "rendered_diagrams_drawio/load-test-erc1155-efficiency.png",
    "Figure 6.21": "rendered_diagrams_drawio/load-test-token-recycling.png",
    "Figure 6.22": "rendered_diagrams_drawio/load-test-restriction-enforcement.png",
    "Figure 6.23": "rendered_diagrams_drawio/load-test-restriction-overhead.png",
    "Figure 6.24": "rendered_diagrams_drawio/load-test-rq-validation.png",
    # Section 7 - Evaluation (8 figures)
    "Figure 7.1": "survey_interview_charts/fig_demographic_overview.png",
    "Figure 7.2": "survey_interview_charts/fig_tokenisation_interest_analysis.png",
    "Figure 7.3": "survey_interview_charts/fig_correlation_matrix.png",
    "Figure 7.4": "survey_interview_charts/fig_interview_demographics.png",
    "Figure 7.5": "survey_interview_charts/fig_thematic_code_frequencies.png",
    "Figure 7.6": "generated_charts/token-decision-framework.png",
    "Figure 7.7": "generated_charts/fig_7_7_volatile_recovery_comparison.png",
    "Figure 7.8": "generated_charts/fig_7_8_testnet_local_comparison.png",
    # Appendix F - Web wireframes (F.1-F.10)
    "Figure F.1": "rendered_diagrams_drawio/wireframe-analytics-dashboard.png",
    "Figure F.2": "rendered_diagrams_drawio/wireframe-governance.png",
    "Figure F.3": "rendered_diagrams_drawio/wireframe-governance-proposal-detail.png",
    "Figure F.4": "rendered_diagrams_drawio/wireframe-kyc.png",
    "Figure F.5": "rendered_diagrams_drawio/wireframe-kyc-admin.png",
    "Figure F.6": "rendered_diagrams_drawio/wireframe-marketplace.png",
    "Figure F.7": "rendered_diagrams_drawio/wireframe-portfolio.png",
    "Figure F.8": "rendered_diagrams_drawio/wireframe-properties-list.png",
    "Figure F.9": "rendered_diagrams_drawio/wireframe-yield-agreement-detail.png",
    "Figure F.10": "rendered_diagrams_drawio/wireframe-yield-agreements-list.png",
    # Appendix F - Mobile wireframes (F.11-F.19)
    "Figure F.11": "rendered_diagrams_drawio/wireframe-mobile-analytics.png",
    "Figure F.12": "rendered_diagrams_drawio/wireframe-mobile-governance.png",
    "Figure F.13": "rendered_diagrams_drawio/wireframe-mobile-kyc.png",
    "Figure F.14": "rendered_diagrams_drawio/wireframe-mobile-kyc-admin.png",
    "Figure F.15": "rendered_diagrams_drawio/wireframe-mobile-marketplace.png",
    "Figure F.16": "rendered_diagrams_drawio/wireframe-mobile-portfolio.png",
    "Figure F.17": "rendered_diagrams_drawio/wireframe-mobile-properties-list.png",
    "Figure F.18": "rendered_diagrams_drawio/wireframe-mobile-register-property.png",
    "Figure F.19": "rendered_diagrams_drawio/wireframe-mobile-yield-agreements-list.png",
    # Appendix G - Additional charts (G.1-G.5)
    "Figure G.1": "survey_interview_charts/fig_cluster_analysis.png",
    "Figure G.2": "survey_interview_charts/fig_feature_importance.png",
    "Figure G.3": "survey_interview_charts/fig_motivations_concerns.png",
    "Figure G.4": "survey_interview_charts/fig_landlord_fintech_comparison.png",
    "Figure G.5": "survey_interview_charts/fig_likert_distributions.png",
}

print("Step 1: Verify all images exist...")
missing = []
for fig, path in FIGURES.items():
    if not os.path.exists(path):
        missing.append(f"{fig}: {path}")
        
if missing:
    print(f"MISSING {len(missing)} images:")
    for m in missing:
        print(f"  {m}")
    exit(1)
else:
    print(f"  All {len(FIGURES)} images verified.")

print("\nStep 2: Convert markdown to fresh Word document...")
subprocess.run(['pandoc', 'DissertationProgressFinal.md', '-o', 'temp_fresh.docx'], check=True)
print("  Done.")

print("\nStep 3: Open and process document...")
doc = Document('temp_fresh.docx')

# Find all Figure placeholder paragraphs
inserted = 0
not_found = []

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Match: **Figure X.X: Title** or Figure X.X: Title
    match = re.match(r'^\*?\*?(Figure\s+[0-9A-GF]+\.[0-9]+):\s*(.+?)\*?\*?$', text)
    
    if match:
        fig_num = match.group(1)
        title = match.group(2).rstrip('*')
        
        if fig_num in FIGURES:
            img_path = FIGURES[fig_num]
            
            if os.path.exists(img_path):
                # Clear the paragraph
                para.clear()
                
                # Add image
                run = para.add_run()
                width = Inches(5.5)
                run.add_picture(img_path, width=width)
                
                # Add line break and caption
                para.add_run('\n')
                caption_run = para.add_run(f"{fig_num}: {title}")
                caption_run.bold = True
                caption_run.font.size = Pt(10)
                
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                inserted += 1
                if inserted <= 5 or inserted % 20 == 0:
                    print(f"  ✓ {fig_num}")
        else:
            not_found.append(fig_num)

print(f"\nStep 4: Save document...")
doc.save('DissertationProgressFinal_v2.2.docx')
os.remove('temp_fresh.docx')

print(f"\n{'='*60}")
print(f"COMPLETE: Inserted {inserted} figures")
print(f"Total figures expected: {len(FIGURES)}")
if not_found:
    print(f"Not found in mapping: {not_found}")
print(f"Output: DissertationProgressFinal_v2.2.docx")
