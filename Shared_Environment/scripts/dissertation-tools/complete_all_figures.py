#!/usr/bin/env python3
"""
Complete the Word document with ALL 108 figures:
1. Replace existing placeholders (46 Mermaid diagrams)
2. Add Section 5 screenshots (23 images)
3. Add Section 7 survey/interview figures (10 images)
4. Add Appendix wireframes (19 images)
5. Add any remaining generated charts (10 images)
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
import re
from pathlib import Path

# Paths
doc_path = "DissertationProgressFinal.docx"
output_path = "DissertationProgressFinal_COMPLETE.docx"

# Image directories
drawio_dir = Path("rendered_diagrams_drawio")
survey_dir = Path("survey_interview_charts")
charts_dir = Path("generated_charts")
screenshots_dir = Path("generated_screenshots")

# Load document
print(f"Loading {doc_path}...")
doc = Document(doc_path)

# Build complete image inventory
all_images = {}

# Draw.io diagrams (65)
for f in drawio_dir.glob("*.png"):
    all_images[f.stem] = f

# Survey/Interview charts (10)
for f in survey_dir.glob("*.png"):
    all_images[f.stem] = f

# Generated charts (10)
if charts_dir.exists():
    for f in charts_dir.glob("*.png"):
        all_images[f.stem] = f

# Generated screenshots (23)
for f in screenshots_dir.glob("*.png"):
    all_images[f.stem] = f

print(f"Total images available: {len(all_images)}")

# Track what's been inserted
inserted_images = set()

# Pattern to extract .mmd filename from placeholder
mmd_pattern = re.compile(r'([a-z0-9-]+)\.mmd', re.IGNORECASE)

def add_figure(para, image_path, caption=None):
    """Insert image and optional caption"""
    para.clear()
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(5.5))
    inserted_images.add(image_path.stem)

# ============================================================
# PHASE 1: Replace existing Mermaid placeholders
# ============================================================
print("\n" + "="*60)
print("PHASE 1: Replacing Mermaid diagram placeholders")
print("="*60)

phase1_count = 0
for para in doc.paragraphs:
    text = para.text
    
    if '[PLACEHOLDER' in text.upper() or '[Placeholder' in text:
        mmd_match = mmd_pattern.search(text)
        if mmd_match:
            filename = mmd_match.group(1)
            if filename in all_images:
                add_figure(para, all_images[filename])
                phase1_count += 1
                print(f"  ✓ {filename}")

print(f"\nPhase 1 complete: {phase1_count} Mermaid diagrams inserted")

# ============================================================
# PHASE 2: Add Section 5 screenshots after figure references
# ============================================================
print("\n" + "="*60)
print("PHASE 2: Adding Section 5 screenshots")
print("="*60)

section5_figures = [
    ("Figure 5.1", "docker-desktop-dashboard", "Docker Desktop Container Status"),
    ("Figure 5.2", "container-resource-allocation", "Container Resource Allocation"),
    ("Figure 5.3", "health-check-configuration", "Health Check Configuration"),
    ("Figure 5.4", "network-isolation-setup", "Network Isolation Setup"),
    ("Figure 5.5", "volume-configuration", "Volume Configuration"),
    ("Figure 5.6", "iteration-1-test-results", "Iteration 1 Test Results"),
    ("Figure 5.7", "smart-contract-deployment", "Smart Contract Deployment"),
    ("Figure 5.8", "iteration-3-test-results", "Iteration 3 Test Results"),
    ("Figure 5.9", "iteration-5-test-results", "Iteration 5 Test Results"),
    ("Figure 5.10", "backend-api-test-results", "Backend API Test Results"),
    ("Figure 5.11", "iteration-7-test-results", "Iteration 7 Test Results"),
    ("Figure 5.12", "frontend-build-output", "Frontend Build Output"),
    ("Figure 5.13", "iteration-10-test-results", "Iteration 10 Test Results"),
    ("Figure 5.14", "mobile-app-screenshots", "Mobile App Screenshots"),
    ("Figure 5.15", "iteration-11-test-results", "Iteration 11 Test Results"),
    ("Figure 5.16", "analytics-dashboard", "Analytics Dashboard"),
    ("Figure 5.17", "iteration-12-test-results", "Iteration 12 Test Results"),
    ("Figure 5.18", "load-testing-results-detailed", "Load Testing Results"),
    ("Figure 5.19", "iteration-14-test-results", "Iteration 14 Test Results"),
    ("Figure 5.20", "diamond-pattern-migration", "Diamond Pattern Migration"),
    ("Figure 5.21", "iteration-15-test-results", "Iteration 15 Test Results"),
    ("Figure 5.22", "iteration-16-test-results", "Iteration 16 Test Results"),
    ("Figure 5.23", "amoy-testnet-deployment", "Amoy Testnet Deployment"),
]

phase2_count = 0
paragraphs_to_process = list(doc.paragraphs)

for i, para in enumerate(paragraphs_to_process):
    text = para.text
    
    for fig_ref, img_name, caption in section5_figures:
        # Check if this paragraph contains the figure reference
        if fig_ref in text and img_name in all_images and img_name not in inserted_images:
            # Check if next paragraph is a placeholder that should be replaced
            if i + 1 < len(paragraphs_to_process):
                next_para = paragraphs_to_process[i + 1]
                if 'Screenshot placeholder' in next_para.text or '[Screenshot' in next_para.text:
                    add_figure(next_para, all_images[img_name])
                    phase2_count += 1
                    print(f"  ✓ {fig_ref}: {caption}")
                    break

# If screenshots weren't found via placeholders, find Section 5 and add them
if phase2_count < 10:
    print(f"  (Adding remaining screenshots to Section 5...)")
    
    # Find "Section 5" or "Implementation" heading
    section5_start = None
    section6_start = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.lower()
        if 'implementation' in text and ('5.' in para.text or 'section 5' in text):
            section5_start = i
        if 'testing' in text and ('6.' in para.text or 'section 6' in text):
            section6_start = i
            break
    
    if section5_start and section6_start:
        # Add any missing Section 5 figures before Section 6
        for fig_ref, img_name, caption in section5_figures:
            if img_name in all_images and img_name not in inserted_images:
                # Create new paragraph with image
                new_para = doc.add_paragraph()
                run = new_para.add_run()
                run.add_picture(str(all_images[img_name]), width=Inches(5.5))
                inserted_images.add(img_name)
                phase2_count += 1
                print(f"  ✓ Added: {fig_ref}: {caption}")

print(f"\nPhase 2 complete: {phase2_count} Section 5 screenshots added")

# ============================================================
# PHASE 3: Add Section 7 survey/interview figures
# ============================================================
print("\n" + "="*60)
print("PHASE 3: Adding Section 7 survey/interview figures")
print("="*60)

section7_figures = [
    ("Figure 7.1", "fig_demographic_overview", "Survey Demographics Overview"),
    ("Figure 7.2", "fig_tokenisation_interest_analysis", "Tokenisation Interest Analysis"),
    ("Figure 7.3", "fig_correlation_matrix", "Spearman Correlation Matrix"),
    ("Figure 7.4", "fig_interview_demographics", "Interview Demographics"),
    ("Figure 7.5", "fig_thematic_code_frequencies", "Theme Frequency Analysis"),
    ("Figure 7.6", "fig_cluster_analysis", "Cluster Analysis"),
    ("Figure 7.7", "fig_motivations_concerns", "Motivations and Concerns"),
    ("Figure 7.8", "fig_feature_importance", "Feature Importance"),
    ("Figure 7.9", "fig_landlord_fintech_comparison", "Landlord vs Expert Comparison"),
    ("Figure 7.10", "fig_likert_distributions", "Likert Scale Distributions"),
]

phase3_count = 0
for fig_ref, img_name, caption in section7_figures:
    if img_name in all_images and img_name not in inserted_images:
        # Find the figure reference in the document
        for para in doc.paragraphs:
            if fig_ref in para.text:
                # Add image after this paragraph
                new_para = doc.add_paragraph()
                run = new_para.add_run()
                run.add_picture(str(all_images[img_name]), width=Inches(5.5))
                inserted_images.add(img_name)
                phase3_count += 1
                print(f"  ✓ {fig_ref}: {caption}")
                break

print(f"\nPhase 3 complete: {phase3_count} Section 7 figures added")

# ============================================================
# PHASE 4: Add Appendix with 19 extra wireframes
# ============================================================
print("\n" + "="*60)
print("PHASE 4: Adding Appendix wireframes")
print("="*60)

# Extra wireframes (19 that weren't in the main document)
appendix_wireframes = [
    ("wireframe-analytics-dashboard", "Analytics Dashboard Wireframe"),
    ("wireframe-governance", "Governance Dashboard Wireframe"),
    ("wireframe-governance-proposal-detail", "Governance Proposal Detail Wireframe"),
    ("wireframe-kyc", "KYC Verification Wireframe"),
    ("wireframe-kyc-admin", "KYC Admin Dashboard Wireframe"),
    ("wireframe-marketplace", "Marketplace Wireframe"),
    ("wireframe-portfolio", "Portfolio Dashboard Wireframe"),
    ("wireframe-properties-list", "Properties List Wireframe"),
    ("wireframe-yield-agreement-detail", "Yield Agreement Detail Wireframe"),
    ("wireframe-yield-agreements-list", "Yield Agreements List Wireframe"),
    ("wireframe-mobile-analytics", "Mobile Analytics Wireframe"),
    ("wireframe-mobile-governance", "Mobile Governance Wireframe"),
    ("wireframe-mobile-kyc", "Mobile KYC Wireframe"),
    ("wireframe-mobile-kyc-admin", "Mobile KYC Admin Wireframe"),
    ("wireframe-mobile-marketplace", "Mobile Marketplace Wireframe"),
    ("wireframe-mobile-portfolio", "Mobile Portfolio Wireframe"),
    ("wireframe-mobile-properties-list", "Mobile Properties List Wireframe"),
    ("wireframe-mobile-register-property", "Mobile Property Registration Wireframe"),
    ("wireframe-mobile-yield-agreements-list", "Mobile Yield Agreements List Wireframe"),
]

# Find or create Appendix section
appendix_found = False
for i, para in enumerate(doc.paragraphs):
    if 'appendix' in para.text.lower() and ('wireframe' in para.text.lower() or 'ui' in para.text.lower()):
        appendix_found = True
        break

# Add Appendix heading if not found
if not appendix_found:
    doc.add_page_break()
    heading = doc.add_heading('Appendix D: Additional UI Wireframes', level=1)

phase4_count = 0

# Add web wireframes section
doc.add_heading('D.1 Web Application Wireframes', level=2)

web_wireframes = [w for w in appendix_wireframes if 'mobile' not in w[0]]
for img_name, caption in web_wireframes:
    if img_name in all_images and img_name not in inserted_images:
        # Add caption
        caption_para = doc.add_paragraph()
        caption_para.add_run(f"Figure D.{phase4_count + 1}: {caption}").bold = True
        
        # Add image
        img_para = doc.add_paragraph()
        run = img_para.add_run()
        run.add_picture(str(all_images[img_name]), width=Inches(5.5))
        
        inserted_images.add(img_name)
        phase4_count += 1
        print(f"  ✓ Figure D.{phase4_count}: {caption}")

# Add mobile wireframes section
doc.add_heading('D.2 Mobile Application Wireframes', level=2)

mobile_wireframes = [w for w in appendix_wireframes if 'mobile' in w[0]]
mobile_count = 0
for img_name, caption in mobile_wireframes:
    if img_name in all_images and img_name not in inserted_images:
        # Add caption
        caption_para = doc.add_paragraph()
        caption_para.add_run(f"Figure D.{phase4_count + mobile_count + 1}: {caption}").bold = True
        
        # Add image
        img_para = doc.add_paragraph()
        run = img_para.add_run()
        run.add_picture(str(all_images[img_name]), width=Inches(5.5))
        
        inserted_images.add(img_name)
        mobile_count += 1
        print(f"  ✓ Figure D.{phase4_count + mobile_count}: {caption}")

phase4_count += mobile_count
print(f"\nPhase 4 complete: {phase4_count} Appendix wireframes added")

# ============================================================
# PHASE 5: Add any remaining images
# ============================================================
print("\n" + "="*60)
print("PHASE 5: Checking for remaining images")
print("="*60)

remaining = []
for name, path in all_images.items():
    if name not in inserted_images:
        remaining.append((name, path))

if remaining:
    print(f"  {len(remaining)} images not yet inserted:")
    for name, path in remaining[:10]:
        print(f"    - {name}")
    if len(remaining) > 10:
        print(f"    ... and {len(remaining) - 10} more")
else:
    print("  All images have been inserted!")

# ============================================================
# Save document
# ============================================================
print(f"\nSaving to {output_path}...")
doc.save(output_path)

total_inserted = len(inserted_images)
print(f"\n{'='*60}")
print(f"FINAL SUMMARY")
print(f"{'='*60}")
print(f"Phase 1 (Mermaid diagrams):     {phase1_count}")
print(f"Phase 2 (Section 5 screenshots): {phase2_count}")
print(f"Phase 3 (Section 7 figures):     {phase3_count}")
print(f"Phase 4 (Appendix wireframes):   {phase4_count}")
print(f"{'='*60}")
print(f"TOTAL FIGURES INSERTED:          {total_inserted}")
print(f"REMAINING (not inserted):        {len(remaining)}")
print(f"\n✅ Document saved: {output_path}")
print(f"   File size: {os.path.getsize(output_path) / 1024 / 1024:.1f} MB")
