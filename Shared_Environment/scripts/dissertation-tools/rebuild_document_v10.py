#!/usr/bin/env python3
"""
Robust Dissertation Document Rebuild Script
Version 10: Handles both media/image*.png references AND Figure placeholders

This script provides a robust pipeline for converting markdown to Word documents
with proper image embedding. It handles two types of image references:
1. Pandoc-style: ![](media/image*.png) - from Word-to-markdown conversion
2. Placeholder-style: **Figure X.Y:** - for charts that need external insertion

Usage:
    python rebuild_document_v10.py [input_md] [output_docx]
    
    Defaults:
    - input_md: DissertationProgressFinal_v2.6.md
    - output_docx: DissertationProgressFinal_v2.6.docx
"""

import subprocess
import os
import re
import sys
import shutil
from pathlib import Path
from datetime import datetime

# =============================================================================
# CONFIGURATION
# =============================================================================

WORKSPACE = Path("/Users/scott/Cursor/RWA-Rental-Yield-Tokenisation_v5.1")

# Default input/output files
DEFAULT_MARKDOWN = WORKSPACE / "DissertationProgressFinal_v2.6.md"
DEFAULT_OUTPUT = WORKSPACE / "DissertationProgressFinal_v2.6.docx"
REFERENCE_DOC = WORKSPACE / "DissertationProgressFinal_v2.5.docx"

# Image directories
MEDIA_DIR = WORKSPACE / "media"  # Extracted from Word document
DIAGRAMS_DIR = WORKSPACE / "rendered_diagrams_drawio"
GENERATED_SCREENSHOTS = WORKSPACE / "generated_screenshots"
GENERATED_CHARTS = WORKSPACE / "generated_charts"
SURVEY_CHARTS = WORKSPACE / "survey_interview_charts"
ARCHITECTURE_DIAGRAMS = WORKSPACE / "Shared_Environment/docs/architecture/diagrams"

# =============================================================================
# FIGURE PLACEHOLDER MAPPING
# =============================================================================

def get_figure_mapping():
    """
    Create mapping from figure placeholders (e.g., **Figure S.1:**) to image files.
    These are figures that need external chart images inserted.
    """
    mapping = {}
    
    # Survey figures (S.1-S.6)
    mapping["S.1"] = SURVEY_CHARTS / "fig_demographic_overview.png"
    mapping["S.2"] = SURVEY_CHARTS / "fig_tokenisation_interest_analysis.png"
    mapping["S.3"] = SURVEY_CHARTS / "fig_correlation_matrix.png"
    mapping["S.4"] = SURVEY_CHARTS / "fig_cluster_analysis.png"
    mapping["S.5"] = SURVEY_CHARTS / "fig_motivations_concerns.png"
    mapping["S.6"] = SURVEY_CHARTS / "fig_feature_importance.png"
    
    # Interview figures (I.1-I.4)
    mapping["I.1"] = SURVEY_CHARTS / "fig_interview_demographics.png"
    mapping["I.2"] = SURVEY_CHARTS / "fig_thematic_code_frequencies.png"
    mapping["I.3"] = SURVEY_CHARTS / "fig_landlord_fintech_comparison.png"
    mapping["I.4"] = SURVEY_CHARTS / "fig_likert_distributions.png"
    
    # Appendix F wireframes (F.1-F.19)
    mapping["F.1"] = DIAGRAMS_DIR / "wireframe-analytics-dashboard.png"
    mapping["F.2"] = DIAGRAMS_DIR / "wireframe-governance.png"
    mapping["F.3"] = DIAGRAMS_DIR / "wireframe-governance-proposal-detail.png"
    mapping["F.4"] = DIAGRAMS_DIR / "wireframe-kyc.png"
    mapping["F.5"] = DIAGRAMS_DIR / "wireframe-kyc-admin.png"
    mapping["F.6"] = DIAGRAMS_DIR / "wireframe-marketplace.png"
    mapping["F.7"] = DIAGRAMS_DIR / "wireframe-portfolio.png"
    mapping["F.8"] = DIAGRAMS_DIR / "wireframe-properties-list.png"
    mapping["F.9"] = DIAGRAMS_DIR / "wireframe-yield-agreement-detail.png"
    mapping["F.10"] = DIAGRAMS_DIR / "wireframe-yield-agreements-list.png"
    mapping["F.11"] = DIAGRAMS_DIR / "wireframe-mobile-analytics.png"
    mapping["F.12"] = DIAGRAMS_DIR / "wireframe-mobile-governance.png"
    mapping["F.13"] = DIAGRAMS_DIR / "wireframe-mobile-kyc.png"
    mapping["F.14"] = DIAGRAMS_DIR / "wireframe-mobile-kyc-admin.png"
    mapping["F.15"] = DIAGRAMS_DIR / "wireframe-mobile-marketplace.png"
    mapping["F.16"] = DIAGRAMS_DIR / "wireframe-mobile-portfolio.png"
    mapping["F.17"] = DIAGRAMS_DIR / "wireframe-mobile-properties-list.png"
    mapping["F.18"] = DIAGRAMS_DIR / "wireframe-mobile-register-property.png"
    mapping["F.19"] = DIAGRAMS_DIR / "wireframe-mobile-yield-agreements-list.png"
    
    # Appendix G charts (G.1-G.5)
    mapping["G.1"] = SURVEY_CHARTS / "fig_cluster_analysis.png"
    mapping["G.2"] = SURVEY_CHARTS / "fig_feature_importance.png"
    mapping["G.3"] = SURVEY_CHARTS / "fig_motivations_concerns.png"
    mapping["G.4"] = SURVEY_CHARTS / "fig_landlord_fintech_comparison.png"
    mapping["G.5"] = SURVEY_CHARTS / "fig_likert_distributions.png"
    
    return mapping

# =============================================================================
# VALIDATION FUNCTIONS
# =============================================================================

def validate_prerequisites():
    """Validate all prerequisites are in place."""
    print("\n" + "=" * 60)
    print("PHASE 1: Validating Prerequisites")
    print("=" * 60)
    
    errors = []
    warnings = []
    
    # Check pandoc
    try:
        result = subprocess.run(["pandoc", "--version"], capture_output=True, text=True)
        if result.returncode == 0:
            version = result.stdout.split('\n')[0]
            print(f"✓ Pandoc installed: {version}")
        else:
            errors.append("Pandoc not working properly")
    except FileNotFoundError:
        errors.append("Pandoc not installed. Install with: brew install pandoc")
    
    # Check python-docx
    try:
        from docx import Document
        print("✓ python-docx installed")
    except ImportError:
        print("⚠ python-docx not installed. Installing...")
        subprocess.run(["pip", "install", "--break-system-packages", "python-docx"], 
                      capture_output=True)
        try:
            from docx import Document
            print("✓ python-docx installed successfully")
        except ImportError:
            errors.append("Failed to install python-docx")
    
    # Check media directory
    if MEDIA_DIR.exists():
        media_count = len(list(MEDIA_DIR.glob("*.png")))
        print(f"✓ Media directory exists with {media_count} images")
    else:
        errors.append(f"Media directory not found: {MEDIA_DIR}")
        print(f"  Run: unzip -o {REFERENCE_DOC} 'word/media/*' -d /tmp/docx_extract && mv /tmp/docx_extract/word/media/* {MEDIA_DIR}/")
    
    # Check reference document
    if REFERENCE_DOC.exists():
        print(f"✓ Reference document exists: {REFERENCE_DOC.name}")
    else:
        warnings.append(f"Reference document not found: {REFERENCE_DOC}")
        print(f"⚠ Reference document not found (styles may not be preserved)")
    
    # Check figure mapping directories
    figure_mapping = get_figure_mapping()
    missing_dirs = set()
    for fig_num, img_path in figure_mapping.items():
        if not img_path.parent.exists():
            missing_dirs.add(img_path.parent)
    
    if missing_dirs:
        for d in missing_dirs:
            warnings.append(f"Directory not found: {d}")
            print(f"⚠ Directory not found: {d}")
    else:
        print(f"✓ All figure mapping directories exist")
    
    if errors:
        print("\n❌ ERRORS (must be fixed):")
        for e in errors:
            print(f"  - {e}")
        return False
    
    if warnings:
        print(f"\n⚠ {len(warnings)} warnings (may cause missing images)")
    
    return True

def validate_markdown_images(markdown_file):
    """Validate all image references in markdown can be resolved."""
    print("\n" + "=" * 60)
    print("PHASE 2: Validating Markdown Image References")
    print("=" * 60)
    
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Find all media/image*.png references
    media_refs = re.findall(r'media/image(\d+)\.png', content)
    unique_media_refs = sorted(set(int(n) for n in media_refs))
    
    # Find all Figure placeholders
    figure_refs = re.findall(r'\*\*Figure\s+([A-Z]?\d+\.?\d*):', content)
    unique_figure_refs = sorted(set(figure_refs), key=lambda x: (x[0].isdigit(), x))
    
    print(f"Found {len(unique_media_refs)} unique media/image references")
    print(f"Found {len(unique_figure_refs)} unique Figure placeholders")
    
    # Check media images exist
    missing_media = []
    for img_num in unique_media_refs:
        img_path = MEDIA_DIR / f"image{img_num}.png"
        if not img_path.exists():
            missing_media.append(img_num)
    
    if missing_media:
        print(f"\n⚠ Missing media images: {missing_media[:10]}{'...' if len(missing_media) > 10 else ''}")
    else:
        print(f"✓ All {len(unique_media_refs)} media images exist")
    
    # Check figure mappings exist
    figure_mapping = get_figure_mapping()
    missing_figures = []
    for fig_num in unique_figure_refs:
        if fig_num in figure_mapping:
            if not figure_mapping[fig_num].exists():
                missing_figures.append(f"{fig_num}: {figure_mapping[fig_num]}")
        else:
            missing_figures.append(f"{fig_num}: No mapping defined")
    
    if missing_figures:
        print(f"\n⚠ Missing figure images:")
        for f in missing_figures[:10]:
            print(f"  - {f}")
        if len(missing_figures) > 10:
            print(f"  ... and {len(missing_figures) - 10} more")
    else:
        print(f"✓ All {len(unique_figure_refs)} figure images mapped and exist")
    
    return {
        'media_refs': unique_media_refs,
        'figure_refs': unique_figure_refs,
        'missing_media': missing_media,
        'missing_figures': missing_figures
    }

# =============================================================================
# CONVERSION FUNCTIONS
# =============================================================================

def convert_markdown_to_docx(markdown_file, output_file, use_reference=True):
    """Convert markdown to Word using pandoc with reference document."""
    print("\n" + "=" * 60)
    print("PHASE 3: Converting Markdown to Word")
    print("=" * 60)
    
    temp_output = output_file.parent / f"temp_{output_file.name}"
    
    # Build pandoc command
    cmd = [
        "pandoc",
        str(markdown_file),
        "-o", str(temp_output),
        "--from", "markdown",
        "--to", "docx",
        "--standalone",
        "--toc",
        "--toc-depth=3",
        f"--resource-path={WORKSPACE}:{MEDIA_DIR}"
    ]
    
    # Add reference document for styles if available
    if use_reference and REFERENCE_DOC.exists():
        cmd.extend(["--reference-doc", str(REFERENCE_DOC)])
        print(f"Using reference document for styles: {REFERENCE_DOC.name}")
    
    print(f"Input: {markdown_file.name}")
    print(f"Output: {temp_output.name}")
    print(f"Running pandoc...")
    
    result = subprocess.run(cmd, capture_output=True, text=True)
    
    if result.returncode != 0:
        print(f"❌ Pandoc error: {result.stderr}")
        return None
    
    print(f"✓ Created temporary document: {temp_output.name}")
    return temp_output

def insert_figure_images(docx_file, output_file):
    """Insert images for Figure placeholders in the Word document."""
    print("\n" + "=" * 60)
    print("PHASE 4: Inserting Figure Placeholder Images")
    print("=" * 60)
    
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document(str(docx_file))
    mapping = get_figure_mapping()
    
    # Pattern to match figure placeholders
    figure_pattern = re.compile(r'^(?:\*\*)?Figure\s+([A-Z]?\d+\.?\d*):(.+?)(?:\*\*)?$')
    
    figures_found = 0
    figures_inserted = 0
    figures_missing = []
    
    print("Scanning paragraphs for Figure placeholders...")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        match = figure_pattern.match(text)
        
        if match:
            fig_num = match.group(1)
            fig_title = match.group(2).strip()
            figures_found += 1
            
            if fig_num in mapping:
                img_path = mapping[fig_num]
                
                if img_path.exists():
                    try:
                        # Clear paragraph and add image
                        para.clear()
                        run = para.add_run()
                        run.add_picture(str(img_path), width=Inches(6.0))
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add caption below (new paragraph would be needed for proper caption)
                        figures_inserted += 1
                        print(f"  ✓ Figure {fig_num}: {fig_title[:40]}...")
                    except Exception as e:
                        print(f"  ✗ Error inserting Figure {fig_num}: {e}")
                        para.add_run(f"Figure {fig_num}: {fig_title}")
                        figures_missing.append(f"{fig_num}: Error - {e}")
                else:
                    print(f"  ✗ Image not found: {img_path.name}")
                    figures_missing.append(f"{fig_num}: {img_path}")
            else:
                # No mapping - might be embedded image already
                pass
    
    print(f"\nFigure placeholder summary:")
    print(f"  Found: {figures_found}")
    print(f"  Inserted: {figures_inserted}")
    print(f"  Missing: {len(figures_missing)}")
    
    # Save the document
    doc.save(str(output_file))
    print(f"✓ Saved document: {output_file.name}")
    
    return {
        'found': figures_found,
        'inserted': figures_inserted,
        'missing': figures_missing
    }

# =============================================================================
# MAIN PIPELINE
# =============================================================================

def rebuild_document(markdown_file=None, output_file=None):
    """Main rebuild pipeline."""
    
    # Set defaults
    markdown_file = Path(markdown_file) if markdown_file else DEFAULT_MARKDOWN
    output_file = Path(output_file) if output_file else DEFAULT_OUTPUT
    
    print("=" * 70)
    print("ROBUST DISSERTATION DOCUMENT REBUILD SCRIPT v10")
    print("=" * 70)
    print(f"Timestamp: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Input:  {markdown_file}")
    print(f"Output: {output_file}")
    
    # Phase 1: Validate prerequisites
    if not validate_prerequisites():
        print("\n❌ Prerequisites validation failed. Fix errors and retry.")
        return False
    
    # Check input file exists
    if not markdown_file.exists():
        print(f"\n❌ Input file not found: {markdown_file}")
        return False
    
    # Phase 2: Validate markdown images
    validation = validate_markdown_images(markdown_file)
    
    # Phase 3: Convert markdown to Word
    temp_docx = convert_markdown_to_docx(markdown_file, output_file)
    if not temp_docx:
        print("\n❌ Conversion failed.")
        return False
    
    # Phase 4: Insert figure placeholder images
    result = insert_figure_images(temp_docx, output_file)
    
    # Cleanup temp file
    if temp_docx.exists():
        os.remove(temp_docx)
        print(f"✓ Cleaned up temporary file")
    
    # Final summary
    print("\n" + "=" * 70)
    print("BUILD COMPLETE")
    print("=" * 70)
    print(f"Output file: {output_file}")
    print(f"File size: {output_file.stat().st_size / 1024 / 1024:.2f} MB")
    print(f"\nImage summary:")
    print(f"  Media images (embedded via pandoc): {len(validation['media_refs'])}")
    print(f"  Figure placeholders inserted: {result['inserted']}")
    
    if validation['missing_media'] or result['missing']:
        print(f"\n⚠ Some images may be missing. Check the document manually.")
    else:
        print(f"\n✓ All images should be present.")
    
    return True

# =============================================================================
# CLI ENTRY POINT
# =============================================================================

if __name__ == "__main__":
    if len(sys.argv) >= 3:
        rebuild_document(sys.argv[1], sys.argv[2])
    elif len(sys.argv) == 2:
        rebuild_document(sys.argv[1])
    else:
        rebuild_document()


