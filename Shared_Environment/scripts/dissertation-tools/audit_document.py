#!/usr/bin/env python3
"""
Comprehensive audit of DissertationProgressFinal_v2.1.docx
- Extract ALL images with their surrounding context
- Extract ALL tables with their surrounding context
- Identify missing Figure/Table labels
- Identify misplaced content
"""

from docx import Document
from docx.shared import Inches
import re
import json

def get_paragraph_text(para):
    """Get text from paragraph"""
    return para.text.strip() if para.text else ""

def audit_document(docx_path):
    doc = Document(docx_path)
    
    images = []
    tables = []
    current_section = "Unknown"
    current_subsection = ""
    
    # Track paragraph index for context
    all_elements = []
    
    # First pass: collect all paragraphs and their types
    for i, para in enumerate(doc.paragraphs):
        text = get_paragraph_text(para)
        style_name = para.style.name if para.style else "Normal"
        
        # Track section headers
        if style_name.startswith("Heading 1") or (text and re.match(r'^[0-9]+\.\s+[A-Z]', text)):
            current_section = text[:80]
            current_subsection = ""
        elif style_name.startswith("Heading 2") or (text and re.match(r'^[0-9]+\.[0-9]+\s+', text)):
            current_subsection = text[:80]
        
        # Check for images in this paragraph
        has_image = False
        for run in para.runs:
            if run._element.xpath('.//a:blip'):
                has_image = True
                break
        
        all_elements.append({
            'index': i,
            'type': 'image_para' if has_image else 'text',
            'text': text,
            'style': style_name,
            'section': current_section,
            'subsection': current_subsection,
            'has_image': has_image
        })
    
    # Second pass: analyze images with context
    print("=" * 80)
    print("IMAGE AUDIT REPORT")
    print("=" * 80)
    
    image_count = 0
    issues = []
    
    for i, elem in enumerate(all_elements):
        if elem['has_image']:
            image_count += 1
            
            # Get surrounding context (3 paragraphs before and after)
            context_before = []
            context_after = []
            
            for j in range(max(0, i-3), i):
                if all_elements[j]['text']:
                    context_before.append(all_elements[j]['text'][:100])
            
            for j in range(i+1, min(len(all_elements), i+4)):
                if all_elements[j]['text']:
                    context_after.append(all_elements[j]['text'][:100])
            
            # Check if there's a Figure label nearby
            has_figure_label = False
            figure_label = ""
            
            # Check in context before (most common placement)
            for ctx in context_before:
                if re.search(r'Figure\s+[0-9]+\.[0-9]+', ctx):
                    has_figure_label = True
                    match = re.search(r'Figure\s+[0-9]+\.[0-9]+[^:]*', ctx)
                    if match:
                        figure_label = match.group(0)
                    break
            
            # Check in context after
            if not has_figure_label:
                for ctx in context_after:
                    if re.search(r'Figure\s+[0-9]+\.[0-9]+', ctx):
                        has_figure_label = True
                        match = re.search(r'Figure\s+[0-9]+\.[0-9]+[^:]*', ctx)
                        if match:
                            figure_label = match.group(0)
                        break
            
            # Check current paragraph text
            if not has_figure_label and elem['text']:
                if re.search(r'Figure\s+[0-9]+\.[0-9]+', elem['text']):
                    has_figure_label = True
                    match = re.search(r'Figure\s+[0-9]+\.[0-9]+[^:]*', elem['text'])
                    if match:
                        figure_label = match.group(0)
            
            status = "✓" if has_figure_label else "✗ MISSING LABEL"
            
            print(f"\n--- Image #{image_count} (Para {i}) ---")
            print(f"Section: {elem['section']}")
            print(f"Subsection: {elem['subsection']}")
            print(f"Status: {status}")
            if figure_label:
                print(f"Label: {figure_label}")
            print(f"Context Before:")
            for ctx in context_before[-2:]:
                print(f"  > {ctx}")
            print(f"Image Paragraph Text: {elem['text'][:100] if elem['text'] else '[No text]'}")
            print(f"Context After:")
            for ctx in context_after[:2]:
                print(f"  < {ctx}")
            
            if not has_figure_label:
                issues.append({
                    'type': 'image',
                    'number': image_count,
                    'para_index': i,
                    'section': elem['section'],
                    'subsection': elem['subsection'],
                    'context_before': context_before,
                    'context_after': context_after
                })
    
    print(f"\n{'=' * 80}")
    print(f"SUMMARY: Found {image_count} images, {len(issues)} without proper Figure labels")
    print("=" * 80)
    
    # Now audit tables
    print("\n" + "=" * 80)
    print("TABLE AUDIT REPORT")
    print("=" * 80)
    
    table_issues = []
    
    # We need to iterate through the document body to find tables in order
    # This is more complex as tables are separate from paragraphs
    
    # Get all tables
    for t_idx, table in enumerate(doc.tables):
        # Try to find context by looking at nearby paragraphs
        # This is approximate since tables aren't in paragraph flow
        
        # Get first cell content as identifier
        first_cell = ""
        if table.rows and table.rows[0].cells:
            first_cell = table.rows[0].cells[0].text[:50]
        
        # Check if there's a table caption in first row
        has_table_label = False
        table_label = ""
        
        for row in table.rows[:2]:  # Check first 2 rows
            for cell in row.cells:
                if re.search(r'Table\s+[0-9]+\.[0-9]+', cell.text):
                    has_table_label = True
                    match = re.search(r'Table\s+[0-9]+\.[0-9]+[^:|\n]*', cell.text)
                    if match:
                        table_label = match.group(0)
                    break
        
        status = "✓" if has_table_label else "✗ MISSING LABEL"
        
        print(f"\n--- Table #{t_idx + 1} ---")
        print(f"Status: {status}")
        if table_label:
            print(f"Label: {table_label}")
        print(f"First cell: {first_cell}")
        print(f"Rows: {len(table.rows)}, Cols: {len(table.columns)}")
        
        if not has_table_label:
            table_issues.append({
                'type': 'table',
                'number': t_idx + 1,
                'first_cell': first_cell,
                'rows': len(table.rows),
                'cols': len(table.columns)
            })
    
    print(f"\n{'=' * 80}")
    print(f"SUMMARY: Found {len(doc.tables)} tables, {len(table_issues)} without proper Table labels")
    print("=" * 80)
    
    return {'image_issues': issues, 'table_issues': table_issues, 'total_images': image_count, 'total_tables': len(doc.tables)}

if __name__ == "__main__":
    result = audit_document("DissertationProgressFinal_v2.1.docx")
    
    # Save issues to JSON for further processing
    with open('document_audit_issues.json', 'w') as f:
        json.dump(result, f, indent=2)
    
    print(f"\n\nIssues saved to document_audit_issues.json")
