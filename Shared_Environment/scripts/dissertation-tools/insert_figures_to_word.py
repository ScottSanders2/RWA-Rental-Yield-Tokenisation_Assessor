#!/usr/bin/env python3
"""
Insert all figures into DissertationProgressFinal.docx
Maps placeholders to actual image files and inserts them.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from pathlib import Path

# Paths
doc_path = "DissertationProgressFinal.docx"
output_path = "DissertationProgressFinal_WithFigures.docx"

# Image directories
drawio_dir = Path("rendered_diagrams_drawio")
survey_dir = Path("survey_interview_charts")
charts_dir = Path("generated_charts")

# Load document
print(f"Loading {doc_path}...")
doc = Document(doc_path)

# Create mapping of figure descriptions to image files
image_mapping = {}

# Map Draw.io diagrams
drawio_files = {f.stem: f for f in drawio_dir.glob("*.png")}
print(f"Found {len(drawio_files)} Draw.io PNGs")

# Map survey/interview charts
survey_files = {f.stem: f for f in survey_dir.glob("*.png")}
print(f"Found {len(survey_files)} Survey/Interview PNGs")

# Map generated charts
chart_files = {f.stem: f for f in charts_dir.glob("*.png")} if charts_dir.exists() else {}
print(f"Found {len(chart_files)} Generated chart PNGs")

# Keywords to image file mapping
keyword_mappings = {
    # Architecture diagrams
    "docker architecture": "docker-architecture-logical",
    "docker.*logical": "docker-architecture-logical",
    "docker.*physical": "docker-architecture-physical",
    "monitoring architecture": "monitoring-architecture",
    "mobile.*architecture": "mobile-architecture",
    "backend.*api.*architecture": "backend-api-architecture",
    "property owner workflow": "property-owner-workflow",
    "yield.*architecture": "yield-architecture",
    "yield.*tokenization.*flow": "yield-tokenization-flow",
    "repayment.*processing": "repayment-processing-flow",
    "pooling.*mechanism": "pooling-mechanism-flow",
    "default.*handling": "default-handling-flow",
    "storage.*layout": "storage-layout-diagram",
    "analytics.*architecture": "analytics-architecture",
    "subgraph.*entity": "subgraph-entity-relationships",
    "governance.*architecture": "governance-architecture",
    "governance.*proposal": "governance-proposal-flow",
    "secondary.*market": "secondary-market-architecture",
    "transfer.*restriction": "transfer-restriction-flow",
    "kyc.*architecture": "kyc-architecture",
    "kyc.*workflow": "kyc-workflow",
    "entity.*relationship|er.*diagram": "er-diagram",
    "erc.*1155.*architecture": "erc1155-architecture",
    "erc.*20.*token": "erc20-token-flow",
    "property.*nft": "property-nft-architecture",
    "share.*transfer": "share-transfer-architecture",
    "use.*case": "use-case-diagram",
    "mobile.*workflow": "mobile-workflow",
    
    # Wireframes
    "wireframe.*dashboard": "wireframe-dashboard",
    "wireframe.*property.*registration": "wireframe-property-registration",
    "wireframe.*yield.*agreement": "wireframe-yield-agreement",
    "wireframe.*mobile.*dashboard": "wireframe-mobile-dashboard",
    "wireframe.*mobile.*yield": "wireframe-mobile-yield-agreement",
    
    # Load testing
    "load.*test.*system": "load-test-system-architecture",
    "load.*test.*workflow": "load-test-workflow",
    "load.*testing.*results": "load-testing-results",
    "load.*test.*gas.*scaling": "load-test-gas-scaling",
    "load.*test.*shareholder": "load-test-shareholder-distribution",
    "load.*test.*success": "load-test-success-rates",
    "load.*test.*latency": "load-test-latency-distribution",
    "load.*test.*metrics": "load-test-metrics-summary",
    "load.*test.*cost": "load-test-cost-comparison",
    "load.*test.*erc1155": "load-test-erc1155-efficiency",
    "load.*test.*token.*recycling": "load-test-token-recycling",
    "load.*test.*restriction.*enforcement": "load-test-restriction-enforcement",
    "load.*test.*restriction.*overhead": "load-test-restriction-overhead",
    "load.*test.*rq": "load-test-rq-validation",
    
    # Survey/Interview
    "demographic.*overview|survey.*demographics": "fig_demographic_overview",
    "tokenisation.*interest": "fig_tokenisation_interest_analysis",
    "correlation.*matrix|spearman": "fig_correlation_matrix",
    "cluster.*analysis": "fig_cluster_analysis",
    "motivations.*concerns": "fig_motivations_concerns",
    "feature.*importance|platform.*feature": "fig_feature_importance",
    "interview.*demographics": "fig_interview_demographics",
    "thematic.*code": "fig_thematic_code_frequencies",
    "landlord.*fintech|landlord.*expert": "fig_landlord_fintech_comparison",
    "likert.*distribution": "fig_likert_distributions",
}

def find_image_for_placeholder(placeholder_text):
    """Find matching image file for a placeholder description."""
    text_lower = placeholder_text.lower()
    
    # Try keyword mappings first
    for pattern, filename in keyword_mappings.items():
        if re.search(pattern, text_lower):
            # Check all directories
            if filename in drawio_files:
                return drawio_files[filename]
            if filename in survey_files:
                return survey_files[filename]
            if filename in chart_files:
                return chart_files[filename]
    
    # Direct filename match
    for files in [drawio_files, survey_files, chart_files]:
        for name, path in files.items():
            if name.replace("-", " ") in text_lower or name.replace("-", "") in text_lower.replace(" ", ""):
                return path
    
    return None

# Find and replace placeholders
placeholder_pattern = re.compile(r'\[PLACEHOLDER:?\s*(Figure|Table|Diagram|Chart|Screenshot)?\s*[-–]?\s*(.+?)\]', re.IGNORECASE)

figures_inserted = 0
figures_not_found = []

print("\nProcessing document paragraphs...")

for i, para in enumerate(doc.paragraphs):
    text = para.text
    match = placeholder_pattern.search(text)
    
    if match:
        placeholder_type = match.group(1) or "Figure"
        description = match.group(2).strip()
        
        # Find matching image
        image_path = find_image_for_placeholder(description)
        
        if image_path and image_path.exists():
            # Clear paragraph and insert image
            para.clear()
            run = para.add_run()
            
            # Insert image (5 inches wide for most diagrams)
            run.add_picture(str(image_path), width=Inches(5.5))
            
            # Add caption below
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(f"{placeholder_type} {i}: {description}")
            caption_run.italic = True
            caption_run.font.size = Pt(10)
            
            figures_inserted += 1
            print(f"  ✓ Inserted: {description[:50]}... -> {image_path.name}")
        else:
            figures_not_found.append(description)
            print(f"  ✗ Not found: {description[:60]}...")

# Save document
print(f"\nSaving to {output_path}...")
doc.save(output_path)

print(f"\n{'='*60}")
print(f"SUMMARY")
print(f"{'='*60}")
print(f"Figures inserted: {figures_inserted}")
print(f"Figures not found: {len(figures_not_found)}")

if figures_not_found:
    print(f"\nMissing figures ({len(figures_not_found)}):")
    for desc in figures_not_found[:20]:
        print(f"  - {desc[:70]}")
    if len(figures_not_found) > 20:
        print(f"  ... and {len(figures_not_found) - 20} more")

print(f"\n✅ Document saved: {output_path}")
