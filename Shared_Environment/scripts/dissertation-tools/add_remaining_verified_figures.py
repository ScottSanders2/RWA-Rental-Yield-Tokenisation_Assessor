#!/usr/bin/env python3
"""
Add remaining verified figures that don't have placeholders in the document
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from pathlib import Path

doc = Document("DissertationProgressFinal_VERIFIED.docx")
output_path = "DissertationProgressFinal_VERIFIED_COMPLETE.docx"

# Remaining figures to add (Section 6 charts and Section 7 survey figures)
remaining_figures = {
    # Section 6 Charts (6.1-6.10)
    "6.1": ("generated_charts/fig_6_1_token_standard_gas_comparison.png", "Token Standard Gas Comparison"),
    "6.2": ("generated_charts/fig_6_2_batch_operation_scaling.png", "Batch Operation Scaling Curve"),
    "6.3": ("generated_charts/fig_6_3_amoy_anvil_variance_heatmap.png", "Amoy vs Anvil Variance Heatmap"),
    "6.4": ("generated_charts/fig_6_4_volatile_simulation_radar.png", "Volatile Simulation Recovery Radar"),
    "6.5": ("generated_charts/fig_6_5_diamond_overhead_boxplot.png", "Diamond Architecture Call Overhead"),
    "6.6": ("generated_charts/fig_6_6_load_testing_scatter.png", "Load Testing Throughput Scatter"),
    "6.7": ("generated_charts/fig_6_7_gas_cost_projections.png", "Gas Cost Projections"),
    "6.8": ("generated_charts/fig_6_8_test_pass_rate_evolution.png", "Test Pass Rate Evolution"),
    "6.9": ("generated_charts/architecture-comparison.png", "Architecture Comparison"),
    "6.10": ("generated_charts/user-workflow-diagrams.png", "User Workflow Diagrams"),
    
    # Section 7 Survey/Interview figures
    "7.1": ("survey_interview_charts/fig_demographic_overview.png", "Survey Demographics Overview"),
    "7.2": ("survey_interview_charts/fig_tokenisation_interest_analysis.png", "Tokenisation Interest Analysis"),
    "7.3": ("survey_interview_charts/fig_correlation_matrix.png", "Spearman Correlation Matrix"),
    "7.4": ("survey_interview_charts/fig_interview_demographics.png", "Interview Demographics"),
    "7.5": ("survey_interview_charts/fig_thematic_code_frequencies.png", "Theme Frequency Analysis"),
    "7.6": ("generated_charts/token-decision-framework.png", "Token Standard Decision Framework"),
    "7.7": ("generated_charts/fig_7_7_volatile_recovery_comparison.png", "Volatile Market Recovery"),
    "7.8": ("generated_charts/fig_7_8_testnet_local_comparison.png", "Testnet vs Local Performance"),
}

# Find Section 6 and add charts
print("Looking for Section 6 and 7 to add remaining figures...")

# Search for figure references and add images after them
for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    for fig_num, (img_path, caption) in remaining_figures.items():
        # Look for figure reference
        if f"Figure {fig_num}:" in text or f"Figure {fig_num} " in text:
            # Check if image already inserted (paragraph would be empty with picture)
            if len(para.runs) > 0 and not any(run.text for run in para.runs):
                continue  # Already has image
            
            # Check if next paragraph has the image
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                # If next para is empty or has placeholder, it might need image
                
            # Add image after this paragraph
            img_file = Path(img_path)
            if img_file.exists():
                new_para = doc.add_paragraph()
                run = new_para.add_run()
                run.add_picture(str(img_file), width=Inches(5.5))
                print(f"  ✓ Added Figure {fig_num}: {caption}")

# Save
doc.save(output_path)
print(f"\n✅ Document saved: {output_path}")
print(f"   File size: {os.path.getsize(output_path) / 1024 / 1024:.1f} MB")
