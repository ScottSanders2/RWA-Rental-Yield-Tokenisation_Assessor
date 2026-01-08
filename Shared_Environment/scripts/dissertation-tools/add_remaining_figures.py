#!/usr/bin/env python3
"""
Add remaining 25 figures to the complete document
"""

from docx import Document
from docx.shared import Inches, Pt
import os
from pathlib import Path

# Load the complete document
doc_path = "DissertationProgressFinal_COMPLETE.docx"
output_path = "DissertationProgressFinal_FINAL.docx"

print(f"Loading {doc_path}...")
doc = Document(doc_path)

# Image directories
drawio_dir = Path("rendered_diagrams_drawio")
survey_dir = Path("survey_interview_charts")
charts_dir = Path("generated_charts")
screenshots_dir = Path("generated_screenshots")

# Remaining images to add
remaining_images = {
    # Survey/Interview (2 remaining)
    "fig_likert_distributions": survey_dir / "fig_likert_distributions.png",
    "fig_landlord_fintech_comparison": survey_dir / "fig_landlord_fintech_comparison.png",
    
    # Generated charts (10)
    "fig_6_1_token_standard_gas_comparison": charts_dir / "fig_6_1_token_standard_gas_comparison.png",
    "fig_6_2_batch_operation_scaling": charts_dir / "fig_6_2_batch_operation_scaling.png",
    "fig_6_3_amoy_anvil_variance_heatmap": charts_dir / "fig_6_3_amoy_anvil_variance_heatmap.png",
    "fig_6_4_volatile_simulation_radar": charts_dir / "fig_6_4_volatile_simulation_radar.png",
    "fig_6_5_diamond_architecture_overhead": charts_dir / "fig_6_5_diamond_architecture_overhead.png",
    "fig_6_6_load_testing_scatter": charts_dir / "fig_6_6_load_testing_scatter.png",
    "fig_6_7_gas_cost_projections": charts_dir / "fig_6_7_gas_cost_projections.png",
    "fig_6_8_test_pass_rate_evolution": charts_dir / "fig_6_8_test_pass_rate_evolution.png",
    "fig_7_6_token_standard_decision_framework": charts_dir / "fig_7_6_token_standard_decision_framework.png",
    "fig_7_7_volatile_recovery_comparison": charts_dir / "fig_7_7_volatile_recovery_comparison.png",
    
    # Screenshots (13 remaining)
    "iteration-1-test-results": screenshots_dir / "iteration-1-test-results.png",
    "iteration-3-test-results": screenshots_dir / "iteration-3-test-results.png",
    "iteration-10-test-results": screenshots_dir / "iteration-10-test-results.png",
    "iteration-11-test-results": screenshots_dir / "iteration-11-test-results.png",
    "iteration-14-test-results": screenshots_dir / "iteration-14-test-results.png",
    "iteration-15-test-results": screenshots_dir / "iteration-15-test-results.png",
    "network-isolation-setup": screenshots_dir / "network-isolation-setup.png",
    "analytics-dashboard": screenshots_dir / "analytics-dashboard.png",
    "load-testing-results-detailed": screenshots_dir / "load-testing-results-detailed.png",
    "diamond-pattern-migration": screenshots_dir / "diamond-pattern-migration.png",
    "amoy-testnet-deployment": screenshots_dir / "amoy-testnet-deployment.png",
    "backend-api-test-results": screenshots_dir / "backend-api-test-results.png",
    "frontend-build-output": screenshots_dir / "frontend-build-output.png",
}

# Add Section 6 charts appendix
doc.add_page_break()
doc.add_heading('Appendix E: Testing and Evaluation Charts', level=1)

print("\n=== Adding Section 6 Testing Charts ===")
section6_charts = [
    ("fig_6_1_token_standard_gas_comparison", "Token Standard Gas Comparison"),
    ("fig_6_2_batch_operation_scaling", "Batch Operation Scaling Analysis"),
    ("fig_6_3_amoy_anvil_variance_heatmap", "Amoy vs Anvil Variance Heatmap"),
    ("fig_6_4_volatile_simulation_radar", "Volatile Market Simulation Radar"),
    ("fig_6_5_diamond_architecture_overhead", "Diamond Architecture Overhead"),
    ("fig_6_6_load_testing_scatter", "Load Testing Performance Scatter"),
    ("fig_6_7_gas_cost_projections", "Gas Cost Projections"),
    ("fig_6_8_test_pass_rate_evolution", "Test Pass Rate Evolution"),
]

chart_count = 0
for img_name, caption in section6_charts:
    img_path = charts_dir / f"{img_name}.png"
    if img_path.exists():
        # Add caption
        caption_para = doc.add_paragraph()
        caption_para.add_run(f"Figure E.{chart_count + 1}: {caption}").bold = True
        
        # Add image
        img_para = doc.add_paragraph()
        run = img_para.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        
        chart_count += 1
        print(f"  ✓ Figure E.{chart_count}: {caption}")

print(f"\nAdded {chart_count} Section 6 charts")

# Add Section 7 additional charts
doc.add_heading('Appendix F: Evaluation Charts', level=1)

print("\n=== Adding Section 7 Evaluation Charts ===")
section7_charts = [
    ("fig_7_6_token_standard_decision_framework", "Token Standard Decision Framework"),
    ("fig_7_7_volatile_recovery_comparison", "Volatile Market Recovery Comparison"),
    ("fig_likert_distributions", "Likert Scale Distributions"),
    ("fig_landlord_fintech_comparison", "Landlord vs FinTech Expert Comparison"),
]

eval_count = 0
for img_name, caption in section7_charts:
    # Check multiple directories
    img_path = None
    if (charts_dir / f"{img_name}.png").exists():
        img_path = charts_dir / f"{img_name}.png"
    elif (survey_dir / f"{img_name}.png").exists():
        img_path = survey_dir / f"{img_name}.png"
    
    if img_path and img_path.exists():
        # Add caption
        caption_para = doc.add_paragraph()
        caption_para.add_run(f"Figure F.{eval_count + 1}: {caption}").bold = True
        
        # Add image
        img_para = doc.add_paragraph()
        run = img_para.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        
        eval_count += 1
        print(f"  ✓ Figure F.{eval_count}: {caption}")

print(f"\nAdded {eval_count} Section 7 evaluation charts")

# Add remaining screenshots appendix
doc.add_heading('Appendix G: Additional Implementation Screenshots', level=1)

print("\n=== Adding Remaining Screenshots ===")
remaining_screenshots = [
    ("iteration-1-test-results", "Iteration 1 Test Results"),
    ("iteration-3-test-results", "Iteration 3 Test Results"),
    ("iteration-10-test-results", "Iteration 10 Test Results"),
    ("iteration-11-test-results", "Iteration 11 Test Results"),
    ("iteration-14-test-results", "Iteration 14 Test Results"),
    ("iteration-15-test-results", "Iteration 15 Test Results"),
    ("network-isolation-setup", "Docker Network Isolation Setup"),
    ("analytics-dashboard", "Analytics Dashboard"),
    ("load-testing-results-detailed", "Load Testing Detailed Results"),
    ("diamond-pattern-migration", "Diamond Pattern Migration"),
    ("amoy-testnet-deployment", "Polygon Amoy Testnet Deployment"),
    ("backend-api-test-results", "Backend API Test Results"),
    ("frontend-build-output", "Frontend Build Output"),
]

screenshot_count = 0
for img_name, caption in remaining_screenshots:
    img_path = screenshots_dir / f"{img_name}.png"
    if img_path.exists():
        # Add caption
        caption_para = doc.add_paragraph()
        caption_para.add_run(f"Figure G.{screenshot_count + 1}: {caption}").bold = True
        
        # Add image
        img_para = doc.add_paragraph()
        run = img_para.add_run()
        run.add_picture(str(img_path), width=Inches(5.5))
        
        screenshot_count += 1
        print(f"  ✓ Figure G.{screenshot_count}: {caption}")

print(f"\nAdded {screenshot_count} remaining screenshots")

# Save document
print(f"\nSaving to {output_path}...")
doc.save(output_path)

total_added = chart_count + eval_count + screenshot_count
print(f"\n{'='*60}")
print(f"FINAL SUMMARY")
print(f"{'='*60}")
print(f"Section 6 charts (Appendix E):    {chart_count}")
print(f"Section 7 charts (Appendix F):    {eval_count}")
print(f"Screenshots (Appendix G):         {screenshot_count}")
print(f"{'='*60}")
print(f"ADDITIONAL FIGURES ADDED:         {total_added}")
print(f"PREVIOUS TOTAL:                   83")
print(f"NEW TOTAL:                        {83 + total_added}")
print(f"\n✅ Document saved: {output_path}")
print(f"   File size: {os.path.getsize(output_path) / 1024 / 1024:.1f} MB")
