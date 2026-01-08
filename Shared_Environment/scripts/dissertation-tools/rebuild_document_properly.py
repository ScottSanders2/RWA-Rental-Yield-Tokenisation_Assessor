#!/usr/bin/env python3
"""
Rebuild DissertationProgressFinal.docx with PROPER figure captions.
Each image MUST have its Figure X.X caption directly associated with it.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re
import os
import json

# Define the complete figure mapping - Figure number -> image file
FIGURE_MAPPING = {
    # Section 4 - System Design (32 figures: 4.1-4.32)
    "Figure 4.1": "rendered_diagrams_drawio/docker-architecture-logical.png",
    "Figure 4.2": "rendered_diagrams_drawio/docker-architecture-physical.png",
    "Figure 4.3": "rendered_diagrams_drawio/monitoring-architecture.png",
    "Figure 4.4": "rendered_diagrams_drawio/mobile-architecture.png",
    "Figure 4.5": "rendered_diagrams_drawio/backend-api-architecture.png",
    "Figure 4.6": "rendered_diagrams_drawio/property-owner-workflow.png",
    "Figure 4.7": "rendered_diagrams_drawio/yield-architecture.png",
    "Figure 4.8": "rendered_diagrams_drawio/yield-tokenization-flow.png",
    "Figure 4.9": "rendered_diagrams_drawio/repayment-processing-flow.png",
    "Figure 4.10": "rendered_diagrams_drawio/pooling-mechanism-flow.png",
    "Figure 4.11": "rendered_diagrams_drawio/default-handling-flow.png",
    "Figure 4.12": "rendered_diagrams_drawio/erc7201-storage-layout.png",
    "Figure 4.13": "rendered_diagrams_drawio/analytics-architecture.png",
    "Figure 4.14": "rendered_diagrams_drawio/subgraph-entity-relationships.png",
    "Figure 4.15": "rendered_diagrams_drawio/governance-architecture.png",
    "Figure 4.16": "rendered_diagrams_drawio/governance-proposal-flow.png",
    "Figure 4.17": "rendered_diagrams_drawio/secondary-market-architecture.png",
    "Figure 4.18": "rendered_diagrams_drawio/transfer-restriction-flow.png",
    "Figure 4.19": "rendered_diagrams_drawio/kyc-architecture.png",
    "Figure 4.20": "rendered_diagrams_drawio/kyc-verification-workflow.png",
    "Figure 4.21": "rendered_diagrams_drawio/er-diagram.png",
    "Figure 4.22": "rendered_diagrams_drawio/erc1155-architecture.png",
    "Figure 4.23": "rendered_diagrams_drawio/erc20-token-flow.png",
    "Figure 4.24": "rendered_diagrams_drawio/property-nft-architecture.png",
    "Figure 4.25": "rendered_diagrams_drawio/share-transfer-architecture.png",
    "Figure 4.26": "rendered_diagrams_drawio/use-case-diagram.png",
    "Figure 4.27": "rendered_diagrams_drawio/mobile-workflow.png",
    "Figure 4.28": "rendered_diagrams_drawio/dashboard-wireframe.png",
    "Figure 4.29": "rendered_diagrams_drawio/property-registration-wireframe.png",
    "Figure 4.30": "rendered_diagrams_drawio/yield-agreement-wireframe.png",
    "Figure 4.31": "rendered_diagrams_drawio/mobile-dashboard-wireframe.png",
    "Figure 4.32": "rendered_diagrams_drawio/mobile-yield-agreement-wireframe.png",
    
    # Section 5 - Implementation (23 figures: 5.1-5.23)
    "Figure 5.1": "generated_screenshots/docker-desktop-dashboard.png",
    "Figure 5.2": "generated_screenshots/grafana-dashboard.png",
    "Figure 5.3": "generated_screenshots/prometheus-targets.png",
    "Figure 5.4": "generated_screenshots/node-exporter-metrics.png",
    "Figure 5.5": "generated_screenshots/multi-environment-orchestration.png",
    "Figure 5.6": "generated_screenshots/backup-directory.png",
    "Figure 5.7": "generated_screenshots/foundry-project-structure.png",
    "Figure 5.8": "generated_screenshots/forge-test-output.png",
    "Figure 5.9": "generated_screenshots/gas-report.png",
    "Figure 5.10": "generated_screenshots/anvil-container-logs.png",
    "Figure 5.11": "generated_screenshots/deployment-script-output.png",
    "Figure 5.12": "generated_screenshots/erc7201-storage-layout.png",
    "Figure 5.13": "generated_screenshots/bytecode-size-verification.png",
    "Figure 5.14": "generated_screenshots/amoy-testnet-deployment.png",
    "Figure 5.15": "generated_screenshots/polygonscan-transaction.png",
    "Figure 5.16": "generated_screenshots/diamond-test-results.png",
    "Figure 5.17": "generated_screenshots/property-registration-form.png",
    "Figure 5.18": "generated_screenshots/yield-agreement-creation.png",
    "Figure 5.19": "generated_screenshots/analytics-dashboard.png",
    "Figure 5.20": "generated_screenshots/simulation-test-results.png",
    "Figure 5.21": "generated_screenshots/gas-comparison-report.png",
    "Figure 5.22": "generated_screenshots/variance-tracking-metrics.png",
    "Figure 5.23": "rendered_diagrams_drawio/yield-tokenization-flow.png",  # Same as 4.8
    
    # Section 6 - Testing (24 figures: 6.1-6.24)
    "Figure 6.1": "generated_charts/token-standard-gas-comparison.png",
    "Figure 6.2": "generated_charts/batch-operation-scaling.png",
    "Figure 6.3": "generated_charts/amoy-anvil-variance-heatmap.png",
    "Figure 6.4": "generated_charts/volatile-simulation-recovery.png",
    "Figure 6.5": "generated_charts/diamond-call-overhead.png",
    "Figure 6.6": "generated_charts/load-testing-throughput.png",
    "Figure 6.7": "generated_charts/gas-cost-projections.png",
    "Figure 6.8": "generated_charts/test-pass-rate-evolution.png",
    "Figure 6.9": "rendered_diagrams_drawio/architecture-comparison.png",
    "Figure 6.10": "rendered_diagrams_drawio/user-workflow-diagrams.png",
    "Figure 6.11": "rendered_diagrams_drawio/load-test-architecture.png",
    "Figure 6.12": "rendered_diagrams_drawio/load-test-workflow.png",
    "Figure 6.13": "rendered_diagrams_drawio/load-test-results-summary.png",
    "Figure 6.14": "rendered_diagrams_drawio/load-test-gas-scaling.png",
    "Figure 6.15": "rendered_diagrams_drawio/load-test-shareholder-distribution.png",
    "Figure 6.16": "rendered_diagrams_drawio/load-test-success-rates.png",
    "Figure 6.17": "rendered_diagrams_drawio/load-test-latency-distribution.png",
    "Figure 6.18": "rendered_diagrams_drawio/load-test-metrics-summary.png",
    "Figure 6.19": "rendered_diagrams_drawio/load-test-cost-comparison.png",
    "Figure 6.20": "rendered_diagrams_drawio/load-test-erc1155-efficiency.png",
    "Figure 6.21": "rendered_diagrams_drawio/load-test-token-recycling.png",
    "Figure 6.22": "rendered_diagrams_drawio/load-test-restriction-enforcement.png",
    "Figure 6.23": "rendered_diagrams_drawio/load-test-restriction-overhead.png",
    "Figure 6.24": "rendered_diagrams_drawio/load-test-research-validation.png",
    
    # Section 7 - Evaluation (8 figures: 7.1-7.8)
    "Figure 7.1": "survey_interview_charts/survey_demographics_composite.png",
    "Figure 7.2": "survey_interview_charts/tokenisation_interest_distribution.png",
    "Figure 7.3": "survey_interview_charts/spearman_correlation_matrix.png",
    "Figure 7.4": "survey_interview_charts/interview_demographics_composite.png",
    "Figure 7.5": "survey_interview_charts/theme_frequency_analysis.png",
    "Figure 7.6": "rendered_diagrams_drawio/token-standard-decision-framework.png",
    "Figure 7.7": "generated_charts/volatile-simulation-recovery.png",
    "Figure 7.8": "generated_charts/testnet-local-performance.png",
}

# Appendix D wireframes (19 figures: D.1-D.19)
APPENDIX_D_FIGURES = {
    "Figure D.1": "rendered_diagrams_drawio/analytics-dashboard-wireframe.png",
    "Figure D.2": "rendered_diagrams_drawio/compliance-dashboard-wireframe.png",
    "Figure D.3": "rendered_diagrams_drawio/error-states-wireframe.png",
    "Figure D.4": "rendered_diagrams_drawio/governance-dashboard-wireframe.png",
    "Figure D.5": "rendered_diagrams_drawio/investor-dashboard-wireframe.png",
    "Figure D.6": "rendered_diagrams_drawio/investor-portfolio-wireframe.png",
    "Figure D.7": "rendered_diagrams_drawio/kyc-verification-wireframe.png",
    "Figure D.8": "rendered_diagrams_drawio/loading-states-wireframe.png",
    "Figure D.9": "rendered_diagrams_drawio/mobile-analytics-wireframe.png",
    "Figure D.10": "rendered_diagrams_drawio/mobile-investor-wireframe.png",
    "Figure D.11": "rendered_diagrams_drawio/mobile-notifications-wireframe.png",
    "Figure D.12": "rendered_diagrams_drawio/mobile-property-details-wireframe.png",
    "Figure D.13": "rendered_diagrams_drawio/mobile-secondary-market-wireframe.png",
    "Figure D.14": "rendered_diagrams_drawio/notification-center-wireframe.png",
    "Figure D.15": "rendered_diagrams_drawio/property-details-wireframe.png",
    "Figure D.16": "rendered_diagrams_drawio/secondary-market-wireframe.png",
    "Figure D.17": "rendered_diagrams_drawio/settings-wireframe.png",
    "Figure D.18": "rendered_diagrams_drawio/transaction-history-wireframe.png",
    "Figure D.19": "rendered_diagrams_drawio/yield-distribution-wireframe.png",
}

# Appendix E charts (5 figures: E.1-E.5)
APPENDIX_E_FIGURES = {
    "Figure E.1": "survey_interview_charts/respondent_cluster_analysis.png",
    "Figure E.2": "survey_interview_charts/platform_feature_importance.png",
    "Figure E.3": "survey_interview_charts/investment_motivations_concerns.png",
    "Figure E.4": "survey_interview_charts/landlord_fintech_comparison.png",
    "Figure E.5": "survey_interview_charts/likert_scale_distributions.png",
}

def get_figure_title(figure_num, md_content):
    """Extract the full figure title from markdown"""
    # Pattern: **Figure X.X: Title**
    pattern = rf'\*\*{re.escape(figure_num)}:\s*([^*]+)\*\*'
    match = re.search(pattern, md_content)
    if match:
        return f"{figure_num}: {match.group(1).strip()}"
    return figure_num

def rebuild_document():
    """Rebuild the document with proper figure captions"""
    
    # Read the markdown to get figure titles
    with open('DissertationProgressFinal.md', 'r') as f:
        md_content = f.read()
    
    # Open the existing document
    doc = Document('DissertationProgressFinal_v2.1.docx')
    
    # Create a Caption style if it doesn't exist
    try:
        caption_style = doc.styles['Caption']
    except KeyError:
        caption_style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        caption_style.font.size = Pt(10)
        caption_style.font.italic = True
    
    # Track what we find
    figures_found = []
    figures_missing_caption = []
    
    # Iterate through paragraphs
    para_idx = 0
    while para_idx < len(doc.paragraphs):
        para = doc.paragraphs[para_idx]
        
        # Check if this paragraph contains an image
        has_image = False
        for run in para.runs:
            if run._element.xpath('.//a:blip'):
                has_image = True
                break
        
        if has_image:
            # Check if previous paragraph has a Figure label
            prev_text = ""
            if para_idx > 0:
                prev_text = doc.paragraphs[para_idx - 1].text.strip()
            
            # Check if next paragraph has a Figure label
            next_text = ""
            if para_idx < len(doc.paragraphs) - 1:
                next_text = doc.paragraphs[para_idx + 1].text.strip()
            
            # Look for Figure X.X pattern
            figure_match = None
            if re.search(r'Figure\s+[0-9]+\.[0-9]+', prev_text):
                figure_match = re.search(r'(Figure\s+[0-9]+\.[0-9]+[^*\n]*)', prev_text)
            elif re.search(r'Figure\s+[0-9]+\.[0-9]+', next_text):
                figure_match = re.search(r'(Figure\s+[0-9]+\.[0-9]+[^*\n]*)', next_text)
            
            if figure_match:
                figures_found.append(figure_match.group(1))
            else:
                figures_missing_caption.append(f"Para {para_idx}")
        
        para_idx += 1
    
    print(f"Figures with captions: {len(figures_found)}")
    print(f"Figures missing captions: {len(figures_missing_caption)}")
    
    return figures_found, figures_missing_caption

if __name__ == "__main__":
    found, missing = rebuild_document()
    print("\n=== Figures Found ===")
    for f in found[:20]:
        print(f"  {f}")
    print("\n=== Missing Caption Locations ===")
    for m in missing[:20]:
        print(f"  {m}")
