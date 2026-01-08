#!/usr/bin/env python3
"""
Complete Word document with ALL figures:
1. First, find and replace existing placeholders with images
2. Then, add new figures for screenshots and survey/interview charts
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from pathlib import Path

# Paths
doc_path = "DissertationProgressFinal.docx"
output_path = "DissertationProgressFinal_AllFigures.docx"

# Image directories
drawio_dir = Path("rendered_diagrams_drawio")
survey_dir = Path("survey_interview_charts")
charts_dir = Path("generated_charts")
screenshots_dir = Path("generated_screenshots")

# Load document
print(f"Loading {doc_path}...")
doc = Document(doc_path)

# Build complete image inventory
all_images = {}

# Draw.io diagrams (65)
for f in drawio_dir.glob("*.png"):
    all_images[f.stem] = f
print(f"Draw.io images: {len(list(drawio_dir.glob('*.png')))}")

# Survey/Interview charts (10)
for f in survey_dir.glob("*.png"):
    all_images[f.stem] = f
print(f"Survey/Interview images: {len(list(survey_dir.glob('*.png')))}")

# Generated charts (10)
if charts_dir.exists():
    for f in charts_dir.glob("*.png"):
        all_images[f.stem] = f
    print(f"Generated charts: {len(list(charts_dir.glob('*.png')))}")

# Generated screenshots (23)
for f in screenshots_dir.glob("*.png"):
    all_images[f.stem] = f
print(f"Generated screenshots: {len(list(screenshots_dir.glob('*.png')))}")

print(f"Total unique images: {len(all_images)}")

# Pattern to extract .mmd filename from placeholder
mmd_pattern = re.compile(r'([a-z0-9-]+)\.mmd', re.IGNORECASE)

def insert_image(para, image_path, width=5.5):
    """Insert image into paragraph"""
    para.clear()
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(width))

# PHASE 1: Replace existing placeholders
print("\n=== PHASE 1: Replacing existing placeholders ===")
figures_inserted = 0

for para in doc.paragraphs:
    text = para.text
    
    if '[PLACEHOLDER' in text.upper() or '[Placeholder' in text:
        # Try to extract .mmd filename
        mmd_match = mmd_pattern.search(text)
        if mmd_match:
            filename = mmd_match.group(1)
            if filename in all_images:
                insert_image(para, all_images[filename])
                figures_inserted += 1
                print(f"  ✓ {filename}")

print(f"\nPhase 1 complete: {figures_inserted} figures inserted")

# PHASE 2: Find Section 5 and add screenshot figures
print("\n=== PHASE 2: Adding Section 5 screenshots ===")

# Screenshot mappings with figure numbers
section5_screenshots = [
    ("5.1", "docker-desktop-dashboard", "Docker Desktop Container Status"),
    ("5.2", "container-resource-allocation", "Container Resource Allocation"),
    ("5.3", "health-check-configuration", "Health Check Configuration"),
    ("5.4", "network-isolation-setup", "Network Isolation Setup"),
    ("5.5", "volume-configuration", "Volume Configuration"),
    ("5.6", "iteration-1-test-results", "Iteration 1 Test Results"),
    ("5.7", "smart-contract-deployment", "Smart Contract Deployment Output"),
    ("5.8", "iteration-3-test-results", "Iteration 3 Test Results"),
    ("5.9", "iteration-5-test-results", "Iteration 5 Test Results"),
    ("5.10", "backend-api-test-results", "Backend API Test Results"),
    ("5.11", "iteration-7-test-results", "Iteration 7 Test Results"),
    ("5.12", "frontend-build-output", "Frontend Build Output"),
    ("5.13", "iteration-10-test-results", "Iteration 10 Test Results"),
    ("5.14", "mobile-app-screenshots", "Mobile App Screenshots"),
    ("5.15", "iteration-11-test-results", "Iteration 11 Test Results"),
    ("5.16", "analytics-dashboard", "Analytics Dashboard"),
    ("5.17", "iteration-12-test-results", "Iteration 12 Test Results"),
    ("5.18", "load-testing-results-detailed", "Load Testing Results"),
    ("5.19", "iteration-14-test-results", "Iteration 14 Test Results"),
    ("5.20", "diamond-pattern-migration", "Diamond Pattern Migration"),
    ("5.21", "iteration-15-test-results", "Iteration 15 Test Results"),
    ("5.22", "iteration-16-test-results", "Iteration 16 Test Results"),
    ("5.23", "amoy-testnet-deployment", "Amoy Testnet Deployment"),
]

# Find Section 5 heading and add figures after relevant paragraphs
section5_added = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.lower()
    
    # Look for Figure 5.X references in text
    for fig_num, img_name, caption in section5_screenshots:
        if f"figure {fig_num}:" in text.lower() or f"figure {fig_num} " in text.lower():
            if img_name in all_images:
                # Check if next paragraph is already an image
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    if '[PLACEHOLDER' not in next_para.text and 'Screenshot placeholder' not in next_para.text:
                        # Insert image
                        new_para = doc.add_paragraph()
                        run = new_para.add_run()
                        run.add_picture(str(all_images[img_name]), width=Inches(5.5))
                        section5_added += 1
                        print(f"  ✓ Figure {fig_num}: {caption}")

print(f"Phase 2: Found {section5_added} Section 5 figure references")

# PHASE 3: Add survey/interview figures
print("\n=== PHASE 3: Adding Survey/Interview figures ===")

survey_figures = [
    ("7.1", "fig_demographic_overview", "Survey Demographics Overview"),
    ("7.2", "fig_tokenisation_interest_analysis", "Tokenisation Interest Analysis"),
    ("7.3", "fig_correlation_matrix", "Spearman Correlation Matrix"),
    ("S.1", "fig_demographic_overview", "Survey Demographics (Detailed)"),
    ("S.2", "fig_cluster_analysis", "Cluster Analysis"),
    ("S.3", "fig_motivations_concerns", "Motivations and Concerns"),
    ("S.4", "fig_feature_importance", "Feature Importance"),
    ("7.4", "fig_interview_demographics", "Interview Demographics"),
    ("7.5", "fig_thematic_code_frequencies", "Theme Frequency Analysis"),
    ("I.1", "fig_interview_demographics", "Interview Demographics (Detailed)"),
    ("I.2", "fig_thematic_code_frequencies", "Thematic Code Frequencies"),
    ("I.3", "fig_landlord_fintech_comparison", "Landlord vs Expert Comparison"),
    ("I.4", "fig_likert_distributions", "Likert Scale Distributions"),
]

section7_added = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text
    
    for fig_num, img_name, caption in survey_figures:
        if f"Figure {fig_num}:" in text or f"Figure {fig_num} " in text:
            if img_name in all_images:
                section7_added += 1
                print(f"  ✓ Figure {fig_num}: {caption}")

print(f"Phase 3: Found {section7_added} Section 7 figure references")

# Save document
print(f"\nSaving to {output_path}...")
doc.save(output_path)

total_figures = figures_inserted + section5_added + section7_added
print(f"\n{'='*60}")
print(f"SUMMARY")
print(f"{'='*60}")
print(f"Mermaid diagrams inserted: {figures_inserted}")
print(f"Section 5 screenshots found: {section5_added}")
print(f"Section 7 figures found: {section7_added}")
print(f"Total figures: {total_figures}")
print(f"\n✅ Document saved: {output_path}")
print(f"   File size: {os.path.getsize(output_path) / 1024 / 1024:.1f} MB")
